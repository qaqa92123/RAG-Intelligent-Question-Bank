import shutil
import json
import csv
import random
import hashlib
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple
from fastapi import FastAPI, Request, File, UploadFile, Form, Cookie, Response
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse, JSONResponse, RedirectResponse
from fastapi.templating import Jinja2Templates
from fastapi.middleware.cors import CORSMiddleware
from langchain.chains.retrieval_qa.base import RetrievalQA
from langchain_core.prompts import PromptTemplate
from langchain_openai import ChatOpenAI
from pymilvus import Collection, connections
from pymilvus.orm import utility
import argparse
import os
import auth_db
from config import RagConfig
from webapp.auth_agent_routes import register_auth_and_agent_routes
os.environ['NLTK_DATA'] = os.path.expanduser("~/nltk_data")
os.environ['UNSTRUCTURED_SKIP_DOWNLOAD'] = '1'
from collections import Counter, defaultdict
from docx import Document
from openpyxl import load_workbook
import sqlite3
import time
from pathlib import Path
import re
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import base64

QUESTION_DB_PATH = 'question_bank.db'
LEGACY_STATS_DB_PATH = 'question_statistics.db'
QUESTION_XLSX_FILES = [
    ('./单选题.xlsx', '单选'),
    ('./多选题.xlsx', '多选'),
]
KP_XLSX_PATH = './知识点列表.xlsx'
_QUESTION_BANK_READY = False
SHORT_ANSWER_CONTENT_FORMAT = 'short_answer_compound'

SHORT_ANSWER_MODULES = [
    {
        'slug': 'database-big-question-type-1',
        'name': '数据库大题题型一',
        'score': '20分',
        'question_total': '3小题',
        'blank_total': '动态生成',
        'choice_total': '0题',
        'question_type': '数据库设计简答与填空',
        'summary': '基于 E-R 图生成数据库系统分析、概念设计补全与物理设计填空题。',
    }
]


def get_short_answer_module(module_slug: str) -> Optional[Dict[str, Any]]:
    for module in SHORT_ANSWER_MODULES:
        if module.get('slug') == module_slug:
            return module
    return None


def _safe_str(value) -> str:
    return str(value).strip() if value is not None else ''


def _read_questions_from_xlsx_for_import(xlsx_path: str, fallback_qtype: str) -> List[Dict[str, Any]]:
    if not os.path.exists(xlsx_path):
        return []

    wb = load_workbook(xlsx_path, data_only=True)
    sheet = wb.worksheets[0]
    header_row = next(sheet.iter_rows(min_row=1, max_row=1, values_only=True), None)
    headers = [_safe_str(h) for h in (header_row or [])]

    def find_idx(names, default_idx=None):
        for name in names:
            if name in headers:
                return headers.index(name)
        return default_idx

    idx_sn = find_idx(['目录', '序号', '编号'], 0)
    idx_qtype = find_idx(['小题题型', '题目类型', '小题类型'], 1)
    idx_big_stem = find_idx(['大题题干', '大题干'], 2)
    idx_stem = find_idx(['小题题干', '题干'], 4)
    idx_answer = find_idx(['正确答案', '答案'], 5)
    idx_analysis = find_idx(['答案解析', '解析'], 6)
    idx_level = find_idx(['难易度', '难度', '难易'], 7)
    idx_kp = find_idx(['知识点', '知识点标签', '知识点代码'], 8)
    idx_opt_start = find_idx(['选项A'], 10)

    rows = []
    for row in sheet.iter_rows(min_row=2, values_only=True):
        if not any(row):
            continue

        qtype_raw = row[idx_qtype] if idx_qtype is not None and idx_qtype < len(row) else fallback_qtype
        qtype = _safe_str(qtype_raw) or fallback_qtype

        options = []
        for oi in range(idx_opt_start, idx_opt_start + 8):
            if oi < len(row) and row[oi] not in (None, ''):
                options.append(_safe_str(row[oi]))

        rows.append({
            'sn': _safe_str(row[idx_sn]) if idx_sn is not None and idx_sn < len(row) else '',
            'question_type': qtype,
            'big_stem': _safe_str(row[idx_big_stem]) if idx_big_stem is not None and idx_big_stem < len(row) else '',
            'stem': _safe_str(row[idx_stem]) if idx_stem is not None and idx_stem < len(row) else '',
            'answer': _safe_str(row[idx_answer]) if idx_answer is not None and idx_answer < len(row) else '',
            'analysis': _safe_str(row[idx_analysis]) if idx_analysis is not None and idx_analysis < len(row) else '',
            'difficulty': _safe_str(row[idx_level]) if idx_level is not None and idx_level < len(row) else '',
            'knowledge_point_code': _safe_str(row[idx_kp]) if idx_kp is not None and idx_kp < len(row) else '',
            'options': options,
            'source_file': os.path.basename(xlsx_path),
        })

    wb.close()
    return rows


def _import_knowledge_points_from_xlsx(conn: sqlite3.Connection, kp_file_path: str):
    if not os.path.exists(kp_file_path):
        return

    wb = load_workbook(kp_file_path, data_only=True)
    sheet = wb.worksheets[0]
    header_row = next(sheet.iter_rows(min_row=1, max_row=1, values_only=True), None)
    headers = [_safe_str(h) for h in (header_row or [])]

    def find_idx(names):
        for name in names:
            if name in headers:
                return headers.index(name)
        return None

    idx_domain = find_idx(['知识领域名称', '知识领域'])
    idx_unit = find_idx(['知识单元名称', '知识单元'])
    idx_code = find_idx(['知识点代码', '知识点编码'])
    idx_name = find_idx(['知识点名称', '知识点'])

    to_insert = []
    for row in sheet.iter_rows(min_row=2, values_only=True):
        if not any(row):
            continue
        code = _safe_str(row[idx_code]) if idx_code is not None and idx_code < len(row) else ''
        name = _safe_str(row[idx_name]) if idx_name is not None and idx_name < len(row) else ''
        if not code and not name:
            continue
        domain_name = _safe_str(row[idx_domain]) if idx_domain is not None and idx_domain < len(row) else '未分类'
        unit_name = _safe_str(row[idx_unit]) if idx_unit is not None and idx_unit < len(row) else '未分类单元'
        to_insert.append((code, name, domain_name or '未分类', unit_name or '未分类单元'))

    wb.close()

    if to_insert:
        conn.executemany(
            '''
            INSERT INTO knowledge_points (knowledge_point_code, knowledge_point_name, knowledge_domain_name, knowledge_unit_name)
            VALUES (?, ?, ?, ?)
            ON CONFLICT(knowledge_point_code) DO UPDATE SET
                knowledge_point_name=excluded.knowledge_point_name,
                knowledge_domain_name=excluded.knowledge_domain_name,
                knowledge_unit_name=excluded.knowledge_unit_name
            ''',
            to_insert,
        )


def sync_legacy_statistics_db(
    conn: sqlite3.Connection,
    legacy_db_path: str = LEGACY_STATS_DB_PATH,
    source_label: str = 'question_statistics.db'
):
    if not os.path.exists(legacy_db_path):
        return

    try:
        legacy_conn = sqlite3.connect(legacy_db_path)
        legacy_cur = legacy_conn.cursor()
    except Exception as e:
        print(f"打开旧库失败，跳过同步: {e}")
        return

    cur = conn.cursor()

    try:
        legacy_cur.execute(
            '''
            SELECT knowledge_point_code, knowledge_point_name,
                   COALESCE(knowledge_domain_name, '未分类') AS knowledge_domain_name,
                   COALESCE(knowledge_unit_name, '未分类单元') AS knowledge_unit_name
            FROM knowledge_points
            '''
        )
        kp_rows = legacy_cur.fetchall()
    except Exception:
        kp_rows = []

    if kp_rows:
        cur.executemany(
            '''
            INSERT INTO knowledge_points (knowledge_point_code, knowledge_point_name, knowledge_domain_name, knowledge_unit_name)
            VALUES (?, ?, ?, ?)
            ON CONFLICT(knowledge_point_code) DO UPDATE SET
                knowledge_point_name=excluded.knowledge_point_name,
                knowledge_domain_name=excluded.knowledge_domain_name,
                knowledge_unit_name=excluded.knowledge_unit_name
            ''',
            [
                (_safe_str(code), _safe_str(name), _safe_str(domain) or '未分类', _safe_str(unit) or '未分类单元')
                for code, name, domain, unit in kp_rows if _safe_str(code)
            ]
        )

    try:
        legacy_cur.execute(
            '''
            SELECT
                catalog,
                COALESCE(sub_question_type, question_type) AS qtype,
                main_question_stem,
                sub_question_stem,
                correct_answer,
                answer_analysis,
                difficulty,
                knowledge_point_code,
                option_a,
                option_b,
                option_c,
                option_d
            FROM single_choice_questions
            '''
        )
        old_q_rows = legacy_cur.fetchall()
    except Exception:
        old_q_rows = []

    if old_q_rows:
        insert_rows = []
        for row in old_q_rows:
            catalog, qtype, main_stem, sub_stem, answer, analysis, difficulty, kp_code, oa, ob, oc, od = row
            options = [_safe_str(x) for x in [oa, ob, oc, od] if _safe_str(x)]
            qtype_s = _safe_str(qtype) or '单选'
            stem_s = _safe_str(sub_stem)
            kp_s = _safe_str(kp_code)
            unique_key_src = '|'.join([
                qtype_s,
                stem_s,
                json.dumps(options, ensure_ascii=False),
                kp_s,
                source_label,
            ])
            unique_key = hashlib.md5(unique_key_src.encode('utf-8')).hexdigest()
            insert_rows.append((
                _safe_str(catalog),
                qtype_s,
                _safe_str(main_stem),
                stem_s,
                _safe_str(answer),
                _safe_str(analysis),
                _safe_str(difficulty),
                kp_s,
                json.dumps(options, ensure_ascii=False),
                source_label,
                unique_key,
            ))

        cur.executemany(
            '''
            INSERT OR IGNORE INTO question_bank
            (sn, question_type, big_stem, stem, answer, analysis, difficulty, knowledge_point_code, options_json, source_file, unique_key)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''',
            insert_rows,
        )

    try:
        legacy_cur.execute(
            '''
            SELECT knowledge_point_code, knowledge_point_name, question_count
            FROM question_statistics
            '''
        )
        stat_rows = legacy_cur.fetchall()
    except Exception:
        stat_rows = []

    if stat_rows:
        cur.executemany(
            '''
            INSERT INTO legacy_question_statistics
            (knowledge_point_code, knowledge_point_name, question_count, source_db)
            VALUES (?, ?, ?, ?)
            ON CONFLICT(knowledge_point_code, source_db) DO UPDATE SET
                knowledge_point_name=excluded.knowledge_point_name,
                question_count=excluded.question_count,
                synced_at=CURRENT_TIMESTAMP
            ''',
            [
                (
                    _safe_str(code),
                    _safe_str(name),
                    int(cnt) if str(cnt).strip().isdigit() else 0,
                    source_label,
                )
                for code, name, cnt in stat_rows if _safe_str(code)
            ]
        )
        cur.executemany(
            '''
            INSERT INTO knowledge_points (knowledge_point_code, knowledge_point_name, knowledge_domain_name, knowledge_unit_name)
            VALUES (?, ?, '未分类', '未分类单元')
            ON CONFLICT(knowledge_point_code) DO UPDATE SET
                knowledge_point_name=CASE
                    WHEN COALESCE(knowledge_points.knowledge_point_name, '') = '' THEN excluded.knowledge_point_name
                    ELSE knowledge_points.knowledge_point_name
                END
            ''',
            [(_safe_str(code), _safe_str(name)) for code, name, _ in stat_rows if _safe_str(code)]
        )

    conn.commit()
    legacy_conn.close()


def ensure_question_bank_db(db_path: str = QUESTION_DB_PATH):
    global _QUESTION_BANK_READY
    if _QUESTION_BANK_READY and os.path.exists(db_path):
        return

    conn = sqlite3.connect(db_path)
    cur = conn.cursor()
    cur.execute(
        '''
        CREATE TABLE IF NOT EXISTS knowledge_points (
            knowledge_point_code TEXT PRIMARY KEY,
            knowledge_point_name TEXT,
            knowledge_domain_name TEXT,
            knowledge_unit_name TEXT
        )
        '''
    )
    cur.execute(
        '''
        CREATE TABLE IF NOT EXISTS question_bank (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            sn TEXT,
            question_type TEXT,
            big_stem TEXT,
            stem TEXT,
            answer TEXT,
            analysis TEXT,
            difficulty TEXT,
            knowledge_point_code TEXT,
            options_json TEXT,
            ai_generated INTEGER DEFAULT 0,
            content_format TEXT DEFAULT 'standard',
            module_slug TEXT,
            payload_json TEXT,
            source_file TEXT,
            unique_key TEXT UNIQUE,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            updated_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
        '''
    )
    cur.execute(
        '''
        CREATE TABLE IF NOT EXISTS legacy_question_statistics (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            knowledge_point_code TEXT,
            knowledge_point_name TEXT,
            question_count INTEGER DEFAULT 0,
            source_db TEXT,
            synced_at TEXT DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(knowledge_point_code, source_db)
        )
        '''
    )

    qb_columns = [row[1] for row in cur.execute('PRAGMA table_info(question_bank)').fetchall()]
    if 'ai_generated' not in qb_columns:
        cur.execute('ALTER TABLE question_bank ADD COLUMN ai_generated INTEGER DEFAULT 0')
    if 'content_format' not in qb_columns:
        cur.execute("ALTER TABLE question_bank ADD COLUMN content_format TEXT DEFAULT 'standard'")
    if 'module_slug' not in qb_columns:
        cur.execute('ALTER TABLE question_bank ADD COLUMN module_slug TEXT')
    if 'payload_json' not in qb_columns:
        cur.execute('ALTER TABLE question_bank ADD COLUMN payload_json TEXT')

    cur.execute('CREATE INDEX IF NOT EXISTS idx_qb_type ON question_bank(question_type)')
    cur.execute('CREATE INDEX IF NOT EXISTS idx_qb_diff ON question_bank(difficulty)')
    cur.execute('CREATE INDEX IF NOT EXISTS idx_qb_kp ON question_bank(knowledge_point_code)')
    cur.execute('CREATE INDEX IF NOT EXISTS idx_qb_ai_generated ON question_bank(ai_generated)')
    cur.execute('CREATE INDEX IF NOT EXISTS idx_qb_content_format ON question_bank(content_format)')
    cur.execute('CREATE INDEX IF NOT EXISTS idx_qb_module_slug ON question_bank(module_slug)')
    cur.execute('CREATE INDEX IF NOT EXISTS idx_legacy_kp ON legacy_question_statistics(knowledge_point_code)')
    conn.commit()

    count = cur.execute('SELECT COUNT(1) FROM question_bank').fetchone()[0]
    if count == 0:
        _import_knowledge_points_from_xlsx(conn, KP_XLSX_PATH)
        insert_rows = []
        for xlsx_path, fallback_type in QUESTION_XLSX_FILES:
            for q in _read_questions_from_xlsx_for_import(xlsx_path, fallback_type):
                unique_key_src = '|'.join([
                    q.get('question_type', ''),
                    q.get('stem', ''),
                    json.dumps(q.get('options', []), ensure_ascii=False),
                    q.get('knowledge_point_code', ''),
                ])
                unique_key = hashlib.md5(unique_key_src.encode('utf-8')).hexdigest()
                insert_rows.append((
                    q.get('sn', ''),
                    q.get('question_type', ''),
                    q.get('big_stem', ''),
                    q.get('stem', ''),
                    q.get('answer', ''),
                    q.get('analysis', ''),
                    q.get('difficulty', ''),
                    q.get('knowledge_point_code', ''),
                    json.dumps(q.get('options', []), ensure_ascii=False),
                    q.get('source_file', ''),
                    unique_key,
                ))

        if insert_rows:
            conn.executemany(
                '''
                INSERT OR IGNORE INTO question_bank
                (sn, question_type, big_stem, stem, answer, analysis, difficulty, knowledge_point_code, options_json, source_file, unique_key)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                insert_rows,
            )
            conn.commit()

    sync_legacy_statistics_db(conn, LEGACY_STATS_DB_PATH)

    conn.close()
    _QUESTION_BANK_READY = True


def read_questions_from_db(question_keyword: str, allowed_kps=None, db_path: str = QUESTION_DB_PATH):
    ensure_question_bank_db(db_path)
    conn = sqlite3.connect(db_path)
    cur = conn.cursor()

    sql = '''
        SELECT sn, question_type, big_stem, stem, answer, analysis, difficulty, knowledge_point_code, options_json, ai_generated
        FROM question_bank
        WHERE question_type LIKE ?
    '''
    params = [f'%{question_keyword}%']

    if allowed_kps:
        placeholders = ','.join(['?'] * len(allowed_kps))
        sql += f' AND knowledge_point_code IN ({placeholders})'
        params.extend(list(allowed_kps))

    sql += ' ORDER BY id ASC'
    cur.execute(sql, params)
    rows = cur.fetchall()
    conn.close()

    candidates = []
    for sn, qtype, big_stem, stem, answer, analysis, difficulty, kp_code, options_json, ai_generated in rows:
        try:
            options = json.loads(options_json) if options_json else []
        except Exception:
            options = []
        candidates.append({
            'sn': sn,
            'level': _safe_str(difficulty),
            'kp': _safe_str(kp_code),
            'type': _safe_str(qtype),
            'stem': _safe_str(stem),
            'big_stem': _safe_str(big_stem),
            'options': options if isinstance(options, list) else [],
            'answer': _safe_str(answer),
            'analysis': _safe_str(analysis),
            'ai_generated': bool(ai_generated),
        })
    return candidates


try:
    from milvus_vector import get_vector_store, config
    from file_process import RagFileProcessor
    from protocol.prompts import prompt_template, generate_similar_question_template, generate_er_diagram_template
    from protocol.mode import ChatRequest, ClearRequest
    RAG_AVAILABLE = True
except Exception as e:
    print(f"警告: RAG 模块导入失败或初始化出错: {e}")
    import traceback
    traceback.print_exc()
    print('RAG 功能将不可用，但试卷生成功能仍可正常使用')
    RAG_AVAILABLE = False
    get_vector_store = None
    config = None


def sanitize_option_text(opt: str) -> str:
    if not isinstance(opt, str):
        return opt
    pattern = re.compile(r'^(?:\s*[A-Za-z]\s*[\.、:：\)]\s*)+', re.UNICODE)
    cleaned = pattern.sub('', opt).strip()
    return cleaned


def normalize_question_text(text: str) -> str:
    text = str(text or '').strip()
    text = re.sub(r'\s+', ' ', text)
    return text


def normalize_question_options(options: List[str]) -> List[str]:
    return [normalize_question_text(opt) for opt in (options or []) if normalize_question_text(opt)]


def build_effective_question_text(big_stem: str, stem: str) -> str:
    big = normalize_question_text(big_stem)
    small = normalize_question_text(stem)
    return ' '.join([part for part in [big, small] if part]).strip()


def is_same_question_content(original_stem: str, original_options: List[str], candidate_stem: str, candidate_options: List[str]) -> bool:
    original_stem_norm = normalize_question_text(original_stem)
    candidate_stem_norm = normalize_question_text(candidate_stem)
    if not original_stem_norm or not candidate_stem_norm:
        return False

    same_stem = original_stem_norm == candidate_stem_norm
    same_options = normalize_question_options(original_options) == normalize_question_options(candidate_options)
    return same_stem and same_options


def is_invalid_similar_question_structure(original_big_stem: str, original_stem: str, candidate_big_stem: str, candidate_stem: str) -> bool:
    original_has_stem = bool(normalize_question_text(original_stem))
    candidate_has_stem = bool(normalize_question_text(candidate_stem))

    if not original_has_stem and candidate_has_stem:
        return True

    original_effective = build_effective_question_text(original_big_stem, original_stem)
    candidate_effective = build_effective_question_text(candidate_big_stem, candidate_stem)
    if original_effective and candidate_effective and original_effective == candidate_effective:
        return True

    if original_has_stem and normalize_question_text(original_stem) == normalize_question_text(candidate_stem):
        if normalize_question_text(original_big_stem) == normalize_question_text(candidate_big_stem):
            return True

    if not original_has_stem and normalize_question_text(original_big_stem) == normalize_question_text(candidate_big_stem):
        return True

    return False


def parse_reference_constraint(constraint: str):
    raw = str(constraint or '').strip()
    if not raw or ':' not in raw:
        return None, None
    upper = raw.upper()
    if 'NOT NULL' in upper or 'PRIMARY KEY' in upper or raw == '不允许为空':
        return None, None
    left, right = raw.split(':', 1)
    return left.strip(), right.strip()


def normalize_data_type(type_text: str) -> str:
    return str(type_text or '').strip().lower()


def normalize_table_description_label(description: str) -> str:
    text = sanitize_generated_text(description)
    text = re.sub(r'[（(]\s*[1MN]\s*:\s*[1MN]\s*[)）]', '', text, flags=re.IGNORECASE)
    text = re.sub(r'[（(][^()（）]*[，,]\s*[1MN]\s*:\s*[1MN][^()（）]*[)）]\s*$', '', text, flags=re.IGNORECASE)
    text = re.sub(r'(个人信息表|信息表|关系表|数据表|表结构|表)$', '', text).strip()
    return text or sanitize_generated_text(description)


def format_table_title(name: str, description: str) -> str:
    desc_label = normalize_table_description_label(description)
    return f"{desc_label}{name or ''}表结构"


def format_constraint_for_display(constraint: str) -> str:
    raw = str(constraint or '').strip()
    upper = raw.upper()
    if 'NOT NULL' in upper or raw == '不允许为空':
        return '不允许为空'
    return ''


def extract_type_length(type_text: str) -> Optional[int]:
    match = re.search(r'\((\d+)\)', str(type_text or ''))
    return int(match.group(1)) if match else None


def clip_text(value: str, max_len: Optional[int]) -> str:
    if max_len is None:
        return value
    return str(value)[:max_len]


def get_table_seed(table_name: str) -> int:
    text = str(table_name or '')
    return sum(ord(ch) for ch in text)


def is_person_name_column(table_name: str, column_name: str, column_desc: str) -> bool:
    lower_table = str(table_name or '').lower()
    lower_name = str(column_name or '').lower()
    lower_desc = str(column_desc or '').lower()
    person_table_keywords = ['teacher', 'doctor', 'guest', 'patient', 'student', 'employee', 'user', 'customer', 'member', 'driver', 'passenger']
    person_column_keywords = ['teacher_name', 'doctor_name', 'guest_name', 'patient_name', 'student_name', 'employee_name', 'user_name', 'customer_name', 'member_name', 'driver_name', 'passenger_name']

    if '姓名' in lower_desc or '名字' in lower_desc:
        return True
    if lower_name in person_column_keywords:
        return True
    if lower_name == 'name' and any(keyword in lower_table for keyword in person_table_keywords):
        return True
    return False


def build_non_person_name_value(table_name: str, column_name: str, column_desc: str, row_index: int, max_len: Optional[int]) -> str:
    lower_table = str(table_name or '').lower()
    lower_name = str(column_name or '').lower()
    lower_desc = str(column_desc or '').lower()

    value_map = {
        'course': ['数据库原理', '操作系统', '数据结构', '软件工程', '人工智能基础'],
        'subject': ['高等数学', '大学英语', '离散数学', '计算机网络', '编译原理'],
        'hotel': ['豪华大床房', '商务双床房', '行政套房', '精品单人房', '园景双床房'],
        'room': ['标准间', '豪华间', '商务房', '观景房', '家庭套房'],
        'type': ['小型汽车', '大型客车', '普通二轮摩托车', '重型卡车', '新能源轿车'],
        'train': ['G101次列车', 'D215次列车', 'K302次列车', 'Z88次列车', 'C620次列车'],
        'flight': ['CA1234航班', 'MU5678航班', 'CZ2468航班', 'HU1357航班', 'FM4321航班']
    }

    for keyword, values in value_map.items():
        if keyword in lower_table or keyword in lower_name:
            return clip_text(values[row_index % len(values)], max_len)

    if '课程' in lower_desc:
        values = value_map['course']
        return clip_text(values[row_index % len(values)], max_len)
    if '房型' in lower_desc or '房间' in lower_desc:
        values = value_map['room']
        return clip_text(values[row_index % len(values)], max_len)
    if '车型' in lower_desc or '类型' in lower_desc:
        values = value_map['type']
        return clip_text(values[row_index % len(values)], max_len)

    label = str(table_name or 'item')
    if lower_name.endswith('_name') or '名称' in lower_desc:
        return clip_text(f'{label}名称{row_index + 1}', max_len)
    return clip_text(f'{label}{column_name}{row_index + 1}', max_len)


def build_column_value(table_name: str, column: Dict[str, Any], row_index: int, generated_tables: Dict[str, List[Dict[str, str]]]) -> str:
    col_name = str(column.get('name', ''))
    col_desc = str(column.get('desc', ''))
    col_type = normalize_data_type(column.get('type', ''))
    constraint = str(column.get('constraint', ''))
    max_len = extract_type_length(col_type)
    lower_name = col_name.lower()
    lower_desc = col_desc.lower()
    table_seed = get_table_seed(table_name)

    ref_table, ref_col = parse_reference_constraint(constraint)
    if ref_table and ref_col:
        ref_rows = generated_tables.get(ref_table, [])
        if ref_rows:
            ref_row = ref_rows[row_index % len(ref_rows)]
            if ref_col in ref_row:
                return str(ref_row[ref_col])

    if 'gender' in lower_name or '性别' in lower_desc:
        values = ['男', '女', '男']
        return clip_text(values[row_index % len(values)], max_len)

    if 'age' in lower_name or '年龄' in lower_desc:
        return str(22 + row_index * 7)

    if 'phone' in lower_name or '电话' in lower_desc or '手机号' in lower_desc:
        phone_base = 13800000000 + (table_seed % 700) * 10
        return clip_text(str(phone_base + row_index + 1), max_len)

    if 'name' in lower_name or '姓名' in lower_desc or '名称' in lower_desc:
        if not is_person_name_column(table_name, col_name, col_desc):
            return build_non_person_name_value(table_name, col_name, col_desc, row_index, max_len)
        values = ['张佳', '李伟', '王浩', '赵敏', '陈晨', '孙涛', '周宁', '吴菲']
        offset = table_seed % len(values)
        return clip_text(values[(row_index + offset) % len(values)], max_len)

    if 'date' in lower_name or '日期' in lower_desc:
        if 'datetime' in col_type or 'timestamp' in col_type or 'time' in lower_name:
            return f'2024-0{(row_index % 9) + 1}-0{(row_index % 8) + 1} 10:0{row_index}:00'
        return f'2024-0{(row_index % 9) + 1}-0{(row_index % 8) + 1}'

    if 'count' in lower_name or '数量' in lower_desc or col_type in {'int', 'integer', 'smallint', 'bigint', 'tinyint'}:
        return str((row_index + 1) * 5)

    if lower_name.endswith('no') or '编号' in lower_desc or '代码' in lower_desc:
        prefix = ''.join(ch for ch in col_name if ch.isalpha())[:1].upper() or ''.join(ch for ch in table_name if ch.isalpha())[:1].upper() or 'X'
        if max_len is not None and max_len <= 2:
            return str(row_index + 1).zfill(max_len)
        if max_len is not None and max_len <= 3:
            return clip_text(f'{prefix}{row_index + 1}', max_len)
        return f'{prefix}{str(row_index + 1).zfill(2)}'

    if 'char' in col_type or 'text' in col_type:
        seed = [table_name, col_name, str(row_index + 1)]
        return clip_text(''.join(seed), max_len)

    return clip_text(f'{table_name}_{col_name}_{row_index + 1}', max_len)


def attach_sample_rows_to_tables(tables: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    generated_tables: Dict[str, List[Dict[str, str]]] = {}
    for table in tables:
        rows = []
        row_count = 3 if len(table.get('columns', [])) >= 3 else 2
        for row_index in range(row_count):
            row_data = {}
            for column in table.get('columns', []):
                row_data[column.get('name', '')] = build_column_value(table.get('name', ''), column, row_index, generated_tables)
            rows.append(row_data)
        table['sample_rows'] = rows
        generated_tables[table.get('name', '')] = rows
    return tables


def is_primary_key_constraint(constraint: str) -> bool:
    return 'PRIMARY KEY' in str(constraint or '').upper()


def is_foreign_key_constraint(constraint: str) -> bool:
    ref_table, ref_col = parse_reference_constraint(constraint)
    return bool(ref_table and ref_col)


def build_structure_table_mask(tables: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    for table in tables:
        for column in table.get('columns', []) or []:
            column['masked_field'] = None

    candidates = []
    for table_index, table in enumerate(tables):
        columns = table.get('columns', []) or []
        foreign_key_count = sum(1 for column in columns if is_foreign_key_constraint(column.get('constraint', '')))
        is_junction_table = foreign_key_count >= 2
        for column_index, column in enumerate(columns):
            priority = 1
            if not is_primary_key_constraint(column.get('constraint', '')) and not is_foreign_key_constraint(column.get('constraint', '')):
                priority = 4
            elif not is_primary_key_constraint(column.get('constraint', '')):
                priority = 3
            elif not is_foreign_key_constraint(column.get('constraint', '')):
                priority = 2

            col_type = normalize_data_type(column.get('type', ''))
            if any(token in col_type for token in ['varchar', 'char', 'text', 'date', 'datetime', 'time']):
                priority += 1
            if is_junction_table:
                priority += 0.5

            candidates.append({
                'priority': priority,
                'table_index': table_index,
                'column_index': column_index
            })

    if not candidates:
        return tables

    max_priority = max(candidate['priority'] for candidate in candidates)
    best_candidates = [candidate for candidate in candidates if candidate['priority'] == max_priority]
    selected = random.choice(best_candidates)
    tables[selected['table_index']]['columns'][selected['column_index']]['masked_field'] = 'type'
    return tables


def build_sample_data_mask(tables: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    for table in tables:
        table['sample_mask'] = None

    referenced_targets = defaultdict(list)
    junction_reference_targets = set()
    for table in tables:
        columns = table.get('columns', []) or []
        foreign_key_columns = [column for column in columns if is_foreign_key_constraint(column.get('constraint', ''))]
        is_junction_table = len(foreign_key_columns) >= 2

        for column in columns:
            ref_table, ref_col = parse_reference_constraint(column.get('constraint', ''))
            if ref_table and ref_col:
                referenced_targets[(ref_table, ref_col)].append({
                    'table_name': table.get('name', ''),
                    'column_name': column.get('name', '')
                })
                if is_junction_table:
                    junction_reference_targets.add((ref_table, ref_col))

    candidates = []
    for table_index, table in enumerate(tables):
        sample_rows = table.get('sample_rows', []) or []
        columns = table.get('columns', []) or []
        if not sample_rows or not columns:
            continue

        foreign_key_columns = [col for col in columns if parse_reference_constraint(col.get('constraint', '')) != (None, None)]
        is_junction_table = len(foreign_key_columns) >= 2

        for row_index, row in enumerate(sample_rows):
            for column in columns:
                column_name = column.get('name', '')
                cell_value = row.get(column_name, '')
                if cell_value in (None, ''):
                    continue

                if junction_reference_targets and (table.get('name', ''), column_name) not in junction_reference_targets:
                    continue

                priority = 1
                if is_primary_key_constraint(column.get('constraint', '')):
                    if referenced_targets.get((table.get('name', ''), column_name)):
                        priority = 6
                    else:
                        priority = 2
                else:
                    ref_table, ref_col = parse_reference_constraint(column.get('constraint', ''))
                    if ref_table and ref_col:
                        priority = 5 if is_junction_table else 3

                if junction_reference_targets and is_primary_key_constraint(column.get('constraint', '')):
                    priority = 10

                if row_index > 0:
                    priority += 0.5

                candidates.append({
                    'priority': priority,
                    'table_index': table_index,
                    'row_index': row_index,
                    'column_name': column_name,
                    'answer': str(cell_value)
                })

    if not candidates:
        return tables

    max_priority = max(candidate['priority'] for candidate in candidates)
    best_candidates = [candidate for candidate in candidates if candidate['priority'] == max_priority]
    selected = random.choice(best_candidates)
    target_table = tables[selected['table_index']]
    target_table['sample_mask'] = {
        'row_index': selected['row_index'],
        'column_name': selected['column_name'],
        'answer': selected['answer']
    }
    return tables


def align_design_questions_with_sample_mask(physical_design_data: Dict[str, Any]) -> Dict[str, Any]:
    normalized = json.loads(json.dumps(physical_design_data or {}))
    tables = normalized.get('tables', []) or []
    design_questions = normalized.get('design_questions', []) or []

    structure_question = None
    for table in tables:
        for column in table.get('columns', []) or []:
            if column.get('masked_field') != 'type':
                continue
            description = normalize_table_description_label(table.get('description', '')) or table.get('name', '')
            structure_question = {
                'question': f"1）{description}{table.get('name', '')}表的{column.get('name', '')}字段数据类型定义 ________。",
                'answer': normalize_data_type(column.get('type', ''))
            }
            break
        if structure_question:
            break

    sample_question = None
    for table in tables:
        sample_mask = table.get('sample_mask') or {}
        if not sample_mask:
            continue

        description = normalize_table_description_label(table.get('description', '')) or table.get('name', '')
        sample_question = {
            'question': f"3）请根据下方样例数据，填写{description}{table.get('name', '')}表中的空缺数据 ________。",
            'answer': sample_mask.get('answer', '')
        }
        break

    if not structure_question and not sample_question:
        normalized['design_questions'] = design_questions
        return normalized

    filtered_questions = []
    for dq in design_questions:
        question_text = sanitize_generated_text(dq.get('question', ''))
        if not question_text:
            continue
        if '字段数据类型定义' in question_text:
            continue
        if '空缺数据' in question_text or '表结构展示时会随机挖空一个数据类型或约束' in question_text:
            continue
        filtered_questions.append(dq)

    if structure_question:
        filtered_questions.insert(0, structure_question)

    if sample_question:
        insert_index = 2 if len(filtered_questions) >= 2 else len(filtered_questions)
        filtered_questions.insert(insert_index, sample_question)

    for idx, dq in enumerate(filtered_questions, start=1):
        question_text = re.sub(r'^\d+[.)）]\s*', '', str(dq.get('question', '')))
        dq['question'] = f"{idx}）{question_text}"

    normalized['design_questions'] = filtered_questions
    return normalized


SCENARIO_LABELS = {
    'medical': '医疗',
    'education': '教育',
    'train': '火车',
    'flight': '飞机',
    'hotel': '酒店',
    'ticketing': '订票',
    'random': '随机'
}


def sanitize_generated_text(text: str) -> str:
    if text is None:
        return ''

    cleaned = str(text)
    cleaned = re.sub(r'[\u200b-\u200f\ufeff\u2060]', '', cleaned)
    cleaned = cleaned.replace('\xa0', ' ')
    cleaned = cleaned.strip()
    return cleaned


def normalize_sql_types_in_text(text: str) -> str:
    if not text:
        return text

    result = sanitize_generated_text(text)

    patterns = [
        r'varchar\s*\(\s*\d+\s*\)',
        r'char\s*\(\s*\d+\s*\)',
        r'nvarchar\s*\(\s*\d+\s*\)',
        r'nchar\s*\(\s*\d+\s*\)',
        r'decimal\s*\(\s*\d+\s*,\s*\d+\s*\)',
        r'numeric\s*\(\s*\d+\s*,\s*\d+\s*\)',
        r'int',
        r'integer',
        r'bigint',
        r'smallint',
        r'tinyint',
        r'date',
        r'datetime',
        r'timestamp',
        r'time',
        r'text',
        r'float',
        r'double',
        r'real',
        r'boolean',
        r'bool'
    ]

    for pattern in patterns:
        result = re.sub(pattern, lambda m: m.group(0).lower(), result, flags=re.IGNORECASE)
    return result


def normalize_er_data_types(er_data: Dict[str, Any]) -> Dict[str, Any]:
    normalized = json.loads(json.dumps(er_data or {}))

    for entity in normalized.get('entities', []):
        for attr in entity.get('attributes', []):
            if 'type' in attr:
                attr['type'] = normalize_data_type(attr.get('type', ''))

    for rel in normalized.get('relationships', []):
        if 'cardinality' in rel and rel.get('cardinality'):
            rel['cardinality'] = str(rel.get('cardinality', '')).strip().upper()
        for attr in rel.get('attributes', []):
            if 'type' in attr:
                attr['type'] = normalize_data_type(attr.get('type', ''))

    return normalized


def has_many_to_many_relationship(er_data: Dict[str, Any]) -> bool:
    for rel in (er_data or {}).get('relationships', []):
        parts = parse_cardinality_parts(rel.get('cardinality', ''))
        if len(parts) == 2 and parts[0] in {'M', 'N'} and parts[1] in {'M', 'N'}:
            return True
    return False


def ensure_many_to_many_relationship(er_data: Dict[str, Any]) -> Dict[str, Any]:
    normalized = normalize_er_data_types(er_data)
    relationships = normalized.get('relationships', [])
    if not relationships:
        return normalized

    if has_many_to_many_relationship(normalized):
        return normalized

    relationships[0]['cardinality'] = 'M:N'
    return normalized


def get_entity_primary_key_attrs(entity: Dict[str, Any]) -> List[Dict[str, Any]]:
    attrs = entity.get('attributes', []) or []
    pk_attrs = [attr for attr in attrs if attr.get('isPrimaryKey')]
    return pk_attrs if pk_attrs else (attrs[:1] if attrs else [])


def parse_cardinality_parts(cardinality: str) -> List[str]:
    return [part.strip().upper() for part in str(cardinality or '').split(':') if part.strip()]


def is_many_marker(part: str) -> bool:
    return str(part or '').strip().upper() in {'M', 'N'}


def find_matching_existing_table(existing_tables: List[Dict[str, Any]], chinese_name: str) -> Optional[Dict[str, Any]]:
    target = sanitize_generated_text(chinese_name)
    if not target:
        return None

    for table in existing_tables:
        description = sanitize_generated_text(table.get('description', ''))
        normalized_description = normalize_table_description_label(description)
        if target == description or target == normalized_description or target in description or description in target:
            return table

    return None


def is_many_to_many_cardinality(cardinality: str) -> bool:
    parts = parse_cardinality_parts(cardinality)
    return len(parts) == 2 and parts[0] in {'M', 'N'} and parts[1] in {'M', 'N'}


def get_one_to_many_entities(rel: Dict[str, Any]) -> Tuple[Optional[str], Optional[str]]:
    parts = parse_cardinality_parts(rel.get('cardinality', ''))
    if len(parts) != 2:
        return None, None

    if parts[0] == '1' and is_many_marker(parts[1]):
        return rel.get('entity1'), rel.get('entity2')
    if is_many_marker(parts[0]) and parts[1] == '1':
        return rel.get('entity2'), rel.get('entity1')
    return None, None


def find_matching_column(columns: List[Dict[str, Any]], name: str = '', desc: str = '') -> Optional[Dict[str, Any]]:
    normalized_name = sanitize_generated_text(name).lower()
    normalized_desc = sanitize_generated_text(desc).lower()

    for column in columns or []:
        column_name = sanitize_generated_text(column.get('name', '')).lower()
        column_desc = sanitize_generated_text(column.get('desc', '')).lower()
        if normalized_name and normalized_name == column_name:
            return column
        if normalized_desc and normalized_desc == column_desc:
            return column

    return None


def append_column_if_missing(target_columns: List[Dict[str, Any]], column: Dict[str, Any]):
    normalized_name = sanitize_generated_text(column.get('name', '')).lower()
    normalized_desc = sanitize_generated_text(column.get('desc', '')).lower()
    for existing in target_columns:
        existing_name = sanitize_generated_text(existing.get('name', '')).lower()
        existing_desc = sanitize_generated_text(existing.get('desc', '')).lower()
        if normalized_name and normalized_name == existing_name:
            return
        if normalized_desc and normalized_desc == existing_desc:
            return
    target_columns.append(column)


def has_ascii_letter(text: str) -> bool:
    return bool(re.search(r'[A-Za-z]', str(text or '')))


def to_snake_case_identifier(text: str) -> str:
    normalized = re.sub(r'[^A-Za-z0-9]+', '_', str(text or ''))
    normalized = re.sub(r'([a-z0-9])([A-Z])', r'\1_\2', normalized)
    normalized = re.sub(r'_+', '_', normalized).strip('_').lower()
    return normalized


def build_english_column_name(source_text: str, fallback_prefix: str = 'field') -> str:
    raw = sanitize_generated_text(source_text)
    if not raw:
        return fallback_prefix
    if has_ascii_letter(raw):
        normalized = to_snake_case_identifier(raw)
        return normalized or fallback_prefix

    exact_map = {
        '就诊时间': 'visit_time',
        '诊断结果': 'diagnosis_result',
        '科室': 'department',
        '治疗时间': 'treatment_time',
        '治疗日期': 'treatment_date',
        '诊疗时间': 'visit_time',
        '诊疗记录': 'diagnosis_record',
        '处方信息': 'prescription_info',
        '主治医生': 'attending_doctor',
        '状态': 'status',
        '备注': 'remark',
        '描述': 'description',
        '类型': 'type',
        '数量': 'count',
        '金额': 'amount'
    }
    if raw in exact_map:
        return exact_map[raw]

    token_map = {
        '就诊': 'visit',
        '诊疗': 'diagnosis',
        '诊断': 'diagnosis',
        '治疗': 'treatment',
        '处方': 'prescription',
        '医生': 'doctor',
        '患者': 'patient',
        '科室': 'department',
        '时间': 'time',
        '日期': 'date',
        '结果': 'result',
        '记录': 'record',
        '信息': 'info',
        '编号': 'code',
        '名称': 'name',
        '类型': 'type',
        '状态': 'status',
        '数量': 'count',
        '金额': 'amount',
        '备注': 'remark',
        '描述': 'description',
        '地址': 'address',
        '电话': 'phone',
        '职称': 'title',
        '年龄': 'age',
        '性别': 'gender'
    }

    parts = []
    remaining = raw
    for chinese, english in sorted(token_map.items(), key=lambda item: len(item[0]), reverse=True):
        if chinese in remaining:
            remaining = remaining.replace(chinese, ' ')
            if english not in parts:
                parts.append(english)

    if raw.endswith('ID') or raw.endswith('Id') or raw.endswith('id'):
        parts.append('id')

    identifier = '_'.join(part for part in parts if part)
    identifier = to_snake_case_identifier(identifier)
    return identifier or fallback_prefix


def build_expected_table_from_entity(entity: Dict[str, Any], existing_table: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    existing_columns = {str(col.get('name', '')).lower(): col for col in (existing_table or {}).get('columns', [])}
    existing_columns_by_desc = {sanitize_generated_text(col.get('desc', '')).lower(): col for col in (existing_table or {}).get('columns', [])}
    table = {
        'name': (existing_table or {}).get('name') or entity.get('name', ''),
        'description': (existing_table or {}).get('description') or entity.get('name', ''),
        'columns': []
    }

    for attr in entity.get('attributes', []):
        attr_name = sanitize_generated_text(attr.get('name', ''))
        existing_col = existing_columns.get(attr_name.lower(), {}) or existing_columns_by_desc.get(attr_name.lower(), {})
        table['columns'].append({
            'name': existing_col.get('name') or attr.get('name', ''),
            'desc': existing_col.get('desc') or attr.get('name', ''),
            'type': normalize_data_type(existing_col.get('type') or attr.get('type', '')),
            'constraint': existing_col.get('constraint') or ('PRIMARY KEY' if attr.get('isPrimaryKey') else ''),
            'masked_field': existing_col.get('masked_field') if existing_col else None
        })

    return table


def build_expected_table_from_relationship(
    rel: Dict[str, Any],
    entity_map: Dict[str, Dict[str, Any]],
    existing_table: Optional[Dict[str, Any]] = None,
    entity_tables_by_name: Optional[Dict[str, Dict[str, Any]]] = None
) -> Dict[str, Any]:
    existing_columns = {str(col.get('name', '')).lower(): col for col in (existing_table or {}).get('columns', [])}
    existing_columns_by_desc = {sanitize_generated_text(col.get('desc', '')).lower(): col for col in (existing_table or {}).get('columns', [])}
    used_names = set()
    columns = []

    for entity_key in ['entity1', 'entity2']:
        entity_name = rel.get(entity_key, '')
        entity = entity_map.get(entity_name, {})
        entity_table = (entity_tables_by_name or {}).get(entity_name, {})
        for pk_attr in get_entity_primary_key_attrs(entity):
            base_name = pk_attr.get('name', '')
            column_name = base_name
            if column_name.lower() in used_names:
                column_name = f"{entity_name}_{base_name}"
            used_names.add(column_name.lower())

            existing_col = existing_columns.get(column_name.lower()) or existing_columns.get(base_name.lower(), {}) or existing_columns_by_desc.get(sanitize_generated_text(pk_attr.get('name', '')).lower(), {})
            referenced_column = find_matching_column(entity_table.get('columns', []), pk_attr.get('name', ''), pk_attr.get('name', '')) or {}
            reference_table_name = entity_table.get('name') or entity_name
            reference_column_name = referenced_column.get('name') or pk_attr.get('name', '')
            columns.append({
                'name': existing_col.get('name') or column_name,
                'desc': existing_col.get('desc') or pk_attr.get('name', ''),
                'type': normalize_data_type(existing_col.get('type') or pk_attr.get('type', '')),
                'constraint': f"{reference_table_name}:{reference_column_name}",
                'masked_field': existing_col.get('masked_field') if existing_col else None
            })

    for attr in rel.get('attributes', []) or []:
        attr_name = sanitize_generated_text(attr.get('name', ''))
        existing_col = existing_columns.get(attr_name.lower(), {}) or existing_columns_by_desc.get(attr_name.lower(), {})
        columns.append({
            'name': existing_col.get('name') or build_english_column_name(attr.get('name', ''), 'relation_attr'),
            'desc': existing_col.get('desc') or attr.get('name', ''),
            'type': normalize_data_type(existing_col.get('type') or attr.get('type', '')),
            'constraint': existing_col.get('constraint') or '',
            'masked_field': existing_col.get('masked_field') if existing_col else None
        })

    return {
        'name': (existing_table or {}).get('name') or rel.get('name', ''),
        'description': (existing_table or {}).get('description') or rel.get('name', ''),
        'columns': columns
    }


def merge_one_to_many_relationship_into_table(
    target_table: Dict[str, Any],
    target_existing_table: Optional[Dict[str, Any]],
    rel: Dict[str, Any],
    one_entity: Dict[str, Any],
    one_table: Dict[str, Any]
) -> Dict[str, Any]:
    target_columns = target_table.get('columns', [])
    existing_columns = (target_existing_table or {}).get('columns', [])

    for pk_attr in get_entity_primary_key_attrs(one_entity):
        referenced_column = find_matching_column(one_table.get('columns', []), pk_attr.get('name', ''), pk_attr.get('name', '')) or {}
        default_name = referenced_column.get('name') or pk_attr.get('name', '')
        default_desc = referenced_column.get('desc') or pk_attr.get('name', '')
        existing_col = find_matching_column(existing_columns, default_name, default_desc) or find_matching_column(target_columns, default_name, default_desc) or {}
        append_column_if_missing(target_columns, {
            'name': existing_col.get('name') or default_name,
            'desc': existing_col.get('desc') or default_desc,
            'type': normalize_data_type(existing_col.get('type') or referenced_column.get('type') or pk_attr.get('type', '')),
            'constraint': existing_col.get('constraint') or f"{one_table.get('name', '')}:{referenced_column.get('name') or pk_attr.get('name', '')}",
            'masked_field': existing_col.get('masked_field') if existing_col else None
        })

    for attr in rel.get('attributes', []) or []:
        attr_name = sanitize_generated_text(attr.get('name', ''))
        existing_col = find_matching_column(existing_columns, attr_name, attr_name) or find_matching_column(target_columns, attr_name, attr_name) or {}
        append_column_if_missing(target_columns, {
            'name': existing_col.get('name') or build_english_column_name(attr.get('name', ''), 'relation_attr'),
            'desc': existing_col.get('desc') or attr.get('name', ''),
            'type': normalize_data_type(existing_col.get('type') or attr.get('type', '')),
            'constraint': existing_col.get('constraint') or '',
            'masked_field': existing_col.get('masked_field') if existing_col else None
        })

    return target_table


def repair_physical_design_tables(physical_design_data: Dict[str, Any], er_data: Dict[str, Any]) -> Dict[str, Any]:
    repaired = json.loads(json.dumps(physical_design_data or {}))
    existing_tables = repaired.get('tables', []) or []
    entity_map = {entity.get('name', ''): entity for entity in (er_data or {}).get('entities', [])}

    repaired_tables = []
    entity_tables_by_name = {}
    existing_entity_tables = {}
    for entity in (er_data or {}).get('entities', []):
        existing_table = find_matching_existing_table(existing_tables, entity.get('name', ''))
        repaired_table = build_expected_table_from_entity(entity, existing_table)
        repaired_tables.append(repaired_table)
        entity_tables_by_name[entity.get('name', '')] = repaired_table
        existing_entity_tables[entity.get('name', '')] = existing_table

    for rel in (er_data or {}).get('relationships', []):
        one_entity_name, many_entity_name = get_one_to_many_entities(rel)
        if not one_entity_name or not many_entity_name:
            continue

        many_table = entity_tables_by_name.get(many_entity_name)
        one_table = entity_tables_by_name.get(one_entity_name)
        if not many_table or not one_table:
            continue

        merge_one_to_many_relationship_into_table(
            many_table,
            existing_entity_tables.get(many_entity_name),
            rel,
            entity_map.get(one_entity_name, {}),
            one_table
        )

    for rel in (er_data or {}).get('relationships', []):
        if is_many_to_many_cardinality(rel.get('cardinality', '')):
            existing_table = find_matching_existing_table(existing_tables, rel.get('name', ''))
            repaired_tables.append(build_expected_table_from_relationship(rel, entity_map, existing_table, entity_tables_by_name))

    repaired['tables'] = repaired_tables
    return repaired


def normalize_physical_design_data(physical_design_data: Dict[str, Any]) -> Dict[str, Any]:
    normalized = json.loads(json.dumps(physical_design_data or {}))

    for table in normalized.get('tables', []):
        for col in table.get('columns', []):
            if 'type' in col:
                col['type'] = normalize_data_type(col.get('type', ''))

    for dq in normalized.get('design_questions', []):
        if 'question' in dq:
            dq['question'] = normalize_sql_types_in_text(dq.get('question', ''))
        if 'answer' in dq:
            dq['answer'] = sanitize_generated_text(normalize_sql_types_in_text(dq.get('answer', '')))

    for rq in normalized.get('relation_questions', []):
        if 'relation_mode' in rq:
            rq['relation_mode'] = normalize_sql_types_in_text(rq.get('relation_mode', ''))
        if 'pk' in rq:
            rq['pk'] = sanitize_generated_text(rq.get('pk', ''))
        if 'fk' in rq:
            rq['fk'] = sanitize_generated_text(rq.get('fk', ''))

    return normalized


def add_underlined_answer_run(paragraph, answer_text: str):
    run = paragraph.add_run(str(answer_text or ''))
    run.font.color.rgb = RGBColor(255, 0, 0)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    u.set(qn('w:color'), '000000')
    run._r.get_or_add_rPr().append(u)
    return run


def add_blank_or_answer(paragraph, answer_text: str, show_answer: bool, blank_text: str = '________'):
    if show_answer:
        return add_underlined_answer_run(paragraph, answer_text)
    return paragraph.add_run(blank_text)


def render_question_with_blank(paragraph, question_text: str, answer_text: str, show_answer: bool):
    text = sanitize_generated_text(question_text)
    answer_text = sanitize_generated_text(answer_text)
    parts = re.split(r'(_{4,}|＿{2,}|[○◯□▢]{1,})', text)
    replaced = False
    for part in parts:
        if re.fullmatch(r'(_{4,}|＿{2,}|[○◯□▢]{1,})', part or ''):
            if not replaced:
                add_blank_or_answer(paragraph, answer_text, show_answer, part)
                replaced = True
            else:
                paragraph.add_run(part)
        else:
            paragraph.add_run(part)


def safe_json_dumps(data: Any) -> str:
    return json.dumps(data, ensure_ascii=False)


def safe_json_loads(raw: Any, default: Any):
    if raw in (None, ''):
        return default
    try:
        return json.loads(raw)
    except Exception:
        return default


def build_short_answer_question_payload(raw_payload: Dict[str, Any]) -> Dict[str, Any]:
    module_slug = _safe_str(raw_payload.get('module_slug')) or 'database-big-question-type-1'
    module = get_short_answer_module(module_slug)
    if not module:
        raise ValueError('无效的综合题模块')

    physical_design_data = raw_payload.get('physical_design_data', {}) or {}
    if not isinstance(physical_design_data, dict):
        physical_design_data = {}

    answers = raw_payload.get('answers', {}) or {}
    if not isinstance(answers, dict):
        answers = {}

    masked_er_data = raw_payload.get('masked_er_data', {}) or {}
    if not isinstance(masked_er_data, dict):
        masked_er_data = {}

    payload = {
        'module_slug': module_slug,
        'module_name': module.get('name', ''),
        'main_title': _safe_str(raw_payload.get('main_title')) or '(一) 信息系统设计（本题3小题，共20分）',
        'background_text': _safe_str(raw_payload.get('background_text')),
        'q1_text': _safe_str(raw_payload.get('q1_text')) or '1. 请问系统有哪几类用户？简述需要为各类用户设计哪些系统功能。',
        'q1_answer': _safe_str(raw_payload.get('q1_answer')) or '简述功能，合理即可',
        'q2_text': _safe_str(raw_payload.get('q2_text')) or '2. 某设计人员给出了该系统数据库概念设计的E-R图（图1），请补充缺失部分的A、B、C处所对应的联系、属性、联系类型，完成E-R图。',
        'er_caption': _safe_str(raw_payload.get('er_caption')) or '图1 系统局部E-R图',
        'q3_intro': _safe_str(raw_payload.get('q3_intro')) or '3. 根据E-R图，完成以下数据库物理设计任务：',
        'q3_sub1': _safe_str(raw_payload.get('q3_sub1')) or '(1) 完善下列关系模式（只列出主键和外键，以及缺失的属性）：',
        'q3_sub2': _safe_str(raw_payload.get('q3_sub2')) or '(2) 根据上述表结构设计，回答下列问题：',
        'question_stem': _safe_str(raw_payload.get('question_stem') or raw_payload.get('question_text')),
        'answers': {str(k): _safe_str(v) for k, v in answers.items()},
        'er_image': _safe_str(raw_payload.get('er_image')),
        'masked_er_data': masked_er_data,
        'physical_design_data': normalize_physical_design_data(physical_design_data),
    }

    return payload


def build_short_answer_question_bank_record(raw_payload: Dict[str, Any]) -> Dict[str, Any]:
    payload = build_short_answer_question_payload(raw_payload)

    summary_parts = [
        payload.get('background_text', ''),
        payload.get('q1_text', ''),
        payload.get('q2_text', ''),
        payload.get('q3_intro', ''),
        payload.get('q3_sub1', ''),
        payload.get('q3_sub2', ''),
    ]
    stem_text = payload.get('main_title', '')
    analysis_text = '\n'.join(part for part in summary_parts if part)
    unique_key_src = '|'.join([
        SHORT_ANSWER_CONTENT_FORMAT,
        payload.get('module_slug', ''),
        payload.get('main_title', ''),
        payload.get('background_text', ''),
        safe_json_dumps(payload.get('physical_design_data', {})),
        safe_json_dumps(payload.get('answers', {})),
    ])

    return {
        'question_type': payload.get('module_name', '综合题'),
        'big_stem': payload.get('background_text', ''),
        'stem': stem_text,
        'answer': '',
        'analysis': analysis_text,
        'difficulty': '',
        'knowledge_point_code': '',
        'options_json': '[]',
        'ai_generated': 1,
        'content_format': SHORT_ANSWER_CONTENT_FORMAT,
        'module_slug': payload.get('module_slug', ''),
        'payload_json': safe_json_dumps(payload),
        'source_file': 'short_answer_module',
        'unique_key': hashlib.md5(unique_key_src.encode('utf-8')).hexdigest(),
    }


def export_short_answer_payload_to_docx(raw_payload: Dict[str, Any]) -> Dict[str, Any]:
    payload = build_short_answer_question_payload(raw_payload)
    background_text = payload.get('background_text', '')
    main_title = payload.get('main_title', '(一) 信息系统设计（本题3小题，共20分）')
    q1_text = payload.get('q1_text', '1. 请问系统有哪几类用户？简述需要为各类用户设计哪些系统功能。')
    q1_answer = payload.get('q1_answer', '简述功能，合理即可')
    q2_text = payload.get('q2_text', '2. 某设计人员给出了该系统数据库概念设计的E-R图（图1），请补充缺失部分的A、B、C处所对应的联系、属性、联系类型，完成E-R图。')
    er_caption = payload.get('er_caption', '图1 系统局部E-R图')
    q3_intro = payload.get('q3_intro', '3. 根据E-R图，完成以下数据库物理设计任务：')
    q3_sub1 = payload.get('q3_sub1', '(1) 完善下列关系模式（只列出主键和外键，以及缺失的属性）：')
    q3_sub2 = payload.get('q3_sub2', '(2) 根据上述表结构设计，回答下列问题：')
    answers = payload.get('answers', {}) or {}
    er_image_data = payload.get('er_image', None)
    physical_design_data = payload.get('physical_design_data', {}) or {}
    show_answer = bool(raw_payload.get('show_answer', False))

    output_dir = './output'
    os.makedirs(output_dir, exist_ok=True)
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    docx_name = f'question_export_{ts}.docx'
    docx_path = os.path.join(output_dir, docx_name)

    doc = Document()

    p = doc.add_paragraph()
    r = p.add_run(main_title)
    r.bold = True
    r.font.size = Pt(12)

    if background_text:
        for line in background_text.split('\n'):
            para = doc.add_paragraph(line)
            para.paragraph_format.space_after = Pt(6)

    doc.add_paragraph('请完成以下系统分析和设计，将你的答案填写到相应题号下的空格中。')

    doc.add_paragraph(q1_text)
    if show_answer:
        p = doc.add_paragraph()
        add_underlined_answer_run(p, q1_answer)

    doc.add_paragraph(q2_text)
    if show_answer and answers:
        q2_keys = ['A', 'B', 'C']
        p = doc.add_paragraph()
        first = True
        for k in q2_keys:
            if k in answers:
                if not first:
                    p.add_run(', ')
                p.add_run(f'{k}: ')
                add_underlined_answer_run(p, answers.get(k, ''))
                first = False

    if er_image_data and isinstance(er_image_data, str) and er_image_data.startswith('data:image'):
        try:
            _, b64 = er_image_data.split(',', 1)
            img_bytes = base64.b64decode(b64)
            img_filename = f'er_{ts}.png'
            img_path = os.path.join(output_dir, img_filename)
            with open(img_path, 'wb') as f:
                f.write(img_bytes)

            doc.add_picture(img_path, width=Inches(5.5))
            last = doc.paragraphs[-1]
            last.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            cap = doc.add_paragraph(er_caption)
            cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cap.runs[0].font.size = Pt(9)
        except Exception as e:
            print(f'插入图片失败: {e}')

    if physical_design_data:
        doc.add_paragraph(q3_intro)

        tables = physical_design_data.get('tables', [])
        for table in tables:
            doc.add_paragraph(format_table_title(table.get('name', ''), table.get('description', '')))
            table_grid = doc.add_table(rows=1, cols=4)
            table_grid.style = 'Table Grid'
            hdr = table_grid.rows[0].cells
            hdr[0].text = '字段名'
            hdr[1].text = '字段描述'
            hdr[2].text = '数据类型'
            hdr[3].text = '属性限制'

            for col in table.get('columns', []):
                row = table_grid.add_row().cells
                row[0].text = col.get('name', '')
                row[1].text = col.get('desc', '')

                masked_field = col.get('masked_field')
                normalized_type = normalize_data_type(col.get('type', ''))

                p_type = row[2].paragraphs[0]
                if masked_field == 'type':
                    add_blank_or_answer(p_type, normalized_type, show_answer, '________')
                    p_type.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                else:
                    row[2].text = normalized_type

                display_constraint = format_constraint_for_display(col.get('constraint', '') or '')
                p_constr = row[3].paragraphs[0]
                if masked_field == 'constraint':
                    add_blank_or_answer(p_constr, display_constraint, show_answer, '________')
                    p_constr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                else:
                    row[3].text = display_constraint

            doc.add_paragraph('')

        for table in tables:
            sample_rows = table.get('sample_rows', []) or []
            if not sample_rows:
                continue

            sample_mask = table.get('sample_mask') or {}
            doc.add_paragraph(f"{table.get('name', '')}表数据")
            sample_table = doc.add_table(rows=1, cols=len(table.get('columns', [])))
            sample_table.style = 'Table Grid'
            header_cells = sample_table.rows[0].cells
            for idx, column in enumerate(table.get('columns', [])):
                header_cells[idx].text = column.get('name', '')

            for row_index, sample_row in enumerate(sample_rows):
                sample_cells = sample_table.add_row().cells
                for idx, column in enumerate(table.get('columns', [])):
                    value = str(sample_row.get(column.get('name', ''), ''))
                    paragraph = sample_cells[idx].paragraphs[0]
                    if sample_mask and row_index == sample_mask.get('row_index') and column.get('name', '') == sample_mask.get('column_name'):
                        add_blank_or_answer(paragraph, sample_mask.get('answer', value), show_answer, '____')
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    else:
                        paragraph.add_run(value)

            doc.add_paragraph('')

        doc.add_paragraph(q3_sub1)
        req_idx = 1
        for rq in physical_design_data.get('relation_questions', []):
            p = doc.add_paragraph(f"   {req_idx}) {rq.get('table_name')}表: {rq.get('table_name')}(" )
            full_relation = rq.get('relation_mode', '')
            content = ''
            if '(' in full_relation:
                try:
                    content = full_relation.split('(', 1)[1].rstrip(')')
                except Exception:
                    content = full_relation
            else:
                content = full_relation

            add_blank_or_answer(p, content, show_answer, '________')
            p.add_run(')')

            p2 = doc.add_paragraph('      主关键字：')
            add_blank_or_answer(p2, rq.get('pk', ''), show_answer, '___________')
            p2.add_run('   外关键字：')
            add_blank_or_answer(p2, rq.get('fk', ''), show_answer, '__________')
            req_idx += 1

        doc.add_paragraph(q3_sub2)
        q_idx = 1
        for dq in physical_design_data.get('design_questions', []):
            clean_question = re.sub(r'^\d+[.)）]\s*', '', dq.get('question', ''))
            p = doc.add_paragraph(f'   {q_idx}) ')
            render_question_with_blank(p, clean_question, dq.get('answer', ''), show_answer)
            q_idx += 1
    else:
        doc.add_paragraph('3. 暂时占位，后续会补充第三问')

    doc.save(docx_path)
    return {'success': True, 'file_url': f'/output/{docx_name}', 'file_path': docx_path}

    if show_answer and answer_text and not replaced:
        paragraph.add_run(' ')
        add_underlined_answer_run(paragraph, answer_text)

    return replaced

def save_questions_to_csv(questions: List[Dict[str, Any]], filename: str = None) -> str:
    """
    将题目列表保存到CSV文件
    
    Args:
        questions: 题目列表
        filename: 文件名，如果为None则使用当前日期
        
    Returns:
        str: 保存的文件路径
    """
    if filename is None:
        filename = f"questions_{datetime.now().strftime('%Y%m%d')}.csv"
    
    # 确保文件保存在特定目录
    output_dir = "./output"
    os.makedirs(output_dir, exist_ok=True)
    file_path = os.path.join(output_dir, filename)
    
    # CSV文件字段
    fieldnames = ["question", "options", "answer", "explanation", "difficulty"]
    
    with open(file_path, 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        writer.writeheader()
        
        for question_data in questions:
            # 将选项列表转换为字符串
            options_str = " | ".join(question_data.get("options", []))
            
            writer.writerow({
                "question": question_data.get("question", ""),
                "options": options_str,
                "answer": question_data.get("answer", ""),
                "explanation": question_data.get("explanation", ""),
                "difficulty": question_data.get("difficulty", "")
            })
    
    return file_path


def parse_llm_response(response_text: str) -> List[Dict[str, Any]]:
    """
    解析大模型返回的JSON格式题目
    
    Args:
        response_text: 大模型返回的文本
        
    Returns:
        List[Dict]: 解析后的题目列表
    """
    try:
        # 清理响应文本，移除可能的markdown代码块标记
        cleaned_text = response_text.strip()
        if cleaned_text.startswith('```json'):
            cleaned_text = cleaned_text[7:]
        if cleaned_text.endswith('```'):
            cleaned_text = cleaned_text[:-3]
        cleaned_text = cleaned_text.strip()
        
        questions = json.loads(cleaned_text)
        if isinstance(questions, list):
            return questions
        else:
            return [questions]
    except json.JSONDecodeError as e:
        print(f"JSON解析错误: {e}")
        print(f"原始响应: {response_text}")
        return []

# ============= 试卷生成相关函数 =============

def load_knowledge_mapping(kp_file_path):
    """优先从 SQLite 加载知识点代码到名称的映射，缺失时回退到 xlsx。"""
    try:
        ensure_question_bank_db()
        conn = sqlite3.connect(QUESTION_DB_PATH)
        cur = conn.cursor()
        cur.execute('SELECT knowledge_point_code, knowledge_point_name FROM knowledge_points')
        rows = cur.fetchall()
        conn.close()
        mapping = {}
        for code, name in rows:
            code_s = _safe_str(code)
            name_s = _safe_str(name)
            if code_s and name_s:
                mapping[code_s] = name_s
        if mapping:
            return mapping
    except Exception as e:
        print(f"读取 SQLite 知识点映射失败，回退到 xlsx: {e}")

    if not os.path.exists(kp_file_path):
        print(f"警告：知识点列表文件 {kp_file_path} 不存在，将使用知识点代码")
        return {}
    
    try:
        wb = load_workbook(kp_file_path, data_only=True)
        sheet = wb.worksheets[0]
        mapping = {}
        
        # 读取表头
        header_row = next(sheet.iter_rows(min_row=1, max_row=1, values_only=True))
        headers = [h.strip() if isinstance(h, str) else h for h in header_row]
        
        # 查找知识点代码和知识点名称的列索引
        kp_code_idx = None
        kp_name_idx = None
        
        for i, header in enumerate(headers):
            if header in ['知识点代码', '知识点编码']:
                kp_code_idx = i
            elif header in ['知识点名称', '知识点']:
                kp_name_idx = i
        
        # 如果找不到标准列名，使用默认位置（知识点代码在第6列，知识点名称在第7列）
        if kp_code_idx is None and len(headers) > 6:
            kp_code_idx = 6
        if kp_name_idx is None and len(headers) > 7:
            kp_name_idx = 7
        
        if kp_code_idx is None or kp_name_idx is None:
            print("警告：无法在知识点列表文件中找到知识点代码和名称列")
            wb.close()
            return {}
        
        # 读取数据行
        for row in sheet.iter_rows(min_row=2, values_only=True):
            if len(row) <= max(kp_code_idx, kp_name_idx):
                continue
            kp_code = str(row[kp_code_idx]).strip() if row[kp_code_idx] else ''
            kp_name = str(row[kp_name_idx]).strip() if row[kp_name_idx] else ''
            if kp_code and kp_name:
                mapping[kp_code] = kp_name
        
        wb.close()
        print(f"成功加载 {len(mapping)} 个知识点映射")
        return mapping
    except Exception as e:
        print(f"加载知识点映射文件时出错: {e}")
        return {}


def parse_kp_hierarchy(kp_file_path):
    """解析知识点列表 xlsx，返回三级树和名称->代码映射。

    返回 (tree, name_to_code)
    tree 格式：[{ 'area': '知识领域名称', 'units':[{'unit':'单元名','points':[{'name':..., 'code':...}, ...]}, ...]}, ...]
    name_to_code: { 知识点名称: 知识点代码 }
    """
    try:
        ensure_question_bank_db()
        conn = sqlite3.connect(QUESTION_DB_PATH)
        cur = conn.cursor()
        cur.execute(
            '''
            SELECT knowledge_domain_name, knowledge_unit_name, knowledge_point_name, knowledge_point_code
            FROM knowledge_points
            ORDER BY knowledge_domain_name, knowledge_unit_name, knowledge_point_code
            '''
        )
        rows = cur.fetchall()
        conn.close()

        if rows:
            tree_map = {}
            name_to_code = {}
            for area, unit, point_name, point_code in rows:
                area = _safe_str(area) or '未分类'
                unit = _safe_str(unit) or '未分类单元'
                point_name = _safe_str(point_name)
                point_code = _safe_str(point_code)
                if not point_name:
                    continue
                if area not in tree_map:
                    tree_map[area] = {}
                if unit not in tree_map[area]:
                    tree_map[area][unit] = []
                tree_map[area][unit].append({'name': point_name, 'code': point_code})
                name_to_code[point_name] = point_code

            tree = []
            for area, units in tree_map.items():
                ulist = []
                for unit, points in units.items():
                    ulist.append({'unit': unit, 'points': points})
                tree.append({'area': area, 'units': ulist})
            if tree:
                return tree, name_to_code
    except Exception as e:
        print(f"从 SQLite 读取知识点层级失败，回退到 xlsx: {e}")

    if not os.path.exists(kp_file_path):
        return [], {}
    try:
        wb = load_workbook(kp_file_path, data_only=True)
        sheet = wb.worksheets[0]
        header_row = next(sheet.iter_rows(min_row=1, max_row=1, values_only=True))
        headers = [h.strip() if isinstance(h, str) else h for h in header_row]

        def find_idx(names):
            for name in names:
                if name in headers:
                    return headers.index(name)
            return None

        idx_area = find_idx(['知识领域名称', '知识领域', '知识领域名称 '])
        idx_unit = find_idx(['知识单元名称', '知识单元', '知识单元名称 '])
        idx_point_name = find_idx(['知识点名称', '知识点', '知识点名称 '])
        idx_point_code = find_idx(['知识点代码', '知识点编码', '知识点代码 '])

        tree_map = {}
        name_to_code = {}
        for row in sheet.iter_rows(min_row=2, values_only=True):
            if not any(row):
                continue
            area = row[idx_area] if idx_area is not None and idx_area < len(row) else ''
            unit = row[idx_unit] if idx_unit is not None and idx_unit < len(row) else ''
            point_name = row[idx_point_name] if idx_point_name is not None and idx_point_name < len(row) else ''
            point_code = row[idx_point_code] if idx_point_code is not None and idx_point_code < len(row) else ''
            if not point_name:
                continue
            area = str(area).strip() if area else '未分类'
            unit = str(unit).strip() if unit else '未分类单元'
            point_name = str(point_name).strip()
            point_code = str(point_code).strip() if point_code else ''

            if area not in tree_map:
                tree_map[area] = {}
            if unit not in tree_map[area]:
                tree_map[area][unit] = []
            tree_map[area][unit].append({'name': point_name, 'code': point_code})
            if point_name:
                name_to_code[point_name] = point_code

        tree = []
        for area, units in tree_map.items():
            ulist = []
            for unit, points in units.items():
                ulist.append({'unit': unit, 'points': points})
            tree.append({'area': area, 'units': ulist})

        wb.close()
        return tree, name_to_code
    except Exception as e:
        print(f"解析知识点文件出错: {e}")
        return [], {}


def load_kp_counts(db_path=QUESTION_DB_PATH):
    """从 SQLite 数据库加载每个知识点的题目数量，返回 name->count 和 code->count 两个映射。"""
    name_counts = {}
    code_counts = {}
    try:
        ensure_question_bank_db(db_path)
        conn = sqlite3.connect(db_path)
        cur = conn.cursor()
        cur.execute(
            '''
            SELECT q.knowledge_point_code, k.knowledge_point_name, COUNT(1) AS question_count
            FROM question_bank q
            LEFT JOIN knowledge_points k ON q.knowledge_point_code = k.knowledge_point_code
            GROUP BY q.knowledge_point_code, k.knowledge_point_name
            '''
        )
        for code, name, cnt in cur.fetchall():
            name_s = str(name).strip() if name else ''
            code_s = str(code).strip() if code else ''
            try:
                cnt_i = int(cnt)
            except Exception:
                cnt_i = 0
            if name_s:
                name_counts[name_s] = cnt_i
            if code_s:
                code_counts[code_s] = cnt_i

        # 融合旧库同步过来的历史统计；若与现有统计冲突，取较大值避免回退
        cur.execute(
            '''
            SELECT knowledge_point_code, knowledge_point_name, question_count
            FROM legacy_question_statistics
            '''
        )
        for code, name, cnt in cur.fetchall():
            name_s = str(name).strip() if name else ''
            code_s = str(code).strip() if code else ''
            try:
                cnt_i = int(cnt)
            except Exception:
                cnt_i = 0
            if name_s:
                name_counts[name_s] = max(name_counts.get(name_s, 0), cnt_i)
            if code_s:
                code_counts[code_s] = max(code_counts.get(code_s, 0), cnt_i)
        conn.close()
    except Exception as e:
        print(f"读取数据库 {db_path} 出错: {e}")
    return name_counts, code_counts

def parse_ratio_arg(s, keys):
    """解析比例字符串并归一化。
    支持两种格式："高:30,中:40,低:30" 或 "30,40,30"（顺序按 keys）。
    返回字典，值为 0..1 之间且总和为 1。
    """
    parts = [p.strip() for p in s.split(',') if p.strip()]
    d = {}
    if all(':' in p for p in parts):
        for p in parts:
            k, v = p.split(':', 1)
            d[k.strip()] = float(v)
    else:
        vals = [float(p) for p in parts]
        for k, v in zip(keys, vals):
            d[k] = v
    total = sum(d.values())
    if total == 0:
        return {k: 1.0 / len(keys) for k in keys}
    return {k: (d.get(k, 0.0) / total) for k in keys}


def compute_counts_from_ratios(n, ratios):
    """把比例分配为整数计数，保证总和为 n。

    使用最大余数法分配多余项，能较公平地将小数部分分配为整数。
    """
    keys = list(ratios.keys())
    raw = {k: ratios[k] * n for k in keys}
    counts = {k: int(raw[k]) for k in keys}
    remainder = n - sum(counts.values())
    frac = sorted(keys, key=lambda k: raw[k] - int(raw[k]), reverse=True)
    i = 0
    while remainder > 0:
        counts[frac[i % len(frac)]] += 1
        remainder -= 1
        i += 1
    return counts


def read_single_choice_xlsx(xlsx_path, allowed_kps=None):
    """兼容旧函数名，实际从 SQLite 读取候选单选题。"""
    return read_questions_from_db('单选', allowed_kps=allowed_kps)


def read_multi_choice_xlsx(xlsx_path, allowed_kps=None):
    """兼容旧函数名，实际从 SQLite 读取候选多选题。"""
    return read_questions_from_db('多选', allowed_kps=allowed_kps)


def find_selection_backtrack(candidates, N, target_diff_counts, target_kp_counts):
    # candidates: 列表，每项包含 'level' 和 'kp'
    # target_kp_counts 为空时，知识点仅用于构建候选池，不参与硬性配额约束。
    M = len(candidates)
    # 预计算后缀可用量用于剪枝（从 i 到末尾各类别可用数量）
    diff_keys = list(target_diff_counts.keys())
    kp_keys = list(target_kp_counts.keys())

    # 为便于访问构建数组
    levels = [c['level'] for c in candidates]
    kps = [c['kp'] for c in candidates]

    # 后缀计数：suffix_diff[i] 表示从 i 到末尾各难度的可用数量
    suffix_diff = [defaultdict(int) for _ in range(M + 1)]
    suffix_kp = [defaultdict(int) for _ in range(M + 1)]
    for i in range(M - 1, -1, -1):
        suffix_diff[i] = suffix_diff[i + 1].copy()
        suffix_kp[i] = suffix_kp[i + 1].copy()
        suffix_diff[i][levels[i]] += 1
        suffix_kp[i][kps[i]] += 1

    target_total = N

    path = []

    # 当前已选计数
    sel_diff = defaultdict(int)
    sel_kp = defaultdict(int)

    found = [None]
    # 记录搜索过程中找到的最长可行前缀（即已选题目最多的路径），用于在无完整解时返回部分结果
    best_path = {'indices': [], 'length': 0}

    def can_still_satisfy(start, sel_diff, sel_kp, slots_left):
        # 检查从 start 到末尾的可用题目能否满足剩余需求
        for d, tgt in target_diff_counts.items():
            need = max(0, tgt - sel_diff.get(d, 0))
            if suffix_diff[start].get(d, 0) < need:
                return False
        for k, tgt in target_kp_counts.items():
            need = max(0, tgt - sel_kp.get(k, 0))
            if suffix_kp[start].get(k, 0) < need:
                return False
        # 还要保证总体可用数足够
        if sum(suffix_diff[start].values()) < slots_left:
            return False
        return True

    def dfs(start):
        # 递归深度优先搜索，尝试从 start 开始选择题目
        if found[0] is not None:
            return True
        if len(path) == target_total:
            # 校验是否完全满足目标计数
            for d, tgt in target_diff_counts.items():
                if sel_diff.get(d, 0) != tgt:
                    return False
            for k, tgt in target_kp_counts.items():
                if sel_kp.get(k, 0) != tgt:
                    return False
            found[0] = path.copy()
            return True
        # 更新当前已选的最佳前缀（尽量保存更多已选题）
        if len(path) > best_path['length']:
            best_path['indices'] = path.copy()
            best_path['length'] = len(path)
        if start >= M:
            return False

        slots_left = target_total - len(path)
        # 剪枝：如果后缀不能满足剩余需求则回溯
        if not can_still_satisfy(start, sel_diff, sel_kp, slots_left):
            return False

        # 依次尝试包含某个候选或跳过
        for i in range(start, M):
            lvl = levels[i]
            kp = kps[i]
            # 如果选了会超出某类目标，则跳过
            if sel_diff[lvl] + 1 > target_diff_counts.get(lvl, 0):
                continue
            if target_kp_counts and sel_kp[kp] + 1 > target_kp_counts.get(kp, 0):
                continue

            # 选择 i
            path.append(i)
            sel_diff[lvl] += 1
            sel_kp[kp] += 1

            if dfs(i + 1):
                return True

            # 回溯
            path.pop()
            sel_diff[lvl] -= 1
            sel_kp[kp] -= 1

        return False

    ok = dfs(0)
    if ok and found[0] is not None:
        return [candidates[i] for i in found[0]]
    # 若未找到完整解，返回在搜索过程中最长的已选题目序列（可能为空）
    if best_path['length'] > 0:
        return [candidates[i] for i in best_path['indices']]
    return None


def save_docx_real_paper(selected, out_path, knowledge_mapping, show_difficulty=True, show_knowledge=True, show_answer=True, show_analysis=True):
    """保存抽题结果到 Word 文档，根据参数控制输出内容。"""
    # 使用模块级 sanitize_option_text
    doc = Document()
    doc.add_heading('单选题（自动抽题结果）', level=1)
    doc.add_paragraph('一、单选题')
    qnum = 1
    prev_big = None
    for c in selected:
        # 获取大题题干和小题题干
        big = c.get('big_stem', '')
        small_stem = c.get('stem', '')
        
        # 构建题号和大题题干行（如果大题题干存在且与上一题不同）
        if big and big != prev_big:
            question_line = f"{qnum}、 {big}"
            doc.add_paragraph(question_line)
            prev_big = big
        else:
            # 如果没有大题题干或与上一题相同，只写题号
            question_line = f"{qnum}、"
            doc.add_paragraph(question_line)
        
        # 构建小题题干行，包含难度和知识点信息
        stem_line = f"      {small_stem}"
        if show_difficulty or show_knowledge:
            details = []
            if show_difficulty:
                details.append(f"难度：{c.get('level','')}")
            if show_knowledge:
                # 使用知识点映射将代码转换为名称
                kp_code = c.get('kp', '')
                kp_name = knowledge_mapping.get(kp_code, kp_code)  # 如果找不到映射，使用原代码
                details.append(f"知识点：{kp_name}")
            if details:
                stem_line += f"（{', '.join(details)}）"
        doc.add_paragraph(stem_line)

        # 选项列表
        opts = c.get('options', [])
        # 按 A,B,C... 写出选项
        for idx, opt in enumerate(opts):
            label = chr(ord('A') + idx)
            clean_opt = sanitize_option_text(opt)
            doc.add_paragraph(f"    {label}、{clean_opt}")
        
        # 在题目后保留一行空白供答题
        doc.add_paragraph("")
        
        # 根据参数决定是否显示答案和解析
        if show_answer and c.get('answer'):
            doc.add_paragraph(f"    答案：{c.get('answer')}")
        if show_analysis and c.get('analysis'):
            doc.add_paragraph(f"    解析：{c.get('analysis')}")
        
        # 在题目之间增加空行
        doc.add_paragraph("")
        qnum += 1
        
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)


def save_docx_mixed_paper(single_list, multi_list, out_path, knowledge_mapping, show_difficulty=True, show_knowledge=True, show_answer=True, show_analysis=True):
    """将单选题和多选题合并保存为一个文档，先写单选题（一）、再写多选题（二）。"""
    doc = Document()
    # 单选题部分
    doc.add_heading('单选题（自动抽题结果）', level=1)
    doc.add_paragraph('一、单选题')
    qnum = 1
    prev_big = None
    for c in single_list:
        big = c.get('big_stem', '')
        small_stem = c.get('stem', '')
        if big and big != prev_big:
            doc.add_paragraph(f"{qnum}、 {big}")
            prev_big = big
        else:
            doc.add_paragraph(f"{qnum}、")
        stem_line = f"      {small_stem}"
        if show_difficulty or show_knowledge:
            details = []
            if show_difficulty:
                details.append(f"难度：{c.get('level','')}")
            if show_knowledge:
                kp_code = c.get('kp','')
                kp_name = knowledge_mapping.get(kp_code, kp_code)
                details.append(f"知识点：{kp_name}")
            if details:
                stem_line += f"（{', '.join(details)}）"
        doc.add_paragraph(stem_line)
        opts = c.get('options', [])
        for idx, opt in enumerate(opts):
            label = chr(ord('A') + idx)
            clean_opt = sanitize_option_text(opt)
            doc.add_paragraph(f"    {label}、{clean_opt}")
        doc.add_paragraph("")
        if show_answer and c.get('answer'):
            doc.add_paragraph(f"    答案：{c.get('answer')}")
        if show_analysis and c.get('analysis'):
            doc.add_paragraph(f"    解析：{c.get('analysis')}")
        doc.add_paragraph("")
        qnum += 1

    # 多选题部分
    doc.add_heading('多选题（自动抽题结果）', level=1)
    doc.add_paragraph('二、多选题')
    qnum_multi = 1
    prev_big = None
    for c in multi_list:
        big = c.get('big_stem', '')
        small_stem = c.get('stem', '')
        if big and big != prev_big:
            doc.add_paragraph(f"{qnum_multi}、 {big}")
            prev_big = big
        else:
            doc.add_paragraph(f"{qnum_multi}、")
        stem_line = f"      {small_stem}"
        if show_difficulty or show_knowledge:
            details = []
            if show_difficulty:
                details.append(f"难度：{c.get('level','')}")
            if show_knowledge:
                kp_code = c.get('kp','')
                kp_name = knowledge_mapping.get(kp_code, kp_code)
                details.append(f"知识点：{kp_name}")
            if details:
                stem_line += f"（{', '.join(details)}）"
        doc.add_paragraph(stem_line)
        opts = c.get('options', [])
        for idx, opt in enumerate(opts):
            label = chr(ord('A') + idx)
            doc.add_paragraph(f"    {label}、{opt}")
        doc.add_paragraph("")
        if show_answer and c.get('answer'):
            doc.add_paragraph(f"    答案：{c.get('answer')}")
        if show_analysis and c.get('analysis'):
            doc.add_paragraph(f"    解析：{c.get('analysis')}")
        doc.add_paragraph("")
        qnum_multi += 1

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)


def format_selected_knowledge_points(kps: str, knowledge_mapping: Dict[str, str], max_items: int = 4) -> str:
    codes = [k.strip() for k in str(kps or '').split(',') if k.strip()]
    if not codes:
        return '未限制'

    names = []
    seen = set()
    for code in codes:
        name = knowledge_mapping.get(code, code)
        if name in seen:
            continue
        seen.add(name)
        names.append(name)

    if len(names) <= max_items:
        return '、'.join(names)
    return '、'.join(names[:max_items]) + '...'


def _build_count_ratio_text(counts: Dict[str, int], ordered_keys: List[str], total: int) -> str:
    parts = []
    for key in ordered_keys:
        value = int(counts.get(key, 0))
        pct = round((value / total) * 100) if total > 0 else 0
        parts.append(f"{key}{value}题({pct}%)")
    return '，'.join(parts)


def build_current_paper_stats(selection: List[Dict[str, Any]]) -> Dict[str, str]:
    total = len(selection or [])
    type_counts = {'单选': 0, '多选': 0}
    diff_counts = {'难': 0, '中': 0, '易': 0}
    kp_names = []
    seen_kps = set()

    for question in selection or []:
        qtype = _safe_str(question.get('type'))
        if '多选' in qtype:
            type_counts['多选'] += 1
        else:
            type_counts['单选'] += 1

        diff = _safe_str(question.get('level'))
        if diff in diff_counts:
            diff_counts[diff] += 1

        kp_name = _safe_str(question.get('kp_name') or question.get('kp'))
        if kp_name and kp_name not in seen_kps:
            seen_kps.add(kp_name)
            kp_names.append(kp_name)

    covered_kps = '、'.join(kp_names[:4]) + ('...' if len(kp_names) > 4 else '') if kp_names else '无'
    return {
        'total': f'当前试卷共 {total} 题',
        'type': _build_count_ratio_text(type_counts, ['单选', '多选'], total),
        'difficulty': _build_count_ratio_text(diff_counts, ['难', '中', '易'], total),
        'knowledge': covered_kps,
    }


def generate_from_namespace(args):
    """使用现有逻辑生成 docx 并返回 (保存路径, selected_list) 或 (None, None)。"""
    knowledge_mapping = load_knowledge_mapping(args.kp_file)

    diff_keys = ['难', '中', '易']
    diff_ratios = parse_ratio_arg(args.difficulty, diff_keys)
    target_diff_counts = compute_counts_from_ratios(args.number, diff_ratios)

    kps = [k.strip() for k in args.kps.split(',') if k.strip()] if getattr(args, 'kps', None) else []
    if not kps:
        allowed_kps = None
    else:
        allowed_kps = kps

    candidates = read_single_choice_xlsx(args.xlsx, allowed_kps=allowed_kps)

    # 知识点只用于过滤候选题目池，不再对各知识点设置精确抽题配额。
    target_kp_counts = {}

    total_available = len(candidates)
    if total_available < args.number:
        print(f"题库中可用题目 {total_available} 少于所需 {args.number}。无法满足。")
        return None

    selection = find_selection_backtrack(candidates, args.number, target_diff_counts, target_kp_counts)

    out_path = args.out
    if selection:
        save_docx_real_paper(selection, out_path, knowledge_mapping,
                             show_difficulty=bool(args.show_difficulty),
                             show_knowledge=bool(args.show_knowledge),
                             show_answer=bool(args.show_answer),
                             show_analysis=bool(args.show_analysis))
        return out_path, selection

    sel = []
    sel_diff = defaultdict(int)
    for c in candidates:
        if len(sel) >= args.number:
            break
        lvl = c['level']
        if sel_diff[lvl] + 1 > target_diff_counts.get(lvl, 0):
            continue
        sel.append(c)
        sel_diff[lvl] += 1

    if sel:
        save_docx_real_paper(sel, out_path, knowledge_mapping,
                             show_difficulty=bool(args.show_difficulty),
                             show_knowledge=bool(args.show_knowledge),
                             show_answer=bool(args.show_answer),
                             show_analysis=bool(args.show_analysis))
        return out_path, sel

    return None, None


def generate_selection_from_args(number, difficulty, kps, xlsx_path, kp_file, allowed_kps=None, is_multi=False):
    """基于参数从指定 xlsx 中选择题目（不保存文档），返回选择列表或 []。"""
    # 复用 generate_from_namespace 的选择逻辑但不保存
    knowledge_mapping = load_knowledge_mapping(kp_file)
    diff_keys = ['难', '中', '易']
    diff_ratios = parse_ratio_arg(difficulty, diff_keys)
    target_diff_counts = compute_counts_from_ratios(number, diff_ratios)

    kps_list = [k.strip() for k in kps.split(',') if k.strip()] if kps else []
    allowed_kps = kps_list if kps_list else None

    if is_multi:
        candidates = read_multi_choice_xlsx(xlsx_path, allowed_kps=allowed_kps)
    else:
        candidates = read_single_choice_xlsx(xlsx_path, allowed_kps=allowed_kps)

    if not candidates:
        return []

    # 知识点只负责过滤题目池，不参与精确配额。
    target_kp_counts = {}

    selection = find_selection_backtrack(candidates, number, target_diff_counts, target_kp_counts)
    if selection:
        return selection

    # 回退贪心
    sel = []
    sel_diff = defaultdict(int)
    for c in candidates:
        if len(sel) >= number:
            break
        lvl = c['level']
        if sel_diff[lvl] + 1 > target_diff_counts.get(lvl, 0):
            continue
        sel.append(c)
        sel_diff[lvl] += 1

    return sel


AGENT_TASK_DIR = os.path.join('output', 'agent_tasks')
AGENT_PLANNER_PROMPT_TEMPLATE = """你是试卷系统中的任务规划智能体。你的职责是把用户的自然语言需求转换为可以执行的结构化组卷参数。

请只返回 JSON，不要输出任何额外说明。JSON 格式如下：
{{
    "intent": "paper.generate",
    "title": "试卷标题",
    "num_questions": 10,
    "knowledge_point_codes": ["32090102"],
    "difficulty": {{"难": 10, "中": 20, "易": 70}},
    "multichoice_percent": 0,
    "show_difficulty": false,
    "show_knowledge": true,
    "show_answer": false,
    "show_analysis": false,
    "reason": "一句话说明为什么这样规划"
}}

约束：
1. intent 只能返回 paper.generate。
2. 如果用户没有明确题量，则 num_questions 返回 10。
3. difficulty 必须同时包含 难、中、易 三个键，三者总和为 100。
4. multichoice_percent 取 0-100 的整数。
5. knowledge_point_codes 只能从候选列表中选择。

可选知识点候选如下：
{knowledge_candidates}

用户需求：
{message}
"""


def build_output_web_path(file_path: Optional[str]) -> Optional[str]:
    if not file_path:
        return None
    return f"/output/{os.path.basename(file_path)}"


def parse_int_value(value: Any, default: int, minimum: Optional[int] = None, maximum: Optional[int] = None) -> int:
    try:
        result = int(value)
    except Exception:
        result = default

    if minimum is not None:
        result = max(minimum, result)
    if maximum is not None:
        result = min(maximum, result)
    return result


def normalize_bool_value(value: Any, default: bool = False) -> bool:
    if value is None:
        return default
    if isinstance(value, bool):
        return value
    text = str(value).strip().lower()
    if text in {'1', 'true', 'yes', 'y', 'on', '是', '显示', '需要'}:
        return True
    if text in {'0', 'false', 'no', 'n', 'off', '否', '不显示', '不要'}:
        return False
    return default


def build_agent_task_id(prefix: str = 'task') -> str:
    return f"{prefix}_{datetime.now().strftime('%Y%m%d_%H%M%S_%f')}"


def ensure_agent_task_dir() -> str:
    os.makedirs(AGENT_TASK_DIR, exist_ok=True)
    return AGENT_TASK_DIR


def save_agent_task_record(task_record: Dict[str, Any]) -> Dict[str, Any]:
    ensure_agent_task_dir()
    task_id = _safe_str(task_record.get('task_id')) or build_agent_task_id('agent')
    task_record['task_id'] = task_id
    task_record['updated_at'] = datetime.now().isoformat()
    task_path = os.path.join(AGENT_TASK_DIR, f'{task_id}.json')
    with open(task_path, 'w', encoding='utf-8') as f:
        json.dump(task_record, f, ensure_ascii=False, indent=2)
    return task_record


def load_agent_task_record(task_id: str) -> Optional[Dict[str, Any]]:
    safe_task_id = re.sub(r'[^A-Za-z0-9_\-]', '', str(task_id or ''))
    if not safe_task_id:
        return None
    task_path = os.path.join(AGENT_TASK_DIR, f'{safe_task_id}.json')
    if not os.path.exists(task_path):
        return None
    with open(task_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def get_agent_runtime_config():
    if config is not None:
        return config
    return RagConfig()


def build_agent_tool_catalog() -> List[Dict[str, Any]]:
    return [
        {
            'name': 'paper.generate',
            'group': 'paper',
            'description': '根据题量、知识点、难度和题型占比生成试卷，返回预览和导出地址。',
            'invoke': {'method': 'POST', 'path': '/api/agent/paper/generate'},
            'parameters': {
                'title': 'str，可选，试卷标题',
                'num_questions': 'int，可选，默认10',
                'knowledge_point_codes': 'string[]，可选，知识点代码列表',
                'difficulty': {'难': 'int', '中': 'int', '易': 'int'},
                'multichoice_percent': 'int，可选，0-100',
                'show_difficulty': 'bool，可选',
                'show_knowledge': 'bool，可选',
                'show_answer': 'bool，可选',
                'show_analysis': 'bool，可选',
            },
        },
        {
            'name': 'paper.preview',
            'group': 'paper',
            'description': '根据 task_id 获取智能体生成的试卷预览与执行轨迹。',
            'invoke': {'method': 'GET', 'path': '/api/agent/tasks/{task_id}'},
            'parameters': {'task_id': 'str，任务编号'},
        },
        {
            'name': 'paper.export',
            'group': 'paper',
            'description': '根据 task_id 重新导出试卷，并允许调整显示答案、解析等选项。',
            'invoke': {'method': 'POST', 'path': '/api/agent/tasks/{task_id}/export'},
            'parameters': {
                'task_id': 'str，任务编号',
                'show_difficulty': 'bool，可选',
                'show_knowledge': 'bool，可选',
                'show_answer': 'bool，可选',
                'show_analysis': 'bool，可选',
            },
        },
        {
            'name': 'paper.retry_backtrack',
            'group': 'paper',
            'description': '在已有试卷基础上排除已删题目并回溯补题。',
            'invoke': {'method': 'POST', 'path': '/retry-backtrack'},
        },
        {
            'name': 'paper.save_edited',
            'group': 'paper',
            'description': '保存用户编辑后的试卷结果为 docx。',
            'invoke': {'method': 'POST', 'path': '/save-edited-paper'},
        },
        {
            'name': 'question.generate_similar',
            'group': 'question',
            'description': '基于原题和知识库生成相似题目。',
            'invoke': {'method': 'POST', 'path': '/generate-similar-question'},
        },
        {
            'name': 'short_answer.generate_er_diagram',
            'group': 'short_answer',
            'description': '生成数据库场景的 E-R 图结构。',
            'invoke': {'method': 'POST', 'path': '/generate-er-diagram'},
        },
        {
            'name': 'short_answer.generate_question_stem',
            'group': 'short_answer',
            'description': '根据 E-R 图生成数据库综合题题干。',
            'invoke': {'method': 'POST', 'path': '/generate-question-stem'},
        },
        {
            'name': 'short_answer.export_docx',
            'group': 'short_answer',
            'description': '导出当前数据库综合题为 docx。',
            'invoke': {'method': 'POST', 'path': '/export-question-docx'},
        },
        {
            'name': 'question_bank.meta',
            'group': 'question_bank',
            'description': '获取题库筛选元数据。',
            'invoke': {'method': 'GET', 'path': '/api/question-bank/meta'},
        },
        {
            'name': 'question_bank.list',
            'group': 'question_bank',
            'description': '分页查询题库列表。',
            'invoke': {'method': 'GET', 'path': '/api/question-bank/questions'},
        },
        {
            'name': 'question_bank.get',
            'group': 'question_bank',
            'description': '查看单个题库题目详情。',
            'invoke': {'method': 'GET', 'path': '/api/question-bank/questions/{question_id}'},
        },
        {
            'name': 'question_bank.create',
            'group': 'question_bank',
            'description': '新增普通题目到题库。',
            'invoke': {'method': 'POST', 'path': '/api/question-bank/questions'},
        },
        {
            'name': 'question_bank.create_short_answer',
            'group': 'question_bank',
            'description': '新增综合简答题到题库。',
            'invoke': {'method': 'POST', 'path': '/api/question-bank/short-answer'},
        },
        {
            'name': 'question_bank.update',
            'group': 'question_bank',
            'description': '更新普通题目。',
            'invoke': {'method': 'PUT', 'path': '/api/question-bank/questions/{question_id}'},
        },
        {
            'name': 'question_bank.update_compound',
            'group': 'question_bank',
            'description': '更新综合简答题。',
            'invoke': {'method': 'PUT', 'path': '/api/question-bank/questions/{question_id}/compound'},
        },
        {
            'name': 'question_bank.export_docx',
            'group': 'question_bank',
            'description': '将已保存的综合简答题导出为 docx。',
            'invoke': {'method': 'POST', 'path': '/api/question-bank/questions/{question_id}/export-docx'},
        },
        {
            'name': 'rag.status',
            'group': 'rag',
            'description': '查询知识库向量集合状态。',
            'invoke': {'method': 'GET', 'path': '/rag/status/'},
        },
        {
            'name': 'rag.create',
            'group': 'rag',
            'description': '上传文档并写入知识库向量。',
            'invoke': {'method': 'POST', 'path': '/rag/create/'},
        },
        {
            'name': 'rag.chat',
            'group': 'rag',
            'description': '对知识库提问并返回答案。',
            'invoke': {'method': 'POST', 'path': '/rag/chat/'},
        },
        {
            'name': 'rag.clear',
            'group': 'rag',
            'description': '清空知识库集合与上传目录。',
            'invoke': {'method': 'POST', 'path': '/rag/clear/'},
        },
        {
            'name': 'uploads.list',
            'group': 'rag',
            'description': '列出已上传知识库文件。',
            'invoke': {'method': 'GET', 'path': '/uploads/'},
        },
        {
            'name': 'uploads.delete',
            'group': 'rag',
            'description': '删除已上传知识库文件与对应向量数据。',
            'invoke': {'method': 'DELETE', 'path': '/uploads/{filename}'},
        },
    ]


def build_agent_trace_entry(step: str, detail: str, status: str = 'completed') -> Dict[str, str]:
    return {
        'step': step,
        'detail': detail,
        'status': status,
        'timestamp': datetime.now().strftime('%H:%M:%S'),
    }


def infer_agent_title_from_text(message: str) -> str:
    text = str(message or '').strip()
    quoted = re.search(r'《([^》]{2,40})》', text)
    if quoted:
        return quoted.group(1)
    titled = re.search(r'(?:标题|卷名|名称)(?:为|是|叫)?[:：]?\s*([^，。；\n]{2,40})', text)
    if titled:
        return titled.group(1).strip()
    return '智能体生成试卷'


def infer_agent_question_count_from_text(message: str, default: int = 10) -> int:
    text = str(message or '')
    match = re.search(r'(\d{1,3})\s*(?:道|题)', text)
    if match:
        return parse_int_value(match.group(1), default, 1, 200)
    return default


def extract_named_percentages(text: str, labels: List[str]) -> Dict[str, int]:
    result: Dict[str, int] = {}
    for label in labels:
        pattern = re.compile(rf'{label}(?:题|度|占比|比例|约|大概|为|是|各占|分别占)?[^\d]{{0,6}}(\d{{1,3}})', re.IGNORECASE)
        match = pattern.search(text)
        if match:
            result[label] = parse_int_value(match.group(1), 0, 0, 100)
    return result


def normalize_percentage_triplet(raw_values: Dict[str, int], ordered_keys: List[str], defaults: Dict[str, int]) -> Dict[str, int]:
    merged = {key: parse_int_value(raw_values.get(key), defaults.get(key, 0), 0, 100) for key in ordered_keys}
    total = sum(merged.values())
    if total <= 0:
        merged = defaults.copy()
        total = sum(merged.values())

    normalized = compute_counts_from_ratios(100, {key: merged.get(key, 0) / total for key in ordered_keys})
    for key in ordered_keys:
        normalized[key] = int(normalized.get(key, 0))
    return normalized


def infer_agent_difficulty_from_text(message: str) -> Dict[str, int]:
    text = str(message or '')
    extracted = extract_named_percentages(text, ['难', '中', '易'])
    if not extracted:
        alias_map = {'难': ['困难'], '中': ['中等'], '易': ['简单', '容易']}
        for target, aliases in alias_map.items():
            for alias in aliases:
                pattern = re.compile(rf'{alias}(?:题|难度|占比|比例|约|大概|为|是|各占|分别占)?[^\d]{{0,6}}(\d{{1,3}})', re.IGNORECASE)
                match = pattern.search(text)
                if match:
                    extracted[target] = parse_int_value(match.group(1), 0, 0, 100)
                    break
    return normalize_percentage_triplet(extracted, ['难', '中', '易'], {'难': 10, '中': 20, '易': 70})


def infer_agent_multichoice_percent_from_text(message: str) -> int:
    text = str(message or '')
    if re.search(r'不需要多选|不要多选|全是单选|仅单选|只要单选', text):
        return 0
    match = re.search(r'多选(?:题)?(?:占比|比例|约|大概|为|是|各占|分别占)?[^\d]{0,6}(\d{1,3})', text)
    if match:
        return parse_int_value(match.group(1), 0, 0, 100)
    dual_match = re.search(r'单选(?:题)?[^\d]{0,6}(\d{1,3}).{0,20}?多选(?:题)?[^\d]{0,6}(\d{1,3})', text)
    if dual_match:
        single_value = parse_int_value(dual_match.group(1), 100, 0, 100)
        multi_value = parse_int_value(dual_match.group(2), 0, 0, 100)
        total = single_value + multi_value
        if total > 0:
            return round((multi_value / total) * 100)
    return 0


def match_knowledge_point_codes(message: str, knowledge_mapping: Dict[str, str]) -> List[str]:
    text = str(message or '').strip()
    if not text or re.search(r'不限知识点|任意知识点|全部知识点|不限制知识点', text):
        return []

    matched = []
    seen = set()
    candidates = sorted(knowledge_mapping.items(), key=lambda item: len(item[1]), reverse=True)
    for code, name in candidates:
        code_s = _safe_str(code)
        name_s = _safe_str(name)
        if not code_s or not name_s:
            continue
        if code_s in text or name_s in text:
            if code_s not in seen:
                seen.add(code_s)
                matched.append(code_s)
    return matched


def try_plan_paper_request_with_llm(message: str, matched_kps: List[str], knowledge_mapping: Dict[str, str]) -> Optional[Dict[str, Any]]:
    runtime_config = get_agent_runtime_config()
    if not getattr(runtime_config, 'api_key', None) or not getattr(runtime_config, 'base_url', None) or not getattr(runtime_config, 'llm_model_name', None):
        return None

    candidate_lines = []
    for code in matched_kps[:12]:
        candidate_lines.append(f'- {code}: {knowledge_mapping.get(code, code)}')
    if not candidate_lines:
        for code, name in list(knowledge_mapping.items())[:20]:
            candidate_lines.append(f'- {code}: {name}')

    prompt_text = AGENT_PLANNER_PROMPT_TEMPLATE.format(
        knowledge_candidates='\n'.join(candidate_lines) if candidate_lines else '- 无可用知识点候选',
        message=message,
    )

    try:
        llm = ChatOpenAI(
            model=runtime_config.llm_model_name,
            api_key=runtime_config.api_key,
            base_url=runtime_config.base_url,
            temperature=0,
        )
        response = llm.invoke(prompt_text)
        parsed = safe_json_loads(response.content, None)
        if not isinstance(parsed, dict):
            return None
        return parsed
    except Exception as e:
        print(f'智能体规划调用 LLM 失败，回退到规则解析: {e}')
        return None


def build_agent_paper_request_from_text(message: str, overrides: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    overrides = overrides or {}
    knowledge_mapping = load_knowledge_mapping('./知识点列表.xlsx')
    matched_kps = match_knowledge_point_codes(message, knowledge_mapping)
    llm_plan = try_plan_paper_request_with_llm(message, matched_kps, knowledge_mapping)

    if isinstance(llm_plan, dict) and _safe_str(llm_plan.get('intent')) == 'paper.generate':
        planned_codes = [code for code in (llm_plan.get('knowledge_point_codes') or []) if code in knowledge_mapping]
        difficulty_raw = llm_plan.get('difficulty') if isinstance(llm_plan.get('difficulty'), dict) else {}
        difficulty = normalize_percentage_triplet(
            {key: parse_int_value(difficulty_raw.get(key), 0, 0, 100) for key in ['难', '中', '易']},
            ['难', '中', '易'],
            {'难': 10, '中': 20, '易': 70}
        )
        request_payload = {
            'title': _safe_str(llm_plan.get('title')) or infer_agent_title_from_text(message),
            'num_questions': parse_int_value(llm_plan.get('num_questions'), infer_agent_question_count_from_text(message), 1, 200),
            'knowledge_point_codes': planned_codes or matched_kps,
            'difficulty': difficulty,
            'multichoice_percent': parse_int_value(llm_plan.get('multichoice_percent'), infer_agent_multichoice_percent_from_text(message), 0, 100),
            'show_difficulty': normalize_bool_value(llm_plan.get('show_difficulty'), False),
            'show_knowledge': normalize_bool_value(llm_plan.get('show_knowledge'), True),
            'show_answer': normalize_bool_value(llm_plan.get('show_answer'), True),
            'show_analysis': normalize_bool_value(llm_plan.get('show_analysis'), True),
            'reason': _safe_str(llm_plan.get('reason')),
        }
    else:
        request_payload = {
            'title': infer_agent_title_from_text(message),
            'num_questions': infer_agent_question_count_from_text(message),
            'knowledge_point_codes': matched_kps,
            'difficulty': infer_agent_difficulty_from_text(message),
            'multichoice_percent': infer_agent_multichoice_percent_from_text(message),
            'show_difficulty': False,
            'show_knowledge': True,
            'show_answer': True,
            'show_analysis': True,
            'reason': '规则解析生成默认组卷参数',
        }

    merged = {**request_payload, **{k: v for k, v in overrides.items() if v is not None}}
    merged['knowledge_point_codes'] = list(dict.fromkeys(merged.get('knowledge_point_codes') or []))
    merged['num_questions'] = parse_int_value(merged.get('num_questions'), 10, 1, 200)
    merged['multichoice_percent'] = parse_int_value(merged.get('multichoice_percent'), 0, 0, 100)
    merged['difficulty'] = normalize_percentage_triplet(
        merged.get('difficulty') if isinstance(merged.get('difficulty'), dict) else {},
        ['难', '中', '易'],
        {'难': 10, '中': 20, '易': 70},
    )
    for key, default in [('show_difficulty', False), ('show_knowledge', True), ('show_answer', True), ('show_analysis', True)]:
        merged[key] = normalize_bool_value(merged.get(key), default)
    return merged


def normalize_agent_paper_payload(payload: Dict[str, Any]) -> Dict[str, Any]:
    payload = payload or {}
    knowledge_mapping = load_knowledge_mapping('./知识点列表.xlsx')
    knowledge_codes = payload.get('knowledge_point_codes') or payload.get('kps') or []
    if isinstance(knowledge_codes, str):
        knowledge_codes = [item.strip() for item in knowledge_codes.split(',') if item.strip()]
    if not isinstance(knowledge_codes, list):
        knowledge_codes = []

    normalized_codes = []
    for item in knowledge_codes:
        item_s = _safe_str(item)
        if not item_s:
            continue
        if item_s in knowledge_mapping:
            normalized_codes.append(item_s)
            continue
        for code, name in knowledge_mapping.items():
            if item_s == name:
                normalized_codes.append(code)
                break

    difficulty = payload.get('difficulty')
    if isinstance(difficulty, str):
        try:
            ratio_map = parse_ratio_arg(difficulty, ['难', '中', '易'])
            difficulty = compute_counts_from_ratios(100, ratio_map)
        except Exception:
            difficulty = None
    if not isinstance(difficulty, dict):
        difficulty = {'难': 10, '中': 20, '易': 70}

    normalized = {
        'title': _safe_str(payload.get('title')) or '智能体生成试卷',
        'num_questions': parse_int_value(payload.get('num_questions'), 10, 1, 200),
        'knowledge_point_codes': list(dict.fromkeys(normalized_codes)),
        'difficulty': normalize_percentage_triplet(difficulty, ['难', '中', '易'], {'难': 10, '中': 20, '易': 70}),
        'multichoice_percent': parse_int_value(payload.get('multichoice_percent'), 0, 0, 100),
        'show_difficulty': normalize_bool_value(payload.get('show_difficulty'), False),
        'show_knowledge': normalize_bool_value(payload.get('show_knowledge'), True),
        'show_answer': normalize_bool_value(payload.get('show_answer'), True),
        'show_analysis': normalize_bool_value(payload.get('show_analysis'), True),
    }
    normalized['difficulty_text'] = f"难:{normalized['difficulty']['难']},中:{normalized['difficulty']['中']},易:{normalized['difficulty']['易']}"
    normalized['type_ratio_text'] = f"单选:{100 - normalized['multichoice_percent']},多选:{normalized['multichoice_percent']}"
    normalized['kps'] = ','.join(normalized['knowledge_point_codes'])
    normalized['kps_display'] = format_selected_knowledge_points(normalized['kps'], knowledge_mapping)
    return normalized


def enrich_questions_with_kp_name(selection: List[Dict[str, Any]], knowledge_mapping: Dict[str, str]) -> List[Dict[str, Any]]:
    enriched = []
    for item in selection or []:
        current = dict(item)
        kp_code = _safe_str(current.get('kp'))
        current['kp_name'] = knowledge_mapping.get(kp_code, kp_code)
        enriched.append(current)
    return enriched


def build_paper_preview_payload(
    normalized_payload: Dict[str, Any],
    out_path: str,
    selection_single: List[Dict[str, Any]],
    selection_multi: List[Dict[str, Any]],
) -> Dict[str, Any]:
    merged_selection = (selection_single or []) + (selection_multi or [])
    return {
        'title': normalized_payload['title'],
        'num_questions': normalized_payload['num_questions'],
        'kps': normalized_payload['kps'],
        'kps_display': normalized_payload['kps_display'],
        'difficulty': normalized_payload['difficulty_text'],
        'type_ratio': normalized_payload['type_ratio_text'],
        'show_difficulty': normalized_payload['show_difficulty'],
        'show_knowledge': normalized_payload['show_knowledge'],
        'show_answer': normalized_payload['show_answer'],
        'show_analysis': normalized_payload['show_analysis'],
        'out_path': build_output_web_path(out_path),
        'selection_single': selection_single or [],
        'selection_multi': selection_multi or [],
        'selection': merged_selection,
        'multi_exists': bool(selection_multi),
        'current_stats': build_current_paper_stats(merged_selection),
        'fulfilled_count': len(merged_selection),
    }


def build_generation_timing(start_time: float) -> Dict[str, Any]:
    elapsed_seconds = max(0.0, time.perf_counter() - start_time)
    return {
        'response_time_ms': int(round(elapsed_seconds * 1000)),
        'response_time_seconds': round(elapsed_seconds, 3),
    }


def print_generation_timing(
    normalized_payload: Dict[str, Any],
    timing: Dict[str, Any],
    success: bool,
    message: str = '',
    task_id: str = '',
    event_label: str = '组卷耗时',
) -> None:
    status_text = '成功' if success else '失败'
    title = _safe_str(normalized_payload.get('title')) or '未命名试卷'
    question_count = normalized_payload.get('num_questions')
    title_text = f'，标题={title}' if title else ''
    count_text = f'，题量={question_count}' if question_count is not None else ''
    task_text = f'，task_id={task_id}' if task_id else ''
    message_text = f'，信息={message}' if message else ''
    print(
        f"[{event_label}] 状态={status_text}{task_text}{title_text}{count_text}，耗时={timing['response_time_seconds']}秒（{timing['response_time_ms']}ms）{message_text}",
        flush=True,
    )


def build_llm_failure_message(error: Exception, action_label: str) -> str:
    raw_message = _safe_str(error)
    raw_lower = raw_message.lower()

    if 'insufficient balance' in raw_lower or ('402' in raw_message and 'balance' in raw_lower):
        return f'{action_label}失败：AI 服务账户余额不足，请检查当前 API Key 对应账户额度或充值后重试。'

    if 'invalid api key' in raw_lower or 'incorrect api key' in raw_lower or 'authentication' in raw_lower:
        return f'{action_label}失败：API Key 无效或认证失败，请检查后端模型配置。'

    if raw_message.startswith(f'{action_label}失败:') or raw_message.startswith(f'{action_label}失败：'):
        return raw_message

    return f'{action_label}失败: {raw_message}' if raw_message else f'{action_label}失败'


def generate_agent_paper_task(payload: Dict[str, Any], request_text: str = '') -> Dict[str, Any]:
    start_time = time.perf_counter()
    normalized = normalize_agent_paper_payload(payload)
    knowledge_mapping = load_knowledge_mapping('./知识点列表.xlsx')
    trace = [
        build_agent_trace_entry('analyze_request', '已解析组卷参数并完成标准化'),
        build_agent_trace_entry('select_candidates', f"知识点范围：{normalized['kps_display']}；目标题量：{normalized['num_questions']}"),
    ]

    safe_title = ''.join(c for c in normalized['title'] if c.isalnum() or c in (' ', '_')).strip().replace(' ', '_')[:50]
    ts = int(time.time())
    out_fname = f"{safe_title}_{ts}.docx" if safe_title else f"agent_paper_{ts}.docx"
    out_path = os.path.join('output', out_fname)

    multi_pct = normalized['multichoice_percent']
    selection_single: List[Dict[str, Any]] = []
    selection_multi: List[Dict[str, Any]] = []

    if multi_pct <= 0:
        selection_single = generate_selection_from_args(
            normalized['num_questions'],
            normalized['difficulty_text'],
            normalized['kps'],
            './单选题.xlsx',
            './知识点列表.xlsx',
            is_multi=False,
        )
        selection_single = enrich_questions_with_kp_name(selection_single, knowledge_mapping)
        if not selection_single:
            timing = build_generation_timing(start_time)
            message = '未能生成符合约束的单选题试卷，请检查知识点范围、题量或难度占比。'
            print_generation_timing(normalized, timing, success=False, message=message)
            return {
                'success': False,
                'message': message,
                'trace': trace,
                'timing': timing,
                **timing,
            }
        save_docx_real_paper(
            selection_single,
            out_path,
            knowledge_mapping,
            show_difficulty=normalized['show_difficulty'],
            show_knowledge=normalized['show_knowledge'],
            show_answer=normalized['show_answer'],
            show_analysis=normalized['show_analysis'],
        )
    else:
        multi_count = round(normalized['num_questions'] * (multi_pct / 100.0))
        multi_count = max(0, min(normalized['num_questions'], multi_count))
        single_count = normalized['num_questions'] - multi_count

        if single_count > 0:
            selection_single = generate_selection_from_args(
                single_count,
                normalized['difficulty_text'],
                normalized['kps'],
                './单选题.xlsx',
                './知识点列表.xlsx',
                is_multi=False,
            )
        if multi_count > 0:
            selection_multi = generate_selection_from_args(
                multi_count,
                normalized['difficulty_text'],
                normalized['kps'],
                './多选题.xlsx',
                './知识点列表.xlsx',
                is_multi=True,
            )

        selection_single = enrich_questions_with_kp_name(selection_single, knowledge_mapping)
        selection_multi = enrich_questions_with_kp_name(selection_multi, knowledge_mapping)
        if not selection_single and not selection_multi:
            timing = build_generation_timing(start_time)
            message = '未能生成符合约束的试卷，请检查题量、知识点或题型占比。'
            print_generation_timing(normalized, timing, success=False, message=message)
            return {
                'success': False,
                'message': message,
                'trace': trace,
                'timing': timing,
                **timing,
            }
        save_docx_mixed_paper(
            selection_single,
            selection_multi,
            out_path,
            knowledge_mapping,
            show_difficulty=normalized['show_difficulty'],
            show_knowledge=normalized['show_knowledge'],
            show_answer=normalized['show_answer'],
            show_analysis=normalized['show_analysis'],
        )

    timing = build_generation_timing(start_time)
    trace.append(build_agent_trace_entry('timing', f"组卷耗时 {timing['response_time_seconds']} 秒（{timing['response_time_ms']} ms）"))
    trace.append(build_agent_trace_entry('export_docx', f'试卷已生成并导出到 {build_output_web_path(out_path)}'))
    preview = build_paper_preview_payload(normalized, out_path, selection_single, selection_multi)
    task_record = save_agent_task_record({
        'task_id': build_agent_task_id('paper'),
        'task_type': 'paper.generate',
        'tool': 'paper.generate',
        'request_text': request_text,
        'normalized_request': normalized,
        'preview': preview,
        'selection_single': selection_single,
        'selection_multi': selection_multi,
        'trace': trace,
        'export': {
            'file_url': build_output_web_path(out_path),
            'file_path': out_path,
        },
        'timing': timing,
        **timing,
        'created_at': datetime.now().isoformat(),
    })
    print_generation_timing(normalized, timing, success=True, task_id=task_record['task_id'])

    return {
        'success': True,
        'task_id': task_record['task_id'],
        'trace': trace,
        'preview': preview,
        'export': task_record['export'],
        'normalized_request': normalized,
        'timing': timing,
        **timing,
    }


def export_agent_paper_task(task_record: Dict[str, Any], export_options: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    export_options = export_options or {}
    normalized_request = task_record.get('normalized_request', {}) or {}
    selection_single = task_record.get('selection_single', []) or []
    selection_multi = task_record.get('selection_multi', []) or []
    merged_selection = (selection_single or []) + (selection_multi or [])
    if not merged_selection:
        return {'success': False, 'message': '任务中没有可导出的题目内容'}

    show_difficulty = normalize_bool_value(export_options.get('show_difficulty'), normalized_request.get('show_difficulty', False))
    show_knowledge = normalize_bool_value(export_options.get('show_knowledge'), normalized_request.get('show_knowledge', True))
    show_answer = normalize_bool_value(export_options.get('show_answer'), normalized_request.get('show_answer', True))
    show_analysis = normalize_bool_value(export_options.get('show_analysis'), normalized_request.get('show_analysis', True))

    knowledge_mapping = load_knowledge_mapping('./知识点列表.xlsx')
    title = _safe_str(normalized_request.get('title')) or '智能体导出试卷'
    safe_title = ''.join(c for c in title if c.isalnum() or c in (' ', '_')).strip().replace(' ', '_')[:50]
    ts = int(time.time())
    out_fname = f"{safe_title}_agent_export_{ts}.docx" if safe_title else f"agent_export_{ts}.docx"
    out_path = os.path.join('output', out_fname)

    if selection_multi:
        save_docx_mixed_paper(
            selection_single,
            selection_multi,
            out_path,
            knowledge_mapping,
            show_difficulty=show_difficulty,
            show_knowledge=show_knowledge,
            show_answer=show_answer,
            show_analysis=show_analysis,
        )
    else:
        save_docx_real_paper(
            selection_single,
            out_path,
            knowledge_mapping,
            show_difficulty=show_difficulty,
            show_knowledge=show_knowledge,
            show_answer=show_answer,
            show_analysis=show_analysis,
        )

    task_record['export'] = {
        'file_url': build_output_web_path(out_path),
        'file_path': out_path,
        'show_difficulty': show_difficulty,
        'show_knowledge': show_knowledge,
        'show_answer': show_answer,
        'show_analysis': show_analysis,
    }
    task_record.setdefault('trace', []).append(build_agent_trace_entry('re_export', f'已按新的显示选项重新导出 {build_output_web_path(out_path)}'))
    save_agent_task_record(task_record)
    return {'success': True, 'export': task_record['export'], 'task_id': task_record.get('task_id')}


def main():
    parser = argparse.ArgumentParser(description='从单选题.xlsx 回溯抽题，尝试满足用户的占比要求')
    parser.add_argument('-n', '--number', type=int, default=10, help='题目总数（所有知识点题目数量之和）')
    parser.add_argument('--difficulty', type=str, default='难:10,中:20,易:70', help='难度占比')
    parser.add_argument('--types', type=str, default='单选:100', help='题型占比（该脚本仅处理单选题表）')
    parser.add_argument('--kps', type=str, default='32090102,32030303,32080102,32090201,32100401', help='知识点列表，用逗号分隔')
    parser.add_argument('--xlsx', type=str, default='./单选题.xlsx', help='单选题 xlsx 文件路径')##默认，用户不能修改
    parser.add_argument('--kp-file', type=str, default='./知识点列表.xlsx', help='知识点映射文件路径')##默认，用户不能修改
    parser.add_argument('--out', type=str, default='./output/from_xlsx.docx', help='输出 docx 路径')##默认，用户不能修改
    parser.add_argument('--show-difficulty', type=int, default=0, choices=[0, 1], help='是否显示难度 (1=显示, 0=不显示)')
    parser.add_argument('--show-knowledge', type=int, default=1, choices=[0, 1], help='是否显示知识点 (1=显示, 0=不显示)')
    parser.add_argument('--show-answer', type=int, default=1, choices=[0, 1], help='是否后显示答案 (1=显示, 0=不显示)')
    parser.add_argument('--show-analysis', type=int, default=1, choices=[0, 1], help='是否显示解析 (1=显示, 0=不显示)')

    args = parser.parse_args()

    # 加载知识点映射
    knowledge_mapping = load_knowledge_mapping(args.kp_file)

    diff_keys = ['难', '中', '易']
    diff_ratios = parse_ratio_arg(args.difficulty, diff_keys)
    target_diff_counts = compute_counts_from_ratios(args.number, diff_ratios)

    # 忽略题型参数，因为本脚本仅处理单选题；确保请求的是单选题
    # 知识点
    kps = [k.strip() for k in args.kps.split(',') if k.strip()]
    if not kps:
        # 稍后从工作簿收集知识点；暂时允许所有
        allowed_kps = None
    else:
        allowed_kps = kps

    candidates = read_single_choice_xlsx(args.xlsx, allowed_kps=allowed_kps)

    # 知识点仅用于过滤候选池，不再额外施加每个知识点的精确计数约束。
    target_kp_counts = {}

    # 快速可行性检查
    total_available = len(candidates)
    if total_available < args.number:
        print(f"题库中可用题目 {total_available} 少于所需 {args.number}。无法满足。")
        # 报告按难度的可用数量
        avail_diff = Counter([c['level'] for c in candidates])
        print('题库按难度可用：', dict(avail_diff))
        return

    # 如果候选题不包含所需的难度类别，报告短缺
    avail_diff = Counter([c['level'] for c in candidates])
    shortages = {}
    for d, tgt in target_diff_counts.items():
        if avail_diff.get(d, 0) < tgt:
            shortages[d] = tgt - avail_diff.get(d, 0)

    if shortages:
        print('初步检测发现题库中缺少以下题目（按难度）：')
        if shortages:
            print('按难度缺少：', shortages)
        print('仍然会尝试回溯搜索可行解（若存在）……')

    print('候选题数：', total_available)
    print('目标难度计数：', target_diff_counts)

    selection = find_selection_backtrack(candidates, args.number, target_diff_counts, target_kp_counts)

    if selection:
        print('找到可行的题目组合，数量：', len(selection))
        save_docx_real_paper(selection, args.out, knowledge_mapping,
                             show_difficulty=bool(args.show_difficulty),
                             show_knowledge=bool(args.show_knowledge),
                             show_answer=bool(args.show_answer),
                             show_analysis=bool(args.show_analysis))
        print('已保存到', args.out)
    else:
        # 未找到完整解时，使用贪心填充作为回退：按候选顺序依次加入不超目标的题目，直到不能再加入
        print('未找到满足全部约束的题目组合，尝试生成部分试卷（贪心回退）。')
        sel = []
        sel_diff = defaultdict(int)
        for c in candidates:
            if len(sel) >= args.number:
                break
            lvl = c['level']
            if sel_diff[lvl] + 1 > target_diff_counts.get(lvl, 0):
                continue
            sel.append(c)
            sel_diff[lvl] += 1

        if sel:
            print(f'已生成部分试卷，共加入 {len(sel)} 道题（将保存并在文档中列出）。')
            save_docx_real_paper(sel, args.out, knowledge_mapping,
                                 show_difficulty=bool(args.show_difficulty),
                                 show_knowledge=bool(args.show_knowledge),
                                 show_answer=bool(args.show_answer),
                                 show_analysis=bool(args.show_analysis))
            print('部分试卷已保存到', args.out)
        else:
            print('无法生成任何符合约束的题目（所有候选均会导致超出目标）。')

        # 报告题库中实际缺少项
        pool_diff = Counter([c['level'] for c in candidates])
        need_report = {}
        for d, tgt in target_diff_counts.items():
            have = pool_diff.get(d, 0)
            if have < tgt:
                need_report.setdefault('difficulty', {})[d] = tgt - have
        if need_report:
            print('题库不足以满足的项（需补充数量）：', need_report)
        else:
            print('题库看似有足够题量，但无组合能同时满足所有约束（可能是难度分布冲突）。')


if __name__ == '__main__':
    import sys

    # 默认行为：如果未提供额外的 CLI 参数，则运行 Web UI。
    # 如果用户提供了任何位置 CLI 参数（例如，通过 CLI 运行生成），则运行 `main()`。
    # 也支持显式使用 `--web` 强制 Web 模式和 `--cli` 强制 CLI 模式。
    if ('--cli' not in sys.argv) and (len(sys.argv) == 1 or '--web' in sys.argv):
        from fastapi import Form
        from fastapi.responses import HTMLResponse
        from fastapi.templating import Jinja2Templates
        import uvicorn

        templates = Jinja2Templates(directory='html')

        # 创建合并后的 FastAPI 应用
        web_app = FastAPI(
            title="试卷生成与知识问答系统",
            description="集成试卷生成和 RAG 知识问答功能的统一API"
        )
        
        # 添加 CORS 中间件支持跨域访问
        web_app.add_middleware(
            CORSMiddleware,
            allow_origins=["*"],
            allow_credentials=True,
            allow_methods=["*"],
            allow_headers=["*"],
        )
        
        web_app.mount('/static', StaticFiles(directory='html'), name='static')
        web_app.mount('/assets', StaticFiles(directory='html/assets'), name='assets')
        # 公开生成的文件以供下载
        os.makedirs('output', exist_ok=True)
        web_app.mount('/output', StaticFiles(directory='output'), name='output')

        get_current_user = register_auth_and_agent_routes(
            web_app,
            templates,
            auth_db,
            build_agent_tool_catalog=build_agent_tool_catalog,
            safe_str=_safe_str,
            generate_agent_paper_task=generate_agent_paper_task,
            load_agent_task_record=load_agent_task_record,
            export_agent_paper_task=export_agent_paper_task,
            build_agent_paper_request_from_text=build_agent_paper_request_from_text,
            build_agent_trace_entry=build_agent_trace_entry,
        )

        # ============= 原有路由（添加认证检查）=============
        @web_app.get('/', response_class=HTMLResponse)
        async def form_page(request: Request, session_id: Optional[str] = Cookie(None)):
            # 检查用户是否登录
            user = get_current_user(session_id)
            if not user:
                return RedirectResponse(url='/login', status_code=303)
            
            # 解析知识点层级以构建前端树形结构
            kp_file = './知识点列表.xlsx'
            kp_tree, kp_map = parse_kp_hierarchy(kp_file)
            # load counts from DB and attach to points
            name_counts, code_counts = load_kp_counts()
            # annotate kp_tree's points with counts (prefer code->count, fallback to name->count)
            for area in kp_tree:
                # accumulate area total from units
                area_total = 0
                for unit in area.get('units', []):
                    unit_total = 0
                    for p in unit.get('points', []):
                        cnt = 0
                        if p.get('code') and p.get('code') in code_counts:
                            cnt = code_counts[p.get('code')]
                        elif p.get('name') and p.get('name') in name_counts:
                            cnt = name_counts[p.get('name')]
                        # ensure every point has a numeric count (0 allowed)
                        p['count'] = int(cnt) if cnt is not None else 0
                        unit_total += p['count']
                    # set unit-level aggregated count (sum of its points)
                    unit['count'] = int(unit_total)
                    area_total += unit_total
                # set area-level aggregated count (sum of its units)
                area['count'] = int(area_total)

            return templates.TemplateResponse('index.html', {
                'request': request,
                'kp_tree': kp_tree,
                'kp_map_json': json.dumps(kp_map, ensure_ascii=False),
                'user': user
            })

        @web_app.post('/submit', response_class=HTMLResponse)
        async def submit(
            request: Request,
            session_id: Optional[str] = Cookie(None),
            title: str = Form(...),
            num_questions: int = Form(10),
            kps: str = Form('32090102,32030303,32080102,32090201,32100401'),
            difficulty_nan: str = Form('10'),
            difficulty_zhong: str = Form('20'),
            difficulty_yi: str = Form('70'),
            show_difficulty: str = Form(None),
            show_knowledge: str = Form(None),
            show_answer: str = Form(None),
            show_analysis: str = Form(None),
            include_multichoice: Optional[str] = Form(None),
            multichoice_percent: int = Form(0),
        ):
            # 检查用户是否登录
            user = get_current_user(session_id)
            if not user:
                return RedirectResponse(url='/login', status_code=303)

            start_time = time.perf_counter()
            
            # 构建 CLI 期望的难度字符串格式：'难:10,中:20,易:70'
            difficulty = f"难:{difficulty_nan},中:{difficulty_zhong},易:{difficulty_yi}"

            # 是否包含多选题
            want_multi = bool(include_multichoice)
            try:
                multi_pct = int(multichoice_percent) if multichoice_percent is not None else 0
            except Exception:
                multi_pct = 0
            if multi_pct < 0:
                multi_pct = 0
            if multi_pct > 100:
                multi_pct = 100
            single_pct = 100 - multi_pct
            type_ratio = f"单选:{single_pct},多选:{multi_pct}"

            # 准备生成器的参数命名空间
            import argparse
            safe_title = ''.join(c for c in title if c.isalnum() or c in (' ', '_')).strip().replace(' ', '_')[:50]
            ts = int(time.time())
            out_fname = f"{safe_title}_{ts}.docx" if safe_title else f"paper_{ts}.docx"
            out_path = os.path.join('output', out_fname)

            args = argparse.Namespace()
            args.number = num_questions
            args.difficulty = difficulty
            args.types = '单选:100'
            args.kps = kps
            args.xlsx = './单选题.xlsx'
            args.kp_file = './知识点列表.xlsx'
            args.out = out_path
            args.show_difficulty = 1 if show_difficulty else 0
            args.show_knowledge = 1 if show_knowledge else 0
            args.show_answer = 1 if show_answer else 0
            args.show_analysis = 1 if show_analysis else 0

            # 如果不需要多选题，沿用既有逻辑
            if not want_multi or multi_pct <= 0:
                saved_path, selection = generate_from_namespace(args)

                # enrich selection with human-readable kp name
                km = load_knowledge_mapping(args.kp_file)
                if selection:
                    for c in selection:
                        c['kp_name'] = km.get(c.get('kp',''), c.get('kp',''))
                current_selection = selection or []
                data = {
                    'title': title,
                    'num_questions': num_questions,
                    'kps': kps,
                    'kps_display': format_selected_knowledge_points(kps, km),
                    'difficulty': difficulty,
                    'type_ratio': type_ratio,
                    'current_stats': build_current_paper_stats(current_selection),
                    'show_difficulty': bool(args.show_difficulty),
                    'show_knowledge': bool(args.show_knowledge),
                    'show_answer': bool(args.show_answer),
                    'show_analysis': bool(args.show_analysis),
                    'out_path': f'/output/{os.path.basename(saved_path)}' if saved_path else None,
                    'selection': selection,
                }
                submit_timing = build_generation_timing(start_time)
                print_generation_timing(
                    {'title': title, 'num_questions': num_questions},
                    submit_timing,
                    success=bool(saved_path),
                    message='' if saved_path else '未生成试卷文件',
                )
                return templates.TemplateResponse('result.html', {'request': request, 'data': data})

            # 否则需要同时生成单选与多选两部分
            # 计算多选题数量（按百分比）
            multi_count = round(num_questions * (multi_pct / 100.0))
            if multi_count < 0:
                multi_count = 0
            single_count = num_questions - multi_count

            # 先生成单选部分
            args_single = argparse.Namespace(**vars(args))
            args_single.number = single_count if single_count > 0 else 0
            args_single.xlsx = './单选题.xlsx'
            args_single.out = os.path.join('output', f"{safe_title}_single_{ts}.docx")
            saved_single_path, selection_single = (None, [])
            if args_single.number > 0:
                selection_single = generate_selection_from_args(args_single.number, args_single.difficulty, args_single.kps, args_single.xlsx, args_single.kp_file, allowed_kps=None, is_multi=False)
                saved_single_path = None

            # 再生成多选部分
            args_multi = argparse.Namespace(**vars(args))
            args_multi.number = multi_count if multi_count > 0 else 0
            args_multi.xlsx = './多选题.xlsx'
            args_multi.out = os.path.join('output', f"{safe_title}_multi_{ts}.docx")
            saved_multi_path, selection_multi = (None, [])
            if args_multi.number > 0:
                # 使用专门的选择函数从多选题表中挑题（不直接保存）
                selection_multi = generate_selection_from_args(args_multi.number, args_multi.difficulty, args_multi.kps, args_multi.xlsx, args_multi.kp_file, allowed_kps=None, is_multi=True)
                saved_multi_path = None

            # 合并并保存最终文档
            km = load_knowledge_mapping(args.kp_file)
            # Ensure kp_name for each
            for c in (selection_single or []):
                c['kp_name'] = km.get(c.get('kp',''), c.get('kp',''))
            for c in (selection_multi or []):
                c['kp_name'] = km.get(c.get('kp',''), c.get('kp',''))

            # 最终保存为单一文档
            final_out = out_path
            save_docx_mixed_paper(selection_single or [], selection_multi or [], final_out, km,
                                 show_difficulty=bool(args.show_difficulty),
                                 show_knowledge=bool(args.show_knowledge),
                                 show_answer=bool(args.show_answer),
                                 show_analysis=bool(args.show_analysis))

            data = {
                'title': title,
                'num_questions': num_questions,
                'kps': kps,
                'kps_display': format_selected_knowledge_points(kps, km),
                'difficulty': difficulty,
                'type_ratio': type_ratio,
                'current_stats': build_current_paper_stats((selection_single or []) + (selection_multi or [])),
                'show_difficulty': bool(args.show_difficulty),
                'show_knowledge': bool(args.show_knowledge),
                'show_answer': bool(args.show_answer),
                'show_analysis': bool(args.show_analysis),
                'out_path': f'/output/{os.path.basename(final_out)}',
                'selection_single': selection_single,
                'selection_multi': selection_multi,
                'selection': (selection_single or []) + (selection_multi or []),
                'multi_exists': True,
            }
            submit_timing = build_generation_timing(start_time)
            print_generation_timing(
                {'title': title, 'num_questions': num_questions},
                submit_timing,
                success=bool(final_out),
                message='' if final_out else '未生成混合试卷文件',
            )
            return templates.TemplateResponse('result.html', {'request': request, 'data': data})
            # load mapping to present kp names in preview
            km = load_knowledge_mapping(args.kp_file)
            if selection:
                # enrich selection with human-readable kp name
                for c in selection:
                    kp_code = c.get('kp', '')
                    c['kp_name'] = km.get(kp_code, kp_code)
            data = {
                'title': title,
                'num_questions': num_questions,
                'kps': kps,
                'difficulty': difficulty,
                'show_difficulty': bool(args.show_difficulty),
                'show_knowledge': bool(args.show_knowledge),
                'show_answer': bool(args.show_answer),
                'show_analysis': bool(args.show_analysis),
                'out_path': f'/output/{os.path.basename(saved_path)}' if saved_path else None,
                'selection': selection,
            }
            return templates.TemplateResponse('result.html', {'request': request, 'data': data})

        @web_app.post('/save-edited-paper', response_class=JSONResponse)
        async def save_edited_paper(request: Request, session_id: Optional[str] = Cookie(None)):
            """保存编辑后的试卷"""
            # 检查用户是否登录
            user = get_current_user(session_id)
            if not user:
                return JSONResponse({'success': False, 'message': '请先登录'})
            
            try:
                body = await request.json()
                selection = body.get('selection', [])
                title = body.get('title', '编辑后试卷')
                show_difficulty = body.get('show_difficulty', False)
                show_knowledge = body.get('show_knowledge', True)
                show_answer = body.get('show_answer', True)
                show_analysis = body.get('show_analysis', True)
                
                if not selection:
                    return JSONResponse({'success': False, 'message': '没有题目可保存'})
                
                # 生成新的文件名
                import time
                safe_title = ''.join(c for c in title if c.isalnum() or c in (' ', '_')).strip().replace(' ', '_')[:50]
                ts = int(time.time())
                out_fname = f"{safe_title}_edited_{ts}.docx" if safe_title else f"paper_edited_{ts}.docx"
                out_path = os.path.join('output', out_fname)
                
                # 加载知识点映射
                knowledge_mapping = load_knowledge_mapping('./知识点列表.xlsx')
                
                # 保存为 Word 文档
                # 如果前端传来了分区（selection_single/selection_multi）且 multi_exists 为 True，使用分区保存函数
                selection_single = body.get('selection_single')
                selection_multi = body.get('selection_multi')
                multi_exists = body.get('multi_exists', False)

                if multi_exists and (selection_single is not None or selection_multi is not None):
                    # 使用分区保存（先单选，再多选）
                    save_docx_mixed_paper(
                        selection_single or [],
                        selection_multi or [],
                        out_path,
                        knowledge_mapping,
                        show_difficulty=show_difficulty,
                        show_knowledge=show_knowledge,
                        show_answer=show_answer,
                        show_analysis=show_analysis
                    )
                else:
                    save_docx_real_paper(
                        selection,
                        out_path,
                        knowledge_mapping,
                        show_difficulty=show_difficulty,
                        show_knowledge=show_knowledge,
                        show_answer=show_answer,
                        show_analysis=show_analysis
                    )
                
                return JSONResponse({
                    'success': True, 
                    'out_path': f'/output/{os.path.basename(out_path)}',
                    'message': '保存成功'
                })
            except Exception as e:
                import traceback
                traceback.print_exc()
                return JSONResponse({'success': False, 'message': str(e)})

        @web_app.post('/retry-backtrack', response_class=JSONResponse)
        async def retry_backtrack(request: Request, session_id: Optional[str] = Cookie(None)):
            """再次回溯组题，排除已删除的题目"""
            # 检查用户是否登录
            user = get_current_user(session_id)
            if not user:
                return JSONResponse({'success': False, 'message': '请先登录'})
            
            try:
                body = await request.json()
                num_questions = body.get('num_questions', 0)
                desired_single = body.get('desired_single', None)
                desired_multi = body.get('desired_multi', None)
                kps = body.get('kps', '')
                difficulty = body.get('difficulty', '难:10,中:20,易:70')
                excluded_ids = body.get('excluded_ids', [])
                existing_questions_single = body.get('existing_questions_single', [])
                existing_questions_multi = body.get('existing_questions_multi', [])
                title = body.get('title', '补充试卷')
                show_difficulty = body.get('show_difficulty', False)
                show_knowledge = body.get('show_knowledge', True)
                show_answer = body.get('show_answer', True)
                show_analysis = body.get('show_analysis', True)
                
                if num_questions <= 0:
                    return JSONResponse({'success': False, 'message': '不需要补充题目'})
                
                # 准备参数
                import argparse
                import time
                args = argparse.Namespace()
                args.number = num_questions
                args.difficulty = difficulty
                args.types = '单选:100'
                args.kps = kps
                args.xlsx = './单选题.xlsx'
                args.kp_file = './知识点列表.xlsx'
                
                # 临时输出路径
                safe_title = ''.join(c for c in title if c.isalnum() or c in (' ', '_')).strip().replace(' ', '_')[:50]
                ts = int(time.time())
                out_fname = f"{safe_title}_retry_{ts}.docx" if safe_title else f"paper_retry_{ts}.docx"
                args.out = os.path.join('output', out_fname)
                
                args.show_difficulty = 1 if show_difficulty else 0
                args.show_knowledge = 1 if show_knowledge else 0
                args.show_answer = 1 if show_answer else 0
                args.show_analysis = 1 if show_analysis else 0
                
                # 读取候选题目并分别为单选/多选分配
                knowledge_mapping = load_knowledge_mapping(args.kp_file)
                diff_keys = ['难', '中', '易']

                # helper to select from a set of candidates with exclusion
                def select_from_candidates(candidates, desired_number):
                    if desired_number <= 0:
                        return []
                    # 解析难度目标
                    diff_ratios = parse_ratio_arg(difficulty, diff_keys)
                    target_diff_counts = compute_counts_from_ratios(desired_number, diff_ratios)

                    target_kp_counts = {}

                    sel = find_selection_backtrack(candidates, desired_number, target_diff_counts, target_kp_counts)
                    if sel:
                        return sel

                    # 贪心回退
                    sel = []
                    sel_diff = defaultdict(int)
                    for c in candidates:
                        if len(sel) >= desired_number:
                            break
                        lvl = c.get('level')
                        if sel_diff[lvl] + 1 > target_diff_counts.get(lvl, 0):
                            continue
                        sel.append(c)
                        sel_diff[lvl] += 1
                    return sel

                # 准备现有题目列表
                existing_single = existing_questions_single or []
                existing_multi = existing_questions_multi or []

                # 读取并过滤单选候选
                kps_list = [k.strip() for k in kps.split(',') if k.strip()] if kps else []
                allowed_kps = kps_list if kps_list else None
                single_candidates = read_single_choice_xlsx('./单选题.xlsx', allowed_kps=allowed_kps)
                multi_candidates = read_multi_choice_xlsx('./多选题.xlsx', allowed_kps=allowed_kps)

                # 过滤掉已排除的题目（通过题干+选项匹配）
                def get_question_id(q):
                    stem = q.get('stem', '')
                    options = '|'.join(q.get('options', []))
                    return f"{stem}_{options}"

                filtered_single = [c for c in single_candidates if get_question_id(c) not in excluded_ids]
                filtered_multi = [c for c in multi_candidates if get_question_id(c) not in excluded_ids]

                # 计算需要补充的各类型数量
                ds = desired_single if desired_single is not None else 0
                dm = desired_multi if desired_multi is not None else 0

                new_single = select_from_candidates(filtered_single, ds)
                new_multi = select_from_candidates(filtered_multi, dm)

                # 为新题添加知识点名称
                for c in new_single + new_multi:
                    kp_code = c.get('kp', '')
                    c['kp_name'] = knowledge_mapping.get(kp_code, kp_code)

                # 合并现有题与新题并保存为分区文档
                merged_single = existing_single + new_single
                merged_multi = existing_multi + new_multi
                # 保存为分区文档（单选 + 多选）
                if merged_single or merged_multi:
                    save_docx_mixed_paper(merged_single, merged_multi, args.out, knowledge_mapping,
                                           show_difficulty=bool(args.show_difficulty),
                                           show_knowledge=bool(args.show_knowledge),
                                           show_answer=bool(args.show_answer),
                                           show_analysis=bool(args.show_analysis))

                return JSONResponse({
                    'success': True,
                    'new_single': new_single,
                    'new_multi': new_multi,
                    'out_path': f'/output/{os.path.basename(args.out)}',
                    'message': f'成功补充 单选 {len(new_single)} 道，多选 {len(new_multi)} 道'
                })
                
            except Exception as e:
                import traceback
                traceback.print_exc()
                return JSONResponse({'success': False, 'message': str(e)})

        @web_app.post('/generate-similar-question')
        async def generate_similar_question(request: Request, session_id: Optional[str] = Cookie(None)):
            """
            AI 生成类似题目
            """
            # 检查用户是否登录
            user = get_current_user(session_id)
            if not user:
                return JSONResponse({'success': False, 'message': '请先登录'})
            
            try:
                data = await request.json()
                original_question = data.get('original_question', {})
                user_requirement = data.get('user_requirement', '无特殊要求')
                
                # 提取原题信息
                original_stem = original_question.get('stem', '')
                original_big_stem = original_question.get('big_stem', '')
                original_options = original_question.get('options', [])
                original_answer = original_question.get('answer', '')
                original_analysis = original_question.get('analysis', '')
                knowledge_point = original_question.get('knowledge_point', '')
                difficulty = original_question.get('difficulty', '中')
                effective_original_text = original_stem or original_big_stem
                query_option_texts = [sanitize_option_text(opt) for opt in original_options if str(opt or '').strip()]
                
                # 如果没有 RAG 功能，返回错误
                if not RAG_AVAILABLE:
                    return JSONResponse({
                        'success': False,
                        'message': 'RAG 功能未启用，无法使用 AI 生成'
                    })
                
                # 使用知识点进行 RAG 检索
                if query_option_texts:
                    query_text = ' '.join([part for part in [knowledge_point] + query_option_texts if str(part or '').strip()]).strip()
                else:
                    query_text = f"{knowledge_point} {effective_original_text}".strip()
                
                # 从 Milvus 检索相关内容
                try:
                    from milvus_vector import get_vector_store as _get_vs
                    current_vector_store = _get_vs() if _get_vs else None
                    if current_vector_store:
                        retriever = current_vector_store.as_retriever(search_kwargs={"k": 3})
                        relevant_docs = retriever.invoke(query_text)
                        context = "\n\n".join([doc.page_content for doc in relevant_docs])
                        context_snippets = [doc.page_content[:200] + "..." for doc in relevant_docs]
                        print("=== Similar Question RAG Query Start ===")
                        print(query_text)
                        print("=== Similar Question RAG Query End ===")
                        print("=== Similar Question RAG Context Start ===")
                        print(context if context else "暂无相关知识库内容")
                        print("=== Similar Question RAG Context End ===")
                    else:
                        context = "暂无相关知识库内容"
                        context_snippets = []
                except Exception as e:
                    print(f"RAG 检索失败: {e}")
                    context = "暂无相关知识库内容"
                    context_snippets = []
                
                # 格式化选项
                options_str = '\n'.join([f"{chr(65+i)}. {opt}" for i, opt in enumerate(original_options)])
                
                # 使用新的 prompt 模板
                formatted_prompt = generate_similar_question_template.format(
                    original_big_stem=original_big_stem,
                    original_stem=original_stem,
                    original_options=options_str,
                    original_answer=original_answer,
                    original_analysis=original_analysis,
                    knowledge_point=knowledge_point,
                    difficulty=difficulty,
                    context=context,
                    user_requirement=user_requirement
                )
                
                # 调用 LLM 生成
                llm = ChatOpenAI(
                    model=config.llm_model_name,
                    api_key=config.api_key,
                    base_url=config.base_url,
                    temperature=0.7
                )
                
                # 解析 JSON 响应
                def parse_llm_response(text: str) -> dict:
                    """解析 LLM 返回的 JSON"""
                    try:
                        # 提取 JSON 部分
                        start = text.find('{')
                        end = text.rfind('}') + 1
                        if start != -1 and end > start:
                            json_str = text[start:end]
                            return json.loads(json_str)
                        return None
                    except Exception as e:
                        print(f"JSON 解析错误: {e}")
                        return None

                retry_instruction = (
                    "\n\n### 额外强制要求\n"
                    "你上一次如果生成了与原题相同、近似改写或仅替换个别字词的题目，本次必须完全避免。"
                    "新题题干不得与原题一致，不能只改选项和解析，至少需要明显改变题目情境、问法或题干表达。"
                )

                parsed_result = None
                duplicate_detected = False
                for attempt in range(2):
                    prompt_to_use = formatted_prompt if attempt == 0 else formatted_prompt + retry_instruction
                    response = llm.invoke(prompt_to_use)
                    response_text = response.content
                    parsed_result = parse_llm_response(response_text)
                    if not parsed_result:
                        continue

                    candidate_big_stem = parsed_result.get('big_stem', '') or ''
                    candidate_stem = parsed_result.get('stem', '') or parsed_result.get('question', '') or ''

                    if not normalize_question_text(original_stem):
                        if not normalize_question_text(candidate_big_stem):
                            candidate_big_stem = candidate_stem
                        candidate_stem = ''

                    candidate_options = parsed_result.get('options', []) or []
                    duplicate_detected = is_invalid_similar_question_structure(
                        original_big_stem,
                        original_stem,
                        candidate_big_stem,
                        candidate_stem,
                    )
                    if not duplicate_detected:
                        duplicate_detected = is_same_question_content(
                            effective_original_text,
                            original_options,
                            build_effective_question_text(candidate_big_stem, candidate_stem),
                            candidate_options,
                        )
                    if not duplicate_detected:
                        parsed_result['big_stem'] = candidate_big_stem
                        parsed_result['stem'] = candidate_stem
                        break
                    print('检测到生成结果与原题重复，准备使用更强约束重试。')
                
                if not parsed_result:
                    return JSONResponse({
                        'success': False,
                        'message': 'LLM 响应解析失败'
                    })

                if duplicate_detected:
                    return JSONResponse({
                        'success': False,
                        'message': '生成结果与原题重复，请重试或补充更具体的改写要求'
                    })
                
                # 构造新题目对象（保持与原题相同的结构）
                new_question = {
                    'stem': parsed_result.get('stem', ''),
                    'options': parsed_result.get('options', []),
                    'answer': parsed_result.get('answer', ''),
                    'analysis': parsed_result.get('analysis') or parsed_result.get('explanation', ''),
                    'big_stem': parsed_result.get('big_stem', ''),
                    'kp': original_question.get('kp', ''),
                    'kp_name': knowledge_point,
                    'level': difficulty
                }
                
                return JSONResponse({
                    'success': True,
                    'new_question': new_question,
                    'context_snippets': context_snippets
                })
                
            except Exception as e:
                import traceback
                traceback.print_exc()
                return JSONResponse({
                    'success': False,
                    'message': f'生成失败: {str(e)}'
                })

        @web_app.get('/short-answer', response_class=HTMLResponse)
        async def short_answer_page(request: Request, session_id: Optional[str] = Cookie(None)):
            # 检查用户是否登录
            user = get_current_user(session_id)
            if not user:
                return RedirectResponse(url='/login', status_code=303)
            return templates.TemplateResponse(
                'short_answer_modules.html',
                {
                    'request': request,
                    'user': user,
                    'modules': SHORT_ANSWER_MODULES,
                }
            )

        @web_app.get('/short-answer/{module_slug}', response_class=HTMLResponse)
        async def short_answer_module_page(request: Request, module_slug: str, session_id: Optional[str] = Cookie(None)):
            # 检查用户是否登录
            user = get_current_user(session_id)
            if not user:
                return RedirectResponse(url='/login', status_code=303)

            module = get_short_answer_module(module_slug)
            if not module:
                return RedirectResponse(url='/short-answer', status_code=303)

            return templates.TemplateResponse(
                'short_answer.html',
                {
                    'request': request,
                    'user': user,
                    'module': module,
                }
            )

        @web_app.post('/generate-er-diagram')
        async def generate_er_diagram(request: Request, session_id: Optional[str] = Cookie(None)):
            """生成 E-R 图结构"""
            # 检查用户是否登录
            user = get_current_user(session_id)
            if not user:
                return JSONResponse({'success': False, 'message': '请先登录'})

            start_time = time.perf_counter()
            
            # 检查 RAG 功能是否可用
            if not RAG_AVAILABLE:
                return JSONResponse({
                    'success': False,
                    'message': 'RAG 功能未启用，无法使用 AI 生成'
                })
            
            try:
                # 获取请求参数
                body = await request.json()
                entity_count = int(body.get('entityCount', '3'))
                relationship_count = int(body.get('relationshipCount', '2'))
                attribute_count = int(body.get('attributeCount', '12'))
                scenario = str(body.get('scenario', 'random')).strip().lower()
                if scenario not in SCENARIO_LABELS:
                    scenario = 'random'
                
                # 根据选择修改prompt
                prompt_text = generate_er_diagram_template
                prompt_text = prompt_text.replace('生成 2-3 个实体', f'生成 {entity_count} 个实体')
                
                # 添加关系数量要求
                prompt_text = prompt_text.replace(
                    '实体之间应该有关系',
                    f'实体之间应该有 {relationship_count} 个关系，且至少包含 1 个多对多（M:N）关系'
                )
                
                # 计算属性分配策略
                # 每个实体至少1个主键，最多4个属性；每个关系0-4个属性
                entity_attrs = min(4, max(1, attribute_count // (entity_count + relationship_count)))
                rel_attrs = max(0, (attribute_count - entity_count * entity_attrs) // relationship_count)
                rel_attrs = min(4, rel_attrs)
                
                # 添加详细的属性分配要求
                attr_instruction = f'''
**重要**：请严格控制属性数量，使得所有实体属性和关系属性的总数接近 {attribute_count} 个。
- 每个实体应该有 {entity_attrs} 个左右的属性（至少1个主键，最多4个）
- 每个关系应该有 {rel_attrs} 个左右的属性（可以是0个，最多4个）
- 总属性数 = {entity_count}个实体的属性 + {relationship_count}个关系的属性 ≈ {attribute_count}个
'''
                
                prompt_text = prompt_text.replace(
                    '每个实体应该有 1-4 个属性，其中一个必须是主键（标记为 isPrimaryKey: true）',
                    f'每个实体应该有约 {entity_attrs} 个属性（至少1个主键标记为 isPrimaryKey: true，最多4个）{attr_instruction}'
                )
                
                prompt_text = prompt_text.replace(
                    '每个关系可以有 0-4 个描述性属性（关系可以没有属性）',
                    f'每个关系应该有约 {rel_attrs} 个属性（可以是0个，最多4个）'
                )

                scenario_instruction = ''
                if scenario == 'random':
                    scenario_instruction = '### 场景要求\n1. 请在医疗、教育、火车、飞机、酒店、订票这些业务场景中随机选择一个生成。\n2. 不要总是重复医院、医生、患者主题，优先保证本次场景与常见医疗题不同。\n'
                else:
                    scenario_instruction = f'### 场景要求\n1. 本次必须使用“{SCENARIO_LABELS[scenario]}”场景生成实体、关系和属性。\n2. 不要切换到其他业务领域，所有实体和关系都必须与“{SCENARIO_LABELS[scenario]}”场景直接相关。\n'

                prompt_text = prompt_text.replace('### 输出格式', scenario_instruction + '### 输出格式')
                
                # 调用 LLM 生成 E-R 图结构
                llm = ChatOpenAI(
                    model=config.llm_model_name,
                    api_key=config.api_key,
                    base_url=config.base_url,
                    temperature=0.8  # 提高温度以增加随机性
                )
                
                response = llm.invoke(prompt_text)
                response_text = response.content
                
                # 解析 JSON 响应
                def parse_llm_response(text: str) -> dict:
                    """解析 LLM 返回的 JSON"""
                    try:
                        # 尝试提取 JSON 部分
                        start = text.find('{')
                        end = text.rfind('}') + 1
                        if start != -1 and end > start:
                            json_str = text[start:end]
                            return json.loads(json_str)
                        return None
                    except Exception as e:
                        print(f"JSON 解析错误: {e}")
                        return None
                
                parsed_result = parse_llm_response(response_text)
                
                if not parsed_result:
                    timing = build_generation_timing(start_time)
                    message = 'LLM 响应解析失败'
                    print_generation_timing({'title': 'E-R图生成'}, timing, success=False, message=message, event_label='E-R图生成耗时')
                    return JSONResponse({
                        'success': False,
                        'message': message,
                        'timing': timing,
                        **timing,
                    })

                parsed_result = ensure_many_to_many_relationship(parsed_result)
                
                # 验证结果结构
                if 'entities' not in parsed_result or 'relationships' not in parsed_result:
                    timing = build_generation_timing(start_time)
                    message = 'E-R 图结构不完整'
                    print_generation_timing({'title': 'E-R图生成'}, timing, success=False, message=message, event_label='E-R图生成耗时')
                    return JSONResponse({
                        'success': False,
                        'message': message,
                        'timing': timing,
                        **timing,
                    })
                
                timing = build_generation_timing(start_time)
                print_generation_timing({'title': 'E-R图生成'}, timing, success=True, event_label='E-R图生成耗时')
                return JSONResponse({
                    'success': True,
                    'data': parsed_result,
                    'timing': timing,
                    **timing,
                })
                
            except Exception as e:
                import traceback
                traceback.print_exc()
                timing = build_generation_timing(start_time)
                message = build_llm_failure_message(e, 'E-R图生成')
                print_generation_timing({'title': 'E-R图生成'}, timing, success=False, message=message, event_label='E-R图生成耗时')
                return JSONResponse({
                    'success': False,
                    'message': message,
                    'timing': timing,
                    **timing,
                })

        @web_app.post('/generate-question-stem')
        async def generate_question_stem(request: Request, session_id: Optional[str] = Cookie(None)):
            """根据E-R图生成题干"""
            # 检查用户是否登录
            user = get_current_user(session_id)
            if not user:
                return JSONResponse({'success': False, 'message': '请先登录'})

            start_time = time.perf_counter()
            
            # 检查 RAG 功能是否可用
            if not RAG_AVAILABLE:
                return JSONResponse({
                    'success': False,
                    'message': 'RAG 功能未启用，无法使用 AI 生成'
                })
            
            try:
                # 获取E-R图数据
                body = await request.json()
                original_er_data = normalize_er_data_types(body.get('er_data', {}))
                if not original_er_data:
                    timing = build_generation_timing(start_time)
                    message = '缺少 er_data'
                    print_generation_timing({'title': '题干生成'}, timing, success=False, message=message, event_label='题干生成耗时')
                    return JSONResponse({'success': False, 'message': message, 'timing': timing, **timing})
                
                # 深拷贝一份用于修改
                masked_er_data = json.loads(json.dumps(original_er_data))
                
                # 1. 生成背景描述
                # 格式化E-R图数据为可读格式
                er_description = "实体列表：\n"
                for entity in original_er_data.get('entities', []):
                    er_description += f"- {entity['name']}：属性包括"
                    attrs = [f"{attr['name']}({attr['type']})" for attr in entity.get('attributes', [])]
                    er_description += "、".join(attrs) + "\n"
                
                er_description += "\n关系列表：\n"
                for rel in original_er_data.get('relationships', []):
                    er_description += f"- {rel['name']}（{rel['entity1']} 与 {rel['entity2']}，类型：{rel['cardinality']}）"
                    if rel.get('attributes'):
                        attrs = [f"{attr['name']}({attr['type']})" for attr in rel['attributes']]
                        er_description += "，属性：" + "、".join(attrs)
                    er_description += "\n"
                
                # 导入背景生成模板
                from protocol.prompts import generate_background_template, generate_physical_design_template
                
                # 填充prompt
                prompt_text = generate_background_template.format(er_data=er_description)
                
                # 调用 LLM 生成背景
                llm = ChatOpenAI(
                    model=config.llm_model_name,
                    api_key=config.api_key,
                    base_url=config.base_url,
                    temperature=0.7
                )
                
                response = llm.invoke(prompt_text)
                background_text = response.content.strip()

                # 生成物理设计题目
                pd_prompt = generate_physical_design_template.format(er_data=er_description)
                pd_response = llm.invoke(pd_prompt)
                pd_content = pd_response.content.strip()
                
                # 清洗 JSON
                if pd_content.strip().startswith("```json"):
                     pd_content = pd_content.strip()[7:]
                if pd_content.strip().endswith("```"):
                     pd_content = pd_content.strip()[:-3]
                
                try:
                    physical_design_data = json.loads(pd_content)
                    physical_design_data = normalize_physical_design_data(physical_design_data)
                    physical_design_data = repair_physical_design_tables(physical_design_data, original_er_data)
                    physical_design_data['tables'] = build_structure_table_mask(physical_design_data.get('tables', []))
                    physical_design_data['tables'] = attach_sample_rows_to_tables(physical_design_data.get('tables', []))
                    physical_design_data['tables'] = build_sample_data_mask(physical_design_data.get('tables', []))
                    physical_design_data = align_design_questions_with_sample_mask(physical_design_data)
                except Exception as e:
                    print(f"物理设计题目解析失败: {e}")
                    physical_design_data = {}

                
                # 2. 随机选择挖空部分 (A, B, C)
                # A, B: 两个属性
                # C: 一个基数 (n 或 m)
                
                all_attributes = []
                # 收集所有实体属性
                for e_idx, entity in enumerate(masked_er_data.get('entities', [])):
                    for a_idx, attr in enumerate(entity.get('attributes', [])):
                        all_attributes.append({
                            'type': 'entity_attr',
                            'entity_idx': e_idx,
                            'attr_idx': a_idx,
                            'name': attr['name']
                        })
                
                # 收集所有关系属性
                for r_idx, rel in enumerate(masked_er_data.get('relationships', [])):
                    for a_idx, attr in enumerate(rel.get('attributes', [])):
                        all_attributes.append({
                            'type': 'rel_attr',
                            'rel_idx': r_idx,
                            'attr_idx': a_idx,
                            'name': attr['name']
                        })
                
                if len(all_attributes) < 2:
                    timing = build_generation_timing(start_time)
                    message = '属性数量不足，无法生成挖空题目'
                    print_generation_timing({'title': '题干生成'}, timing, success=False, message=message, event_label='题干生成耗时')
                    return JSONResponse({'success': False, 'message': message, 'timing': timing, **timing})
                
                # 随机选择两个属性
                selected_attrs = random.sample(all_attributes, 2)
                attr_a = selected_attrs[0]
                attr_b = selected_attrs[1]
                
                # 执行挖空替换 A
                if attr_a['type'] == 'entity_attr':
                    masked_er_data['entities'][attr_a['entity_idx']]['attributes'][attr_a['attr_idx']]['name'] = 'A'
                else:
                    masked_er_data['relationships'][attr_a['rel_idx']]['attributes'][attr_a['attr_idx']]['name'] = 'A'

                # 执行挖空替换 B
                if attr_b['type'] == 'entity_attr':
                    masked_er_data['entities'][attr_b['entity_idx']]['attributes'][attr_b['attr_idx']]['name'] = 'B'
                else:
                    masked_er_data['relationships'][attr_b['rel_idx']]['attributes'][attr_b['attr_idx']]['name'] = 'B'
                
                # 选择一个关系进行基数挖空 C
                relationships = masked_er_data.get('relationships', [])
                if not relationships:
                    timing = build_generation_timing(start_time)
                    message = '关系数量不足，无法生成挖空题目'
                    print_generation_timing({'title': '题干生成'}, timing, success=False, message=message, event_label='题干生成耗时')
                    return JSONResponse({'success': False, 'message': message, 'timing': timing, **timing})
                
                # 优先选择 M:N 或 1:N 关系
                target_rel_idx = -1
                original_c_val = ''
                
                # 尝试找 M:N
                mn_rels = [i for i, r in enumerate(relationships) if 'M' in r['cardinality'] or 'N' in r['cardinality']]
                if mn_rels:
                    target_rel_idx = random.choice(mn_rels)
                else:
                    target_rel_idx = random.randint(0, len(relationships) - 1)
                
                rel = relationships[target_rel_idx]
                parts = rel['cardinality'].split(':')
                
                # 决定挖哪一边
                if len(parts) == 2:
                    if parts[1] in ['N', 'M']:
                        original_c_val = parts[1]
                        parts[1] = 'C'
                    elif parts[0] in ['N', 'M']:
                        original_c_val = parts[0]
                        parts[0] = 'C'
                    else:
                        target_part = random.randint(0, 1)
                        original_c_val = parts[target_part]
                        parts[target_part] = 'C'
                    
                    # 更新基数
                    masked_er_data['relationships'][target_rel_idx]['cardinality'] = ":".join(parts)
                else:
                     original_c_val = "未知"

                # 3. 构造完整题目
                question_text = f"{background_text}\n\n"
                question_text += "1. 请问系统有哪几类用户？简述需要为各类用户设计哪些系统功能。（结构化题目无需更改）\n\n"
                question_text += "2. 某设计人员给出了该系统数据库概念设计的E-R图（图1），请补充缺失部分的A、B、C处所对应的联系、属性、联系类型，完成E-R图。（需要在E-R图上选两个属性，一个n或m标记进行挖空）\n"
                
                # 处理物理设计题目的挖空逻辑
                masked_tables = []
                if physical_design_data:
                    import copy
                    masked_tables = copy.deepcopy(physical_design_data.get('tables', []))
                    design_questions = physical_design_data.get('design_questions', [])

                    question_text += "\n3. 根据E-R图，完成以下数据库物理设计任务：\n"
                    
                    # 关系模式题目
                    question_text += "(1) 完善下列关系模式（只列出主键和外键，以及缺失的属性）：\n"
                    req_idx = 1
                    for rq in physical_design_data.get('relation_questions', []):
                         question_text += f"   {req_idx}) {rq.get('table_name')}表: {rq.get('table_name')}(                       )\n"
                         question_text += "      主关键字：___________   外关键字：___________\n"
                         req_idx += 1
                    
                    question_text += "(2) 根据上述表结构设计，回答下列问题：\n"
                    q_idx = 1
                    import re
                    for dq in design_questions:
                        # 去除题目开头可能存在的数字编号（如 "1)", "1）", "1." 等）
                        clean_question = re.sub(r'^\d+[.)）]\s*', '', dq.get('question', '')) 
                        question_text += f"   {q_idx}) {clean_question}\n"
                        q_idx += 1

                # 附带答案（仅供参考/调试）
                answers = {
                    "A": attr_a['name'],
                    "B": attr_b['name'],
                    "C": original_c_val
                }
                
                # 添加物理设计答案
                if physical_design_data:
                    for i, rq in enumerate(physical_design_data.get('relation_questions', [])):
                         # 提取括号内的内容作为答案
                         full_relation = rq.get('relation_mode', '')
                         relation_content = full_relation
                         if '(' in full_relation and ')' in full_relation:
                            parts = full_relation.split('(', 1)
                            relation_content = parts[1].rstrip(')')

                         answers[f"Q3-Rel-{i+1}-Mode"] = relation_content
                         answers[f"Q3-Rel-{i+1}-PK"] = rq.get('pk', '')
                         answers[f"Q3-Rel-{i+1}-FK"] = rq.get('fk', '')
                    
                    for i, dq in enumerate(physical_design_data.get('design_questions', [])):
                         answers[f"Q3-Design-{i+1}"] = dq.get('answer', '')
                    
                    # 添加表数据挖空答案
                    for t_idx, table in enumerate(masked_tables):
                        sample_mask = table.get('sample_mask') or {}
                        if sample_mask:
                            answers[f"Q3-Sample-{table.get('name')}-{sample_mask.get('row_index')}-{sample_mask.get('column_name')}"] = sample_mask.get('answer', '')

                timing = build_generation_timing(start_time)
                print_generation_timing({'title': '题干生成'}, timing, success=True, event_label='题干生成耗时')
                return JSONResponse({
                    'success': True, 
                    'question_stem': question_text,
                    'masked_er_data': masked_er_data,
                    'physical_design_data': {
                        'tables': masked_tables, # 发送挖空后的表结构
                        'design_questions': physical_design_data.get('design_questions', []),
                        'relation_questions': physical_design_data.get('relation_questions', [])
                    },
                    'answers': answers,
                    'background_text': background_text,
                    'timing': timing,
                    **timing,
                })
            except Exception as e:
                import traceback
                traceback.print_exc()
                timing = build_generation_timing(start_time)
                message = build_llm_failure_message(e, '题干生成')
                print_generation_timing({'title': '题干生成'}, timing, success=False, message=message, event_label='题干生成耗时')
                return JSONResponse({
                    'success': False,
                    'message': message,
                    'timing': timing,
                    **timing,
                })

        @web_app.post('/export-question-docx')
        async def export_question_docx(request: Request, session_id: Optional[str] = Cookie(None)):
            """导出当前题目为 .docx 文件，接收前端发送的题目信息和画布图片（dataURL）"""
            user = get_current_user(session_id)
            if not user:
                return JSONResponse({'success': False, 'message': '请先登录'})

            try:
                body = await request.json()
                background_text = body.get('background_text', '')
                # Extend to receive custom question texts
                main_title = body.get('main_title', '(一) 信息系统设计（本题3小题，共20分）')
                
                q1_text = body.get('q1_text', '1. 请问系统有哪几类用户？简述需要为各类用户设计哪些系统功能。')
                q1_answer = body.get('q1_answer', '简述功能，合理即可')
                
                q2_text = body.get('q2_text', '2. 某设计人员给出了该系统数据库概念设计的E-R图（图1），请补充缺失部分的A、B、C处所对应的联系、属性、联系类型，完成E-R图。')
                
                er_caption = body.get('er_caption', '图1 系统局部E-R图')
                
                q3_intro = body.get('q3_intro', '3. 根据E-R图，完成以下数据库物理设计任务：')
                q3_sub1 = body.get('q3_sub1', '(1) 完善下列关系模式（只列出主键和外键，以及缺失的属性）：')
                q3_sub2 = body.get('q3_sub2', '(2) 根据上述表结构设计，回答下列问题：')
                
                answers = body.get('answers', {}) or {}
                er_image_data = body.get('er_image', None)
                physical_design_data = body.get('physical_design_data', {})
                show_answer = body.get('show_answer', False)

                # 准备输出目录
                output_dir = './output'
                os.makedirs(output_dir, exist_ok=True)
                ts = datetime.now().strftime('%Y%m%d_%H%M%S')
                docx_name = f'question_export_{ts}.docx'
                docx_path = os.path.join(output_dir, docx_name)

                doc = Document()

                # 标题
                p = doc.add_paragraph()
                r = p.add_run(main_title)
                r.bold = True
                r.font.size = Pt(12)

                # 背景文字
                if background_text:
                    for line in background_text.split('\n'):
                        para = doc.add_paragraph(line)
                        para.paragraph_format.space_after = Pt(6)

                # 引导语
                doc.add_paragraph('请完成以下系统分析和设计，将你的答案填写到相应题号下的空格中。')

                # 问题1
                doc.add_paragraph(q1_text)
                if show_answer:
                    p = doc.add_paragraph()
                    add_underlined_answer_run(p, q1_answer)

                # 问题2
                doc.add_paragraph(q2_text)
                
                # Q2 答案
                if show_answer and answers:
                    q2_keys = ['A', 'B', 'C']
                    # 直接在一行或多行显示 A: x, B: x, C: x
                    p = doc.add_paragraph()
                    first = True
                    for k in q2_keys:
                         if k in answers:
                            if not first:
                                p.add_run(', ')
                            p.add_run(f'{k}: ')
                            add_underlined_answer_run(p, answers.get(k, ''))
                            first = False

                # 插入 E-R 图图片（居中）
                if er_image_data and isinstance(er_image_data, str) and er_image_data.startswith('data:image'):
                    try:
                        header, b64 = er_image_data.split(',', 1)
                        img_bytes = base64.b64decode(b64)
                        # Save temp image
                        img_filename = f'er_{ts}.png'
                        img_path = os.path.join(output_dir, img_filename)
                        with open(img_path, 'wb') as f:
                            f.write(img_bytes)

                        pic = doc.add_picture(img_path, width=Inches(5.5))
                        last = doc.paragraphs[-1]
                        last.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        cap = doc.add_paragraph(er_caption)
                        cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        cap.runs[0].font.size = Pt(9)
                    except Exception as e:
                        print(f'插入图片失败: {e}')

                # 问题3
                if physical_design_data:
                    doc.add_paragraph(q3_intro)
                    
                    # 1. 表结构表格
                    tables = physical_design_data.get('tables', [])
                    for table in tables:
                        doc.add_paragraph(format_table_title(table.get('name', ''), table.get('description', '')))
                        table_grid = doc.add_table(rows=1, cols=4)
                        table_grid.style = 'Table Grid'
                        hdr = table_grid.rows[0].cells
                        hdr[0].text = '字段名'
                        hdr[1].text = '字段描述'
                        hdr[2].text = '数据类型'
                        hdr[3].text = '属性限制'
                        
                        for col in table.get('columns', []):
                            row = table_grid.add_row().cells
                            row[0].text = col.get('name', '')
                            row[1].text = col.get('desc', '')
                            
                            masked_field = col.get('masked_field')
                            normalized_type = normalize_data_type(col.get('type', ''))
                            
                            # Type column (now index 2)
                            p_type = row[2].paragraphs[0]
                            if masked_field == 'type':
                                add_blank_or_answer(p_type, normalized_type, show_answer, '________')
                                p_type.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            else:
                                row[2].text = normalized_type
                                
                            # Constraint column 
                            display_constraint = format_constraint_for_display(col.get('constraint', '') or '')
                                
                            p_constr = row[3].paragraphs[0]
                            if masked_field == 'constraint':
                                add_blank_or_answer(p_constr, display_constraint, show_answer, '________')
                                p_constr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            else:
                                row[3].text = display_constraint

                        doc.add_paragraph('')

                    for table in tables:
                        sample_rows = table.get('sample_rows', []) or []
                        if not sample_rows:
                            continue

                        sample_mask = table.get('sample_mask') or {}

                        doc.add_paragraph(f"{table.get('name', '')}表数据")
                        sample_table = doc.add_table(rows=1, cols=len(table.get('columns', [])))
                        sample_table.style = 'Table Grid'
                        header_cells = sample_table.rows[0].cells
                        for idx, column in enumerate(table.get('columns', [])):
                            header_cells[idx].text = column.get('name', '')

                        for row_index, sample_row in enumerate(sample_rows):
                            sample_cells = sample_table.add_row().cells
                            for idx, column in enumerate(table.get('columns', [])):
                                value = str(sample_row.get(column.get('name', ''), ''))
                                paragraph = sample_cells[idx].paragraphs[0]
                                if sample_mask and row_index == sample_mask.get('row_index') and column.get('name', '') == sample_mask.get('column_name'):
                                    add_blank_or_answer(paragraph, sample_mask.get('answer', value), show_answer, '____')
                                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                else:
                                    paragraph.add_run(value)

                        doc.add_paragraph('')
                    
                    # 2. 关系模式
                    doc.add_paragraph(q3_sub1)
                    req_idx = 1
                    for rq in physical_design_data.get('relation_questions', []):
                         p = doc.add_paragraph(f"   {req_idx}) {rq.get('table_name')}表: {rq.get('table_name')}(")
                         
                         full_relation = rq.get('relation_mode', '')
                         content = ''
                         if '(' in full_relation:
                             # Try to extract content inside (), assuming format T(A,B,C)
                             try:
                                 content = full_relation.split('(', 1)[1].rstrip(')')
                             except:
                                 content = full_relation
                         else:
                             content = full_relation

                         add_blank_or_answer(p, content, show_answer, '________')
                         
                         p.add_run(")")
                         
                         p2 = doc.add_paragraph(f"      主关键字：")
                         add_blank_or_answer(p2, rq.get('pk', ''), show_answer, '___________')
                         
                         p2.add_run("   外关键字：")
                         add_blank_or_answer(p2, rq.get('fk', ''), show_answer, '__________')
                             
                         req_idx += 1

                    # 3. 设计问题
                    doc.add_paragraph(q3_sub2)
                    q_idx = 1
                    import re
                    for dq in physical_design_data.get('design_questions', []):
                        # 去除题目开头可能存在的数字编号
                        clean_question = re.sub(r'^\d+[.)）]\s*', '', dq.get('question', ''))
                        p = doc.add_paragraph(f"   {q_idx}) ")
                        render_question_with_blank(p, clean_question, dq.get('answer', ''), show_answer)
                            
                        q_idx += 1

                else:
                    doc.add_paragraph('3. 暂时占位，后续会补充第三问')

                # 保存 docx
                doc.save(docx_path)

                return JSONResponse({'success': True, 'file_url': f'/output/{docx_name}', 'file_path': docx_path})

            except Exception as e:
                import traceback
                traceback.print_exc()
                return JSONResponse({'success': False, 'message': str(e)})

        @web_app.get('/knowledge-base', response_class=HTMLResponse)
        async def knowledge_base_page(request: Request, session_id: Optional[str] = Cookie(None)):
            # 检查用户是否登录
            user = get_current_user(session_id)
            if not user:
                return RedirectResponse(url='/login', status_code=303)
            """知识库管理页面"""
            return templates.TemplateResponse('knowledge_base.html', {'request': request})

        @web_app.get('/question-bank', response_class=HTMLResponse)
        async def question_bank_page(request: Request, session_id: Optional[str] = Cookie(None)):
            user = get_current_user(session_id)
            if not user:
                return RedirectResponse(url='/login', status_code=303)
            ensure_question_bank_db()
            return templates.TemplateResponse('question_bank.html', {'request': request, 'user': user})

        @web_app.get('/api/question-bank/meta', response_class=JSONResponse)
        async def question_bank_meta(session_id: Optional[str] = Cookie(None)):
            user = get_current_user(session_id)
            if not user:
                return JSONResponse({'success': False, 'message': '请先登录'}, status_code=401)

            ensure_question_bank_db()
            conn = sqlite3.connect(QUESTION_DB_PATH)
            cur = conn.cursor()
            cur.execute('SELECT DISTINCT question_type FROM question_bank WHERE question_type IS NOT NULL AND TRIM(question_type) <> "" ORDER BY question_type')
            question_types = [row[0] for row in cur.fetchall()]
            cur.execute('SELECT DISTINCT difficulty FROM question_bank WHERE difficulty IS NOT NULL AND TRIM(difficulty) <> "" ORDER BY difficulty')
            difficulties = [row[0] for row in cur.fetchall()]
            cur.execute(
                '''
                SELECT knowledge_point_code, knowledge_point_name
                FROM knowledge_points
                WHERE knowledge_point_code IS NOT NULL AND TRIM(knowledge_point_code) <> ''
                ORDER BY knowledge_point_code
                '''
            )
            kp_rows = cur.fetchall()
            conn.close()
            return JSONResponse({
                'success': True,
                'question_types': question_types,
                'difficulties': difficulties,
                'knowledge_points': [{'code': _safe_str(c), 'name': _safe_str(n)} for c, n in kp_rows]
            })

        @web_app.get('/api/question-bank/questions', response_class=JSONResponse)
        async def list_question_bank_questions(request: Request, session_id: Optional[str] = Cookie(None)):
            user = get_current_user(session_id)
            if not user:
                return JSONResponse({'success': False, 'message': '请先登录'}, status_code=401)

            qp = request.query_params
            keyword = _safe_str(qp.get('keyword', ''))
            question_type = _safe_str(qp.get('question_type', ''))
            difficulty = _safe_str(qp.get('difficulty', ''))
            kp_code = _safe_str(qp.get('kp_code', ''))
            try:
                page = max(1, int(qp.get('page', 1)))
            except Exception:
                page = 1
            try:
                page_size = min(100, max(1, int(qp.get('page_size', 20))))
            except Exception:
                page_size = 20

            ensure_question_bank_db()
            conn = sqlite3.connect(QUESTION_DB_PATH)
            cur = conn.cursor()

            where_clauses = ['1=1']
            params = []
            if question_type:
                where_clauses.append('question_type LIKE ?')
                params.append(f'%{question_type}%')
            if difficulty:
                where_clauses.append('difficulty = ?')
                params.append(difficulty)
            if kp_code:
                where_clauses.append('knowledge_point_code = ?')
                params.append(kp_code)
            if keyword:
                where_clauses.append('(stem LIKE ? OR big_stem LIKE ? OR analysis LIKE ?)')
                kw = f'%{keyword}%'
                params.extend([kw, kw, kw])

            where_sql = ' AND '.join(where_clauses)
            total = cur.execute(f'SELECT COUNT(1) FROM question_bank WHERE {where_sql}', params).fetchone()[0]
            offset = (page - 1) * page_size
            rows = cur.execute(
                f'''
                  SELECT q.id, q.question_type, q.big_stem, q.stem, q.answer, q.analysis, q.difficulty,
                      q.knowledge_point_code, q.options_json, q.sn, q.ai_generated, q.content_format,
                      q.module_slug,
                       COALESCE(k.knowledge_point_name, '') AS knowledge_point_name
                FROM question_bank q
                LEFT JOIN knowledge_points k ON q.knowledge_point_code = k.knowledge_point_code
                WHERE {where_sql}
                ORDER BY q.id DESC
                LIMIT ? OFFSET ?
                ''',
                params + [page_size, offset]
            ).fetchall()
            conn.close()

            items = []
            for row in rows:
                options = []
                try:
                    options = json.loads(row[8]) if row[8] else []
                except Exception:
                    options = []
                items.append({
                    'id': row[0],
                    'question_type': _safe_str(row[1]),
                    'big_stem': _safe_str(row[2]),
                    'stem': _safe_str(row[3]),
                    'answer': _safe_str(row[4]),
                    'analysis': _safe_str(row[5]),
                    'difficulty': _safe_str(row[6]),
                    'knowledge_point_code': _safe_str(row[7]),
                    'options': options if isinstance(options, list) else [],
                    'sn': _safe_str(row[9]),
                    'ai_generated': bool(row[10]),
                    'content_format': _safe_str(row[11]) or 'standard',
                    'module_slug': _safe_str(row[12]),
                    'knowledge_point_name': _safe_str(row[13]),
                    'preview_text': _safe_str(row[3]) or _safe_str(row[2]),
                })

            return JSONResponse({
                'success': True,
                'items': items,
                'pagination': {
                    'page': page,
                    'page_size': page_size,
                    'total': total,
                    'total_pages': (total + page_size - 1) // page_size if page_size else 0,
                }
            })

        @web_app.get('/api/question-bank/questions/{question_id}', response_class=JSONResponse)
        async def get_question_bank_question(question_id: int, session_id: Optional[str] = Cookie(None)):
            user = get_current_user(session_id)
            if not user:
                return JSONResponse({'success': False, 'message': '请先登录'}, status_code=401)

            ensure_question_bank_db()
            conn = sqlite3.connect(QUESTION_DB_PATH)
            cur = conn.cursor()
            row = cur.execute(
                '''
                SELECT q.id, q.question_type, q.big_stem, q.stem, q.answer, q.analysis, q.difficulty,
                       q.knowledge_point_code, q.options_json, q.sn, q.ai_generated, q.content_format,
                       q.module_slug, q.payload_json, COALESCE(k.knowledge_point_name, '')
                FROM question_bank q
                LEFT JOIN knowledge_points k ON q.knowledge_point_code = k.knowledge_point_code
                WHERE q.id = ?
                ''',
                (question_id,)
            ).fetchone()
            conn.close()

            if not row:
                return JSONResponse({'success': False, 'message': '题目不存在'}, status_code=404)

            options = safe_json_loads(row[8], [])
            payload = safe_json_loads(row[13], {})
            return JSONResponse({
                'success': True,
                'item': {
                    'id': row[0],
                    'question_type': _safe_str(row[1]),
                    'big_stem': _safe_str(row[2]),
                    'stem': _safe_str(row[3]),
                    'answer': _safe_str(row[4]),
                    'analysis': _safe_str(row[5]),
                    'difficulty': _safe_str(row[6]),
                    'knowledge_point_code': _safe_str(row[7]),
                    'options': options if isinstance(options, list) else [],
                    'sn': _safe_str(row[9]),
                    'ai_generated': bool(row[10]),
                    'content_format': _safe_str(row[11]) or 'standard',
                    'module_slug': _safe_str(row[12]),
                    'payload': payload if isinstance(payload, dict) else {},
                    'knowledge_point_name': _safe_str(row[14]),
                }
            })

        @web_app.post('/api/question-bank/short-answer', response_class=JSONResponse)
        async def create_short_answer_question_bank_question(request: Request, session_id: Optional[str] = Cookie(None)):
            user = get_current_user(session_id)
            if not user:
                return JSONResponse({'success': False, 'message': '请先登录'}, status_code=401)

            try:
                body = await request.json()
                record = build_short_answer_question_bank_record(body)
            except ValueError as e:
                return JSONResponse({'success': False, 'message': str(e)}, status_code=400)
            except Exception as e:
                return JSONResponse({'success': False, 'message': f'保存参数无效: {str(e)}'}, status_code=400)

            ensure_question_bank_db()
            conn = sqlite3.connect(QUESTION_DB_PATH)
            cur = conn.cursor()
            cur.execute(
                '''
                INSERT OR IGNORE INTO question_bank
                (sn, question_type, big_stem, stem, answer, analysis, difficulty, knowledge_point_code,
                 options_json, ai_generated, content_format, module_slug, payload_json, source_file, unique_key)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    '',
                    record['question_type'],
                    record['big_stem'],
                    record['stem'],
                    record['answer'],
                    record['analysis'],
                    record['difficulty'],
                    record['knowledge_point_code'],
                    record['options_json'],
                    record['ai_generated'],
                    record['content_format'],
                    record['module_slug'],
                    record['payload_json'],
                    record['source_file'],
                    record['unique_key'],
                )
            )

            if cur.rowcount == 0:
                existing = cur.execute('SELECT id FROM question_bank WHERE unique_key = ?', (record['unique_key'],)).fetchone()
                conn.commit()
                conn.close()
                return JSONResponse({'success': False, 'message': '题库中已存在相同综合题', 'id': existing[0] if existing else None}, status_code=409)

            new_id = cur.lastrowid
            conn.commit()
            conn.close()
            return JSONResponse({'success': True, 'message': '已保存至题库', 'id': new_id})

        @web_app.post('/api/question-bank/questions', response_class=JSONResponse)
        async def create_question_bank_question(request: Request, session_id: Optional[str] = Cookie(None)):
            user = get_current_user(session_id)
            if not user:
                return JSONResponse({'success': False, 'message': '请先登录'}, status_code=401)

            body = await request.json()
            question_type = _safe_str(body.get('question_type')) or '单选'
            big_stem = _safe_str(body.get('big_stem'))
            stem = _safe_str(body.get('stem'))
            answer = _safe_str(body.get('answer'))
            analysis = _safe_str(body.get('analysis'))
            difficulty = _safe_str(body.get('difficulty'))
            knowledge_point_code = _safe_str(body.get('knowledge_point_code'))
            knowledge_point_name = _safe_str(body.get('knowledge_point_name'))
            options = body.get('options', []) or []

            if not isinstance(options, list):
                return JSONResponse({'success': False, 'message': 'options 必须是数组'}, status_code=400)
            if not big_stem and not stem:
                return JSONResponse({'success': False, 'message': '大题干和小题干不能同时为空'}, status_code=400)

            cleaned_options = [_safe_str(opt) for opt in options if _safe_str(opt)]
            unique_key_src = '|'.join([
                question_type,
                big_stem,
                stem,
                json.dumps(cleaned_options, ensure_ascii=False),
                knowledge_point_code,
                'ai_generated',
            ])
            unique_key = hashlib.md5(unique_key_src.encode('utf-8')).hexdigest()

            ensure_question_bank_db()
            conn = sqlite3.connect(QUESTION_DB_PATH)
            cur = conn.cursor()

            if knowledge_point_code:
                cur.execute(
                    '''
                    INSERT INTO knowledge_points (knowledge_point_code, knowledge_point_name, knowledge_domain_name, knowledge_unit_name)
                    VALUES (?, ?, '未分类', '未分类单元')
                    ON CONFLICT(knowledge_point_code) DO UPDATE SET
                        knowledge_point_name=CASE
                            WHEN COALESCE(excluded.knowledge_point_name, '') <> '' THEN excluded.knowledge_point_name
                            ELSE knowledge_points.knowledge_point_name
                        END
                    ''',
                    (knowledge_point_code, knowledge_point_name),
                )

            cur.execute(
                '''
                INSERT OR IGNORE INTO question_bank
                (sn, question_type, big_stem, stem, answer, analysis, difficulty, knowledge_point_code, options_json, ai_generated, source_file, unique_key)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''',
                (
                    '',
                    question_type,
                    big_stem,
                    stem,
                    answer,
                    analysis,
                    difficulty,
                    knowledge_point_code,
                    json.dumps(cleaned_options, ensure_ascii=False),
                    1,
                    'ai_generated',
                    unique_key,
                )
            )

            if cur.rowcount == 0:
                existing = cur.execute('SELECT id FROM question_bank WHERE unique_key = ?', (unique_key,)).fetchone()
                conn.commit()
                conn.close()
                return JSONResponse({'success': False, 'message': '题库中已存在相同 AI 题目', 'id': existing[0] if existing else None}, status_code=409)

            new_id = cur.lastrowid
            conn.commit()
            conn.close()
            return JSONResponse({'success': True, 'message': '已保存至题库', 'id': new_id})

        @web_app.put('/api/question-bank/questions/{question_id}', response_class=JSONResponse)
        async def update_question_bank_question(question_id: int, request: Request, session_id: Optional[str] = Cookie(None)):
            user = get_current_user(session_id)
            if not user:
                return JSONResponse({'success': False, 'message': '请先登录'}, status_code=401)

            body = await request.json()
            allowed_fields = {
                'question_type', 'big_stem', 'stem', 'answer', 'analysis',
                'difficulty', 'knowledge_point_code', 'options', 'ai_generated'
            }
            updates = {k: body[k] for k in body.keys() if k in allowed_fields}
            if not updates:
                return JSONResponse({'success': False, 'message': '没有可更新字段'}, status_code=400)

            options = updates.get('options', None)
            if options is not None:
                if not isinstance(options, list):
                    return JSONResponse({'success': False, 'message': 'options 必须是数组'}, status_code=400)
                updates['options_json'] = json.dumps([_safe_str(x) for x in options], ensure_ascii=False)
                del updates['options']

            if 'stem' in updates and not _safe_str(updates.get('stem')):
                return JSONResponse({'success': False, 'message': '题干不能为空'}, status_code=400)

            ensure_question_bank_db()
            conn = sqlite3.connect(QUESTION_DB_PATH)
            cur = conn.cursor()

            existing = cur.execute('SELECT id, content_format FROM question_bank WHERE id = ?', (question_id,)).fetchone()
            if not existing:
                conn.close()
                return JSONResponse({'success': False, 'message': '题目不存在'}, status_code=404)
            if _safe_str(existing[1]) == SHORT_ANSWER_CONTENT_FORMAT:
                conn.close()
                return JSONResponse({'success': False, 'message': '该综合题请通过预览页重新生成后再保存，不支持在此表单中直接编辑'}, status_code=400)

            set_parts = []
            params = []
            for field, value in updates.items():
                set_parts.append(f'{field} = ?')
                params.append(_safe_str(value) if field != 'options_json' else value)
            set_parts.append('updated_at = CURRENT_TIMESTAMP')
            params.append(question_id)

            set_sql = ', '.join(set_parts)
            sql = f'UPDATE question_bank SET {set_sql} WHERE id = ?'
            cur.execute(sql, params)
            conn.commit()
            conn.close()
            return JSONResponse({'success': True, 'message': '更新成功'})

        @web_app.put('/api/question-bank/questions/{question_id}/compound', response_class=JSONResponse)
        async def update_compound_question_bank_question(question_id: int, request: Request, session_id: Optional[str] = Cookie(None)):
            user = get_current_user(session_id)
            if not user:
                return JSONResponse({'success': False, 'message': '请先登录'}, status_code=401)

            try:
                body = await request.json()
                record = build_short_answer_question_bank_record(body)
            except ValueError as e:
                return JSONResponse({'success': False, 'message': str(e)}, status_code=400)
            except Exception as e:
                return JSONResponse({'success': False, 'message': f'更新参数无效: {str(e)}'}, status_code=400)

            ensure_question_bank_db()
            conn = sqlite3.connect(QUESTION_DB_PATH)
            cur = conn.cursor()

            existing = cur.execute('SELECT id, content_format FROM question_bank WHERE id = ?', (question_id,)).fetchone()
            if not existing:
                conn.close()
                return JSONResponse({'success': False, 'message': '题目不存在'}, status_code=404)
            if _safe_str(existing[1]) != SHORT_ANSWER_CONTENT_FORMAT:
                conn.close()
                return JSONResponse({'success': False, 'message': '当前题目不是可编辑的大题'}, status_code=400)

            duplicated = cur.execute(
                'SELECT id FROM question_bank WHERE unique_key = ? AND id <> ?',
                (record['unique_key'], question_id)
            ).fetchone()
            if duplicated:
                conn.close()
                return JSONResponse({'success': False, 'message': '题库中已存在相同大题', 'id': duplicated[0]}, status_code=409)

            cur.execute(
                '''
                UPDATE question_bank
                SET question_type = ?,
                    big_stem = ?,
                    stem = ?,
                    answer = ?,
                    analysis = ?,
                    difficulty = ?,
                    knowledge_point_code = ?,
                    options_json = ?,
                    ai_generated = ?,
                    content_format = ?,
                    module_slug = ?,
                    payload_json = ?,
                    source_file = ?,
                    unique_key = ?,
                    updated_at = CURRENT_TIMESTAMP
                WHERE id = ?
                ''',
                (
                    record['question_type'],
                    record['big_stem'],
                    record['stem'],
                    record['answer'],
                    record['analysis'],
                    record['difficulty'],
                    record['knowledge_point_code'],
                    record['options_json'],
                    record['ai_generated'],
                    record['content_format'],
                    record['module_slug'],
                    record['payload_json'],
                    record['source_file'],
                    record['unique_key'],
                    question_id,
                )
            )

            conn.commit()
            conn.close()
            return JSONResponse({'success': True, 'message': '大题更新成功'})

        @web_app.post('/api/question-bank/questions/{question_id}/export-docx', response_class=JSONResponse)
        async def export_question_bank_question_docx(question_id: int, request: Request, session_id: Optional[str] = Cookie(None)):
            user = get_current_user(session_id)
            if not user:
                return JSONResponse({'success': False, 'message': '请先登录'}, status_code=401)

            body = await request.json()
            show_answer = bool(body.get('show_answer', False))

            ensure_question_bank_db()
            conn = sqlite3.connect(QUESTION_DB_PATH)
            cur = conn.cursor()
            row = cur.execute('SELECT content_format, payload_json FROM question_bank WHERE id = ?', (question_id,)).fetchone()
            conn.close()

            if not row:
                return JSONResponse({'success': False, 'message': '题目不存在'}, status_code=404)
            if _safe_str(row[0]) != SHORT_ANSWER_CONTENT_FORMAT:
                return JSONResponse({'success': False, 'message': '当前仅支持导出已保存的综合题'}, status_code=400)

            payload = safe_json_loads(row[1], {})
            if not isinstance(payload, dict) or not payload:
                return JSONResponse({'success': False, 'message': '题目内容缺失，无法导出'}, status_code=400)

            try:
                payload['show_answer'] = show_answer
                result = export_short_answer_payload_to_docx(payload)
                return JSONResponse(result)
            except Exception as e:
                import traceback
                traceback.print_exc()
                return JSONResponse({'success': False, 'message': f'导出失败: {str(e)}'}, status_code=500)

        # ============= RAG API 路由 =============
        
        @web_app.get("/api/", response_class=JSONResponse)
        async def read_api_home():
            """API 首页"""
            return JSONResponse({
                "status": "ok",
                "message": "试卷生成与知识问答系统 API",
                "modules": {
                    "paper_generation": {
                        "description": "试卷生成功能",
                        "endpoints": ["GET /", "POST /submit"]
                    },
                    "rag": {
                        "description": "RAG 知识问答功能",
                        "available": RAG_AVAILABLE,
                        "endpoints": ["POST /rag/create/", "POST /rag/chat/", "POST /rag/clear/", "GET /rag/status/", "GET /uploads/", "DELETE /uploads/{filename}"]
                    }
                }
            })
        
        if RAG_AVAILABLE:
            @web_app.post("/rag/chat/")
            async def chat(request: ChatRequest):
                """
                根据用户问题，从向量库检索并返回回答，并将生成的题目保存到CSV文件。
                """
                print(f"Q: {request.question}")
                try:
                    # 重新获取 vector_store 实例
                    from milvus_vector import get_vector_store
                    current_vector_store = get_vector_store()
                    
                    # 初始化 OpenAI Chat 模型
                    llm = ChatOpenAI(model=config.llm_model_name, api_key=config.api_key, base_url=config.base_url)
                    # 定义 Prompt 模板
                    qa_prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])

                    # 定义搜索参数
                    search_kwargs = {"score_threshold": 0.3, "k": 5}
                    retriever = current_vector_store.as_retriever(search_type="similarity_score_threshold", search_kwargs=search_kwargs)
                    qa_chain = RetrievalQA.from_chain_type(
                        llm=llm,
                        chain_type="stuff",
                        retriever=retriever,
                        chain_type_kwargs={"prompt": qa_prompt},
                        return_source_documents=True
                    )
                    result = qa_chain.invoke({"query": request.question})
                    answer_result = result.get("result", "")
                    print(f"A: {answer_result}")

                    # 解析大模型返回的题目
                    questions = parse_llm_response(answer_result)
                    
                    # 保存题目到CSV文件
                    if questions:
                        csv_filename = f"questions_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
                        saved_file_path = save_questions_to_csv(questions, csv_filename)
                        print(f"题目已保存到: {saved_file_path}")
                    else:
                        print("未解析到有效的题目数据")
                        saved_file_path = None

                    # 检查是否有有效的源文档
                    source_documents = result.get("source_documents", [])
                    if not source_documents:
                        print("警告：未找到相关的源文档")
                        if not answer_result or "no answer" in answer_result.lower() or "not available" in answer_result.lower():
                            answer_result = "抱歉，我目前没有找到与您问题相关的信息。请尝试换一种方式提问或咨询其他问题。"
                    
                    source = {"source_documents": [{"content": doc.page_content, "metadata": doc.metadata} for doc in source_documents]}
                    print(f"source: {source}")

                    return JSONResponse({
                        "status": "success",
                        "answer": answer_result,
                        "questions_count": len(questions),
                        "saved_file": saved_file_path
                    })
                except Exception as e:
                    print(f"错误: {str(e)}")
                    return JSONResponse({"status": "error", "message": str(e)})

            @web_app.post("/rag/clear/")
            async def clear_knowledge(req: ClearRequest):
                """
                清空 Milvus 知识库集合，并删除指定目录中的文件
                """
                collection_name = req.collection_name or config.milvus_collection_name
                host = req.host or config.milvus_host
                port = int(req.port or config.milvus_port)
                folder = "./upload_files"
                try:
                    connections.connect("default", host=host, port=port)
                    if utility.has_collection(collection_name):
                        collection = Collection(name=collection_name)
                        collection.drop()
                        print(f"Collection '{collection_name}' 成功删除.")
                    else:
                        print(f"Collection '{collection_name}' 不存在.")
                    connections.disconnect("default")

                    for filename in os.listdir(folder):
                        file_path = os.path.join(folder, filename)
                        if os.path.isfile(file_path) or os.path.islink(file_path):
                            os.unlink(file_path)
                        elif os.path.isdir(file_path):
                            shutil.rmtree(file_path)

                    return JSONResponse({"status": "success", "message": f"知识库清空: Collection '{collection_name}' 删除, 文件夹 '{folder}' 清空."})
                except Exception as e:
                    return JSONResponse({"status": "error", "message": f"知识库清空失败. 原因: {str(e)}"})

            @web_app.get("/rag/status/")
            async def kb_status(collection_name: str = None, host: str = None, port: int = None):
                """
                返回知识库的基本信息，例如集合是否存在和实体数量
                """
                collection_name = collection_name or config.milvus_collection_name
                host = host or config.milvus_host
                port = port or config.milvus_port
                try:
                    connections.connect("default", host=host, port=port)
                    if utility.has_collection(collection_name):
                        collection = Collection(name=collection_name)
                        num_entities = collection.num_entities
                        connections.disconnect("default")
                        return JSONResponse({"status": "success", "collection_exists": True, "num_entities": num_entities})
                    else:
                        connections.disconnect("default")
                        return JSONResponse({"status": "success", "collection_exists": False, "num_entities": 0})
                except Exception as e:
                    return JSONResponse({"status": "error", "message": str(e)})

            @web_app.post("/rag/create/")
            async def create_knowledge(file: UploadFile = File(...)):
                """
                上传文件到指定目录后，处理文件内容并添加到向量库。
                """
                folder = './upload_files'
                os.makedirs(folder, exist_ok=True)
                file_path = os.path.join(folder, file.filename)

                try:
                    print(f"开始处理文件: {file.filename}")
                    
                    # 保存文件
                    with open(file_path, "wb") as f:
                        content = await file.read()
                        f.write(content)
                    print(f"文件已保存到: {file_path}")

                    # 初始化文件处理器
                    file_processor = RagFileProcessor(chunk_size=64)
                    # 处理文件内容并插入到向量库
                    text_datas = file_processor.get_data(file_path=file_path)
                    print(f"文件处理完成，共 {len(text_datas['texts'])} 个文本块")

                    # 重新获取 vector_store 实例，确保 collection 存在
                    from milvus_vector import get_vector_store
                    current_vector_store = get_vector_store()
                    
                    # 连接到 Milvus 并添加向量
                    result = current_vector_store.add_texts(**text_datas)
                    print(f"向量已添加到 Milvus，返回结果: {result}")
                    
                    return JSONResponse({
                        "status": "success",
                        "message": f"文件 '{file.filename}' 上传成功并添加至向量数据库（{len(text_datas['texts'])} 个文本块）",
                    })
                except Exception as e:
                    import traceback
                    error_detail = traceback.format_exc()
                    print(f"上传文件失败: {error_detail}")
                    
                    # 如果失败，尝试删除已保存的文件
                    if os.path.exists(file_path):
                        try:
                            os.remove(file_path)
                            print(f"已清理失败的文件: {file_path}")
                        except:
                            pass
                    
                    return JSONResponse({
                        "status": "error",
                        "message": f"文件 '{file.filename}' 处理失败: {str(e)}",
                    })

            @web_app.get("/uploads/")
            async def list_uploads():
                """列出已上传的文件"""
                folder = './upload_files'
                os.makedirs(folder, exist_ok=True)
                files = os.listdir(folder)
                return JSONResponse({"status": "success", "files": files})

            @web_app.delete("/uploads/{filename}")
            async def delete_upload(filename: str):
                """删除已上传的文件及其在向量库中的数据"""
                folder = './upload_files'
                file_path = os.path.join(folder, filename)
                
                if not os.path.exists(file_path):
                    return JSONResponse({"status": "error", "message": f"文件 {filename} 不存在"}, status_code=404)
                
                try:
                    # 从 Milvus 中删除该文件对应的向量数据
                    # 使用文件名作为 metadata 过滤条件
                    collection_name = config.milvus_collection_name
                    connections.connect("default", host=config.milvus_host, port=config.milvus_port)
                    
                    if utility.has_collection(collection_name):
                        collection = Collection(name=collection_name)
                        collection.load()
                        
                        # 查询该文件对应的实体 ID
                        # 使用 source metadata 字段来查找
                        expr = f'source == "{filename}"'
                        try:
                            # 先查询出所有匹配的 ID
                            results = collection.query(
                                expr=expr,
                                output_fields=["pk"],
                                limit=10000
                            )
                            
                            if results:
                                # 提取所有 ID 并删除
                                ids_to_delete = [str(item["pk"]) for item in results]
                                if ids_to_delete:
                                    collection.delete(expr=f'pk in [{",".join(ids_to_delete)}]')
                                    print(f"从向量库删除了 {len(ids_to_delete)} 条记录")
                        except Exception as query_error:
                            print(f"查询向量时出错: {query_error}")
                            # 即使查询失败，仍然删除文件
                        
                        collection.release()
                    
                    connections.disconnect("default")
                    
                    # 删除本地文件
                    os.remove(file_path)
                    
                    return JSONResponse({
                        "status": "success",
                        "message": f"文件 {filename} 及其向量数据已删除"
                    })
                    
                except Exception as e:
                    print(f"删除文件时出错: {str(e)}")
                    # 尝试至少删除文件
                    try:
                        if os.path.exists(file_path):
                            os.remove(file_path)
                        return JSONResponse({
                            "status": "warning",
                            "message": f"文件已删除，但向量数据删除失败: {str(e)}"
                        })
                    except:
                        return JSONResponse({
                            "status": "error",
                            "message": f"删除失败: {str(e)}"
                        }, status_code=500)
        else:
            # RAG 功能不可用时的占位路由
            @web_app.get("/rag/status/")
            async def rag_unavailable():
                return JSONResponse({"status": "error", "message": "RAG 功能未启用，请检查依赖模块是否正确安装"})

        # 直接运行应用（不通过模块导入）。
        # 注意：reload=True 需要导入字符串（模块:应用），这在文件名以数字开头（例如 `1.py`）时不可能。
        # 禁用重新加载以避免出现 "必须将应用程序作为导入字符串传递" 错误。
        uvicorn.run(web_app, host='127.0.0.1', port=8000, reload=False)
    else:
        # 运行 CLI 生成
        main()