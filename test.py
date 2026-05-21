import argparse
import os
from collections import Counter, defaultdict
from docx import Document
from openpyxl import load_workbook

def parse_ratio_arg(s, keys):
    """解析比例字符串并归一化。
    支持两种格式："高:30,中:40,低:30" 或 "30,40,30"（顺序按 keys）。
    返回字典，值为 0..1 之间且总和为 1。
    """
    parts = [p.strip() for p in s.split(',') if p.strip()]
    d = {}
    if all(':' in p for p in parts):
        for p in parts:
            k, v = p.split(':', 1)
            d[k.strip()] = float(v)
    else:
        vals = [float(p) for p in parts]
        for k, v in zip(keys, vals):
            d[k] = v
    total = sum(d.values())
    if total == 0:
        return {k: 1.0 / len(keys) for k in keys}
    return {k: (d.get(k, 0.0) / total) for k in keys}


def compute_counts_from_ratios(n, ratios):
    """把比例分配为整数计数，保证总和为 n。

    使用最大余数法分配多余项，能较公平地将小数部分分配为整数。
    """
    keys = list(ratios.keys())
    raw = {k: ratios[k] * n for k in keys}
    counts = {k: int(raw[k]) for k in keys}
    remainder = n - sum(counts.values())
    frac = sorted(keys, key=lambda k: raw[k] - int(raw[k]), reverse=True)
    i = 0
    while remainder > 0:
        counts[frac[i % len(frac)]] += 1
        remainder -= 1
        i += 1
    return counts


def read_single_choice_xlsx(xlsx_path, allowed_kps=None):
    """从 Excel 中读取候选单选题。

    支持表头：目录、题目类型、大题题干、小题题型、小题题干、正确答案、答案解析、难易度、知识点、选项数、选项A..D
    只会收集小题题型中包含"单选"的小题。
    allowed_kps 非空时只保留知识点在该列表内的题目。
    返回候选题的字典列表，每项包含 'sn','level','kp','type'。
    """
    wb = load_workbook(xlsx_path, data_only=True)
    sheet = wb.worksheets[0]
    candidates = []
    # 读取表头并容错匹配列索引
    header_row = next(sheet.iter_rows(min_row=1, max_row=1, values_only=True))
    headers = [h.strip() if isinstance(h, str) else h for h in header_row]

    def find_idx(names):
        for name in names:
            if name in headers:
                return headers.index(name)
        return None

    idx_level = find_idx(['难易度', '难度', '难易'])
    idx_kp = find_idx(['知识点', '知识点标签'])
    idx_qtype = find_idx(['小题题型', '题目类型', '小题类型'])
    idx_big_stem = find_idx(['大题题干', '大题干'])
    idx_sn = find_idx(['目录', '序号', '编号'])

    # 兜底默认索引（当列名匹配失败时）
    header_len = len(headers)
    # 根据表头顺序，优先使用固定列索引（更稳健）：
    # 目录=0, 题目类型=1, ..., 难易度=7, 知识点=8
    if header_len >= 9:
        idx_sn = 0
        idx_qtype = 1
        idx_level = 7
        idx_kp = 8
        idx_big_stem = 2
    else:
        if idx_qtype is None and header_len > 1:
            idx_qtype = 1
        if idx_level is None and header_len > 7:
            idx_level = 7
        if idx_kp is None and header_len > 8:
            idx_kp = 8
        if idx_sn is None and header_len > 0:
            idx_sn = 0

    # 小题题干列通常在第4索引（第5列），选项从第10列开始（A=10,B=11,C=12,D=13）
    for row in sheet.iter_rows(min_row=2, values_only=True):
        if not any(row):
            continue
        level = row[idx_level] if idx_level is not None and idx_level < len(row) else None
        qtype = row[idx_qtype] if idx_qtype is not None and idx_qtype < len(row) else '单选'
        category = row[idx_kp] if idx_kp is not None and idx_kp < len(row) else None
        sn = row[idx_sn] if idx_sn is not None and idx_sn < len(row) else None
        # 小题题干位于第5列（索引4）；大题题干一般位于第3列（索引2）
        stem = row[4] if len(row) > 4 else ''
        big_stem = row[idx_big_stem] if idx_big_stem is not None and idx_big_stem < len(row) else ''
        answer = row[5] if len(row) > 5 else ''
        analysis = row[6] if len(row) > 6 else ''
        option_count = row[9] if len(row) > 9 else None
        # 读取选项 A-D（若存在）
        opts = []
        for oi in range(10, 14):
            if oi < len(row) and row[oi]:
                opts.append(str(row[oi]))
        # 仅保留小题题型包含"单选"的项
        if qtype is None:
            continue
        qtype_str = str(qtype).strip()
        if '单选' not in qtype_str:
            continue
        kp_str = str(category).strip() if category else ''
        if allowed_kps and kp_str not in allowed_kps:
            continue
        candidates.append({
            'sn': sn,
            'level': str(level).strip() if level else '',
            'kp': kp_str,
            'type': qtype_str,
            'stem': str(stem).strip() if stem else '',
            'big_stem': str(big_stem).strip() if big_stem else '',
            'options': opts,
            'answer': str(answer).strip() if answer else '',
            'analysis': str(analysis).strip() if analysis else '',
        })
    wb.close()
    return candidates


def find_selection_backtrack(candidates, N, target_diff_counts, target_kp_counts):
    # candidates: 列表，每项包含 'level' 和 'kp'
    M = len(candidates)
    # 预计算后缀可用量用于剪枝（从 i 到末尾各类别可用数量）
    diff_keys = list(target_diff_counts.keys())
    kp_keys = list(target_kp_counts.keys())

    # 为便于访问构建数组
    levels = [c['level'] for c in candidates]
    kps = [c['kp'] for c in candidates]

    # 后缀计数：suffix_diff[i] 表示从 i 到末尾各难度的可用数量
    suffix_diff = [defaultdict(int) for _ in range(M + 1)]
    suffix_kp = [defaultdict(int) for _ in range(M + 1)]
    for i in range(M - 1, -1, -1):
        suffix_diff[i] = suffix_diff[i + 1].copy()
        suffix_kp[i] = suffix_kp[i + 1].copy()
        suffix_diff[i][levels[i]] += 1
        suffix_kp[i][kps[i]] += 1

    target_total = N

    path = []

    # 当前已选计数
    sel_diff = defaultdict(int)
    sel_kp = defaultdict(int)

    found = [None]
    # 记录搜索过程中找到的最长可行前缀（即已选题目最多的路径），用于在无完整解时返回部分结果
    best_path = {'indices': [], 'length': 0}

    def can_still_satisfy(start, sel_diff, sel_kp, slots_left):
        # 检查从 start 到末尾的可用题目能否满足剩余需求
        for d, tgt in target_diff_counts.items():
            need = max(0, tgt - sel_diff.get(d, 0))
            if suffix_diff[start].get(d, 0) < need:
                return False
        for k, tgt in target_kp_counts.items():
            need = max(0, tgt - sel_kp.get(k, 0))
            if suffix_kp[start].get(k, 0) < need:
                return False
        # 还要保证总体可用数足够
        if sum(suffix_diff[start].values()) < slots_left:
            return False
        return True

    def dfs(start):
        # 递归深度优先搜索，尝试从 start 开始选择题目
        if found[0] is not None:
            return True
        if len(path) == target_total:
            # 校验是否完全满足目标计数
            for d, tgt in target_diff_counts.items():
                if sel_diff.get(d, 0) != tgt:
                    return False
            for k, tgt in target_kp_counts.items():
                if sel_kp.get(k, 0) != tgt:
                    return False
            found[0] = path.copy()
            return True
        # 更新当前已选的最佳前缀（尽量保存更多已选题）
        if len(path) > best_path['length']:
            best_path['indices'] = path.copy()
            best_path['length'] = len(path)
        if start >= M:
            return False

        slots_left = target_total - len(path)
        # 剪枝：如果后缀不能满足剩余需求则回溯
        if not can_still_satisfy(start, sel_diff, sel_kp, slots_left):
            return False

        # 依次尝试包含某个候选或跳过
        for i in range(start, M):
            lvl = levels[i]
            kp = kps[i]
            # 如果选了会超出某类目标，则跳过
            if sel_diff[lvl] + 1 > target_diff_counts.get(lvl, 0):
                continue
            if sel_kp[kp] + 1 > target_kp_counts.get(kp, 0):
                continue

            # 选择 i
            path.append(i)
            sel_diff[lvl] += 1
            sel_kp[kp] += 1

            if dfs(i + 1):
                return True

            # 回溯
            path.pop()
            sel_diff[lvl] -= 1
            sel_kp[kp] -= 1

        return False

    ok = dfs(0)
    if ok and found[0] is not None:
        return [candidates[i] for i in found[0]]
    # 若未找到完整解，返回在搜索过程中最长的已选题目序列（可能为空）
    if best_path['length'] > 0:
        return [candidates[i] for i in best_path['indices']]
    return None


def save_docx_real_paper(selected, out_path, show_difficulty=True, show_knowledge=True, show_answer=True, show_analysis=True):
    """保存抽题结果到 Word 文档，根据参数控制输出内容。"""
    doc = Document()
    doc.add_heading('单选题（自动抽题结果）', level=1)
    doc.add_paragraph('一、单选题')
    qnum = 1
    prev_big = None
    for c in selected:
        # 获取大题题干和小题题干
        big = c.get('big_stem', '')
        small_stem = c.get('stem', '')
        
        # 构建题号和大题题干行（如果大题题干存在且与上一题不同）
        if big and big != prev_big:
            question_line = f"{qnum}、 {big}"
            doc.add_paragraph(question_line)
            prev_big = big
        else:
            # 如果没有大题题干或与上一题相同，只写题号
            question_line = f"{qnum}、"
            doc.add_paragraph(question_line)
        
        # 构建小题题干行，包含难度和知识点信息
        stem_line = f"      {small_stem}"
        if show_difficulty or show_knowledge:
            details = []
            if show_difficulty:
                details.append(f"难度：{c.get('level','')}")
            if show_knowledge:
                details.append(f"知识点：{c.get('kp','')}")
            if details:
                stem_line += f"（{', '.join(details)}）"
        doc.add_paragraph(stem_line)

        # 选项列表
        opts = c.get('options', [])
        # 按 A,B,C... 写出选项
        for idx, opt in enumerate(opts):
            label = chr(ord('A') + idx)
            doc.add_paragraph(f"    {label}、{opt}")
        
        # 在题目后保留一行空白供答题
        doc.add_paragraph("")
        
        # 根据参数决定是否显示答案和解析
        if show_answer and c.get('answer'):
            doc.add_paragraph(f"    答案：{c.get('answer')}")
        if show_analysis and c.get('analysis'):
            doc.add_paragraph(f"    解析：{c.get('analysis')}")
        
        # 在题目之间增加空行
        doc.add_paragraph("")
        qnum += 1
        
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)


def main():
    parser = argparse.ArgumentParser(description='从单选题.xlsx 回溯抽题，尝试满足用户的占比要求')
    parser.add_argument('-n', '--number', type=int, default=10, help='题目总数（所有知识点题目数量之和）')
    parser.add_argument('--difficulty', type=str, default='难:10,中:20,易:70', help='难度占比')
    parser.add_argument('--types', type=str, default='单选:100', help='题型占比（该脚本仅处理单选题表）')
    parser.add_argument('--kps', type=str, default='32090102,32030303,32080102,32090201,32100401', help='知识点列表，用逗号分隔')
    parser.add_argument('--xlsx', type=str, default='./单选题.xlsx', help='单选题 xlsx 文件路径')
    parser.add_argument('--out', type=str, default='./output/from_xlsx.docx', help='输出 docx 路径')
    parser.add_argument('--show-difficulty', type=int, default=0, choices=[0, 1], help='是否显示难度 (1=显示, 0=不显示)')
    parser.add_argument('--show-knowledge', type=int, default=1, choices=[0, 1], help='是否显示知识点 (1=显示, 0=不显示)')
    parser.add_argument('--show-answer', type=int, default=1, choices=[0, 1], help='是否后显示答案 (1=显示, 0=不显示)')
    parser.add_argument('--show-analysis', type=int, default=1, choices=[0, 1], help='是否显示解析 (1=显示, 0=不显示)')

    args = parser.parse_args()

    diff_keys = ['难', '中', '易']
    diff_ratios = parse_ratio_arg(args.difficulty, diff_keys)
    target_diff_counts = compute_counts_from_ratios(args.number, diff_ratios)

    # 忽略题型参数，因为本脚本仅处理单选题；确保请求的是单选题
    # 知识点
    kps = [k.strip() for k in args.kps.split(',') if k.strip()]
    if not kps:
        # 稍后从工作簿收集知识点；暂时允许所有
        allowed_kps = None
    else:
        allowed_kps = kps

    candidates = read_single_choice_xlsx(args.xlsx, allowed_kps=allowed_kps)

    # 如果用户没有提供知识点，从候选题的知识点均匀创建目标知识点计数
    if not kps:
        all_kps = sorted({c['kp'] for c in candidates})
        if not all_kps:
            print('警告：题库中没有知识点信息或题目为空')
            target_kp_counts = {}
        else:
            # 当未指定知识点时，均匀分配题目总数到各个知识点
            ratio = {kp: 1.0 for kp in all_kps}
            target_kp_counts = compute_counts_from_ratios(args.number, ratio)
    else:
        # 当用户指定了知识点，按比例分配题目总数
        ratio = {kp: 1.0 for kp in kps}  # 均匀分配
        target_kp_counts = compute_counts_from_ratios(args.number, ratio)

    # 快速可行性检查
    total_available = len(candidates)
    if total_available < args.number:
        print(f"题库中可用题目 {total_available} 少于所需 {args.number}。无法满足。")
        # 报告按难度/知识点的可用数量
        avail_diff = Counter([c['level'] for c in candidates])
        avail_kp = Counter([c['kp'] for c in candidates])
        print('题库按难度可用：', dict(avail_diff))
        print('题库按知识点可用：', dict(avail_kp))
        return

    # 如果候选题不包含所需的难度类别，报告短缺
    avail_diff = Counter([c['level'] for c in candidates])
    shortages = {}
    for d, tgt in target_diff_counts.items():
        if avail_diff.get(d, 0) < tgt:
            shortages[d] = tgt - avail_diff.get(d, 0)

    avail_kp = Counter([c['kp'] for c in candidates])
    shortages_kp = {}
    for k, tgt in target_kp_counts.items():
        if avail_kp.get(k, 0) < tgt:
            shortages_kp[k] = tgt - avail_kp.get(k, 0)

    if shortages or shortages_kp:
        print('初步检测发现题库中缺少以下题目（按类别/知识点）：')
        if shortages:
            print('按难度缺少：', shortages)
        if shortages_kp:
            print('按知识点缺少：', shortages_kp)
        print('仍然会尝试回溯搜索可行解（若存在）……')

    print('候选题数：', total_available)
    print('目标难度计数：', target_diff_counts)
    print('目标知识点计数：', target_kp_counts)

    selection = find_selection_backtrack(candidates, args.number, target_diff_counts, target_kp_counts)

    if selection:
        print('找到可行的题目组合，数量：', len(selection))
        save_docx_real_paper(selection, args.out, 
                             show_difficulty=bool(args.show_difficulty),
                             show_knowledge=bool(args.show_knowledge),
                             show_answer=bool(args.show_answer),
                             show_analysis=bool(args.show_analysis))
        print('已保存到', args.out)
    else:
        # 未找到完整解时，使用贪心填充作为回退：按候选顺序依次加入不超目标的题目，直到不能再加入
        print('未找到满足全部约束的题目组合，尝试生成部分试卷（贪心回退）。')
        sel = []
        sel_diff = defaultdict(int)
        sel_kp = defaultdict(int)
        for c in candidates:
            if len(sel) >= args.number:
                break
            lvl = c['level']
            kp = c['kp']
            if sel_diff[lvl] + 1 > target_diff_counts.get(lvl, 0):
                continue
            if sel_kp[kp] + 1 > target_kp_counts.get(kp, 0):
                continue
            sel.append(c)
            sel_diff[lvl] += 1
            sel_kp[kp] += 1

        if sel:
            print(f'已生成部分试卷，共加入 {len(sel)} 道题（将保存并在文档中列出）。')
            save_docx_real_paper(sel, args.out,
                                 show_difficulty=bool(args.show_difficulty),
                                 show_knowledge=bool(args.show_knowledge),
                                 show_answer=bool(args.show_answer),
                                 show_analysis=bool(args.show_analysis))
            print('部分试卷已保存到', args.out)
        else:
            print('无法生成任何符合约束的题目（所有候选均会导致超出目标）。')

        # 报告题库中实际缺少项
        pool_diff = Counter([c['level'] for c in candidates])
        pool_kp = Counter([c['kp'] for c in candidates])
        need_report = {}
        for d, tgt in target_diff_counts.items():
            have = pool_diff.get(d, 0)
            if have < tgt:
                need_report.setdefault('difficulty', {})[d] = tgt - have
        for k, tgt in target_kp_counts.items():
            have = pool_kp.get(k, 0)
            if have < tgt:
                need_report.setdefault('knowledge', {})[k] = tgt - have
        if need_report:
            print('题库不足以满足的项（需补充数量）：', need_report)
        else:
            print('题库看似有足够题量，但无组合能同时满足所有约束（可能是交叉约束冲突）。')


if __name__ == '__main__':
    main()